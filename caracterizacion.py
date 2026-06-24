import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import os
import datetime
import re
import io
import unicodedata

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

# ==========================================
# 1. CONFIGURACIÓN E IDENTIDAD
# ==========================================
st.set_page_config(page_title="Análisis Institucional CFTe-AP", layout="wide", page_icon="📊")
st.title("🏛️ Dashboard Integral de Caracterización y Progresión")
st.markdown("---")

COLOR_BASE = '#E5E7E9'
COLOR_SI = '#A9DFBF'
COLOR_NO = '#F5B7B1'
COLORES_PASTEL = ['#AED6F1', '#A9DFBF', '#F9E79F', '#F5CBA7', '#D2B4DE', '#A3E4D7', '#F5B041', '#EB984E', '#CCD1D1']

MAPA_SINO = {'Sí': COLOR_SI, 'No': COLOR_NO}
MAPA_SEXO = {'M': '#AED6F1', 'F': '#FADBD8'}

MOTIVOS_COLORS = {
    "CAMBIO DE INSTITUCIÓN": "#AED6F1", "CAMBIO DE CARRERA": "#A9DFBF", "VOCACIONAL": "#F9E79F",
    "PROBLEMAS DE SALUD": "#F5B7B1", "PROBLEMAS FAMILIARES": "#D2B4DE", "TRASLADO": "#A3E4D7",
    "SITUACIÓN ECONÓMICA": "#F5CBA7", "JORNADA LABORAL": "#EB984E", "PROBLEMAS PERSONALES": "#CCD1D1"
}

CARRERAS_MAP = {
    "180": "TNS EN ASISTENCIA JURÍDICA",
    "184": "TNS EN PROYECTOS ELÉCTRICOS DE DISTRIBUCIÓN",
    "201": "TNS EN ENFERMERÍA CON MENCIÓN EN GERONTOLOGÍA",
    "202": "TNS EN INFORMÁTICA Y APLICACIONES TECNOLÓGICAS",
    "204": "TNS EN GESTIÓN DE COMERCIO EXTERIOR",
    "283": "TNS EN EDUCACIÓN PARVULARIA Y PRIMER - SEGUNDO AÑO DE EDUCACIÓN BÁSICA",
    "380": "TNS EN DEPORTES Y RECREACIÓN",
    "554": "TNS EN ADMINISTRACIÓN PÚBLICA",
    "559": "TNS EN AGRÍCOLA",
    "602": "TNS EN EDUCACIÓN ESPECIAL",
    "603": "TNS EN ADMINISTRACIÓN DE EMPRESAS",
    "608": "TNS EN TRABAJO SOCIAL",
    "609": "TNS EN GEOLOGÍA",
    "613": "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA",
    "615": "TNS EN FABRICACIÓN Y MONTAJE DE ESTRUCTURAS METÁLICAS",
    "187": "TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES",
    "206": "TNS EN GESTIÓN CONTABLE",
    "207": "TNS EN GESTIÓN DE RECURSOS HUMANOS",
    "612": "TNS EN LABORATORIO CLÍNICO, BANCO DE SANGRE E IMAGENOLOGÍA",
    "153": "TNS EN OBRAS CIVILES",
    "188": "TNS EN MANTENIMIENTO ELECTROMECANICO CON MENCION EN ELECTROMOVILIDAD",
    "208": "TNS EN ADMINISTRACION PUBLICA CON MENCION EN GESTION JURIDICA",
    "209": "TNS EN ESTETICA CON MENCION EN MASOTERAPIA",
    "210": "TNS EN VETERINARIA",
    "284": "TNS EN EDUCACION PARVULARIA PRIMER CICLO DE EDUCACION BASICA",
    "219": "TNS EN ESTÉTICA Y COSMETOLOGIA MENCIÓN MASOTERAPIA"
}

LISTA_MOTIVOS = {
    "CAMBIO DE INSTITUCIÓN": ["INGRESO A OTRA INSTITUCIÓN", "CAMBIO DE INSTITUCIÓN", "CAMBIO A OTRA INSTITUCIÓN", "INGRESO A OTRA IES", "INGRESO OTRA IES", "INGRESO A LAS FUERZAS ARMADAS"],
    "CAMBIO DE CARRERA": ["CAMBIO DE CARRERA", "CAMBIO DE CARRERA EN LA INSTITUCIÓN", "INGRESO OTRA CARRERA"],
    "VOCACIONAL": ["NO ME GUSTO LA CARRERA", "DESCONFORMIDAD CON LA CARRERA Y SU METODOLOGÍA", "DISCONFORMIDAD DE LA CARRERA", "INCONFORMIDAD CON LA CARRERA", "LA CARRERA NO CUMPLIÓ MIS EXPECTATIVAS", "LA CARRERA NO CUMPLE MIS EXPECTATIVAS", "VOCACIONAL", "VOCACIONAL INGRESO 2024", "PROBLEMA VOCACIONAL", "VOCACIONAL IGRESO 2024", "FALTA DE CARGA ACADÉMICA", "SIN CARGA ACADÉMICA"],
    "PROBLEMAS DE SALUD": ["PROBLEMAS DE SALUD", "SALUD MENTAL", "PROBLEMAS DE SALUD PERSONAL", "PROBLEMAS PÉRSONALES Y DE SALUD", "PROBLEMA DE SALUD", "PROBLAMAS DE SALUD", "PROBLEMAS DE SALUD MENTAL", "EMBARAZO", "EMBARAZO DE ALTO RIESGO"],
    "PROBLEMAS FAMILIARES": ["PROBLEMA DE SALUD DE FAMILIAR", "PROBLEMA DE SALUD FAMILIAR", "PROBLEMAS DE SALUD DE FAMILIAR", "PROBLEMAS FAMILIARES (SALUD)", "SALUD FAMILIAR", "CUIDADO DE HIJO O FAMILIAR", "PROBLEMAS FAMILIARES", "PROBLEMAS FAMILIAR", "FALTA RED DE APOYO FAMILIAR", "FALTA DE REDES DE APOYO FAMILIAR", "FALTA DE REDES DE APOYO FAMILIARES", "FALTA DE RED DE APOYO FAMILIAR", "FALTAS DE REDES DE APOYO FAMILIAR", "FALTA DE REDE DE APOYO FAMILIAR", "FALTA DE REDES DE APOYO FAMILIAR (HIJO)", "FALTA DE REDES DE APOYO FAMILIAR (HIJOS)", "CUIDADO DE HIJO/A", "FALLECIMIENTO DE HIJO", "CUIDADO DE HIJOS", "CUIDADO DE HIJO", "CUIDADO DE HIJA", "FALTA DE REDES DE APOYO CON LOS HIJOS"],
    "TRASLADO": ["TRASLADO", "TRASLADO DE CIUDAD", "TRASLADO DE CUIDAD", "TRASLADO DE FAMILIA", "TRASLDO DE CIUDAD", "CAMBIO DE CUIDAD"],
    "SITUACIÓN ECONÓMICA": ["CAMBIO DE SITUACIÓN ECONÓMICA", "CAMBIO DE SITUACION ECONOMICA", "CAMBUIO DE SITUACIÓN ECONÓMICA", "CAMBIO DE SITUACIÓN ECONJÓMICA", "SIN APOYO ECONÓMICO", "CAMBIO D SITUACIÓN ECONÓMICA", "CAMBIO DE SITUACIÓN ACONÓMICA"],
    "JORNADA LABORAL": ["JORNADA DE TRABAJO", "INCOMPATIBILIDAD DE HORARIOS CON EL TRABAJO", "HORARIO LABORAL", "JORNADA LABORAL", "PROBLEMAS LABORALES", "PROBLEMAS LABORAL", "JORNADA LABORL", "INCOMPATIBILIDAD JORNADA LABORAL V/S ESTUDIOS", "LABORAL", "CAMBIO DE JORNADA LABORAL", "JORMADA LABORAL"],
    "PROBLEMAS PERSONALES": ["PROBLEMAS PERSONALES", "MOTIVOS PERSONALES", "PERSONALES", "PROBLEMA PERSONAL", "PERSONAL", "SITUACIÓN PERSONAL", "MPTIVOS PERSONALES", "CAMBIO DE PLAN", "POR CAMBIO DE PLAN"]
}

def normalizar_texto(texto):
    if not texto or str(texto).upper() == "NONE": return ""
    texto = str(texto).upper().strip()
    return ''.join(c for c in unicodedata.normalize('NFD', texto) if unicodedata.category(c) != 'Mn')

MAPEO_EXACTO = {normalizar_texto(v): k for k, l in LISTA_MOTIVOS.items() for v in l}

def agrupar_motivo_retiro(motivo):
    if not motivo or str(motivo).strip() == "":
        return "No Especificado / Otro Formal"
    motivo_norm = normalizar_texto(motivo)
    if motivo_norm in MAPEO_EXACTO:
        return MAPEO_EXACTO[motivo_norm]
    for categoria, synonym_list in LISTA_MOTIVOS.items():
        if any(normalizar_texto(syn) in motivo_norm for syn in synonym_list):
            return categoria
    for categoria in LISTA_MOTIVOS.keys():
        if normalizar_texto(categoria) in motivo_norm:
            return categoria
    return "No Especificado / Otro Formal"

# Diccionario de Especialidades EM y Pertinencia
DICT_FAMILIAS_EM = {
    "Administración y Recursos Humanos": ["ADMINISTRACIÓN DE EMPRESAS", "ADMINISTRACIÓN", "ADMINISTRACION", "ADMINISTRACIÓN DE EMPRESA", "ADMINISTRATIVO", "TECNICO EN ADMINISTRACION", "ADM. EMPRESA", "ADMINISTRACION DE EMPRSAS", "TEC. ADMINISTRACION DE EMPRESAS", "ADMINISTRACIÓN EMPRESA", "TECNICO EN ADMINISTRACION DE EMPRESAS", "TECNICO EN ADMINISTRACIÓN EMPRESA", "ADMINISTRACION DE EMPRESA", "ADMINISTRACION MENCION RECURSOS HUMANOS", "Tecnico en Administración de Empresas", "TECNICO ADMINISTRACION DE EMPRESAS Y RECURSOS HUMA", "ADMINISTRACIÓN MENCIÓN DE RECURSOS HUMANOS", "ADMINISTRACIÒN DE EMPRESA", "CONTROL DE GESTIÓN Y LOGISTICA", "MENCIÓN EN RRHH", "ADMINISTRACIÓN MENCIÓN RECURSOS HUMANOS", "ADMINITRACIÓN DE EMPRESA", "ADMINISTRACIÓN CON RECUSOS HUMANOS", "ADMINISTRACIÒN DE EMPRESA EN RECURSO HUMANO", "ADMINISTRACIÓN CON MENCIÓN EN RRHH", "ADMINISTRACIÓN EN RRHH", "ADMINISTRACIÓN DE EMPRESA MENCION EN RRHHH", "ADMINISTRACIÓN DE EMPRESAS Y COMERCIO", "ADMINISTRACION DE EMPRESAS MENCION LOGISTICA", "ADMNISTRACION DE EMPRESAS", "TECNICO ADMINISTRATIVO", "ADMISNISTRACION DE EMPRESAS", "TÉCNICO ADMINISTRACIÓN", "ADM. EMPRESAS", "ADMINISTRA DE EMPRESAS", "ADMINISTRACION DE EMPRESAS CON MENCION LOGISTICA", "ADMINISTRACION CON MENCION LOGISTICA", "ADMINISTRACION MENCION LOGISTICA", "TÉCNICO ADMINISTRACIÓN CON MENCIÓN EN RR.HH.", "ADMINISTRACION DE EMPRESAS CON MENCION EN RECURSOS", "ADMINISTRACIÓN CON MENCIÓN EN RECURSOS HUMANOS", "ADMIRACIÓN DE EMPRESAS", "Administración de empresas Mencion RRHH", "TÉCNICO EN ADMINISTRACIÓN DE EMPRESAS", "TECNICO ADMISNISTRATIVO", "ADMINISTRACIÓN DE EMPRESA CON MENCION EN RRHH", "ADINITRACIÓN DE EMPRESA", "ADMNISTRACION", "ADM EMPRESA", "LOGISTICA"],
    "Secretariado": ["SECRETARIADO ADMINISTRATIVO", "SECRETARIADO EJECUTIVO", "SECRETARIADO", "SECRETARIADO EJECUTIVO CON MENCION EN COMPUTACIÓN", "TECNICO EN SECRETARIADO", "SECREATARIADO ADMINISTRATIVO", "SECRETARIA EJECUTIVA", "SECRETARIADO COMERCIAL", "SECRETARIO ADMINISTRATIVO", "SECRETARIADO ADMINISTRATIVO MENCIÓN COMPUTACIÓN", "Secretariado Administrativo C/M en Computación", "SECRETARIADO NIVEL MEDIO", "SECRETARIADP", "SECRETARIADO BILINGUE", "SECRETARIADA", "TECNICO SECRETARIADO", "SECRETARIADO ADMINISTRATIVA"],
    "Contabilidad": ["TECNICO EN CONTABILIDAD", "CONTABILIDAD", "TECNICO CONTADOR", "CONTADOR", "CONTABILIDAD TECNICA", "CONTADOR DE NIVEL MEDIO", "TÉCNICO CONTABLE", "CONTABILIDAD GENERAL", "ADMINISTRATIVO CONTABLE", "AYUDANTE CONTABLE"],
    "Ventas y Comercio": ["MARCKETING", "MARKETING", "VENTA Y PUBLICIDAD", "TECNICO EN VENTA", "VENTAS", "VENTAS Y PUBLICIDAD", "MARKETING Y CONTROL DE EXISTENCIAS", "VENTAS Y SERVICIO AL CLIENTE", "VENTAS Y PUBLICACIDAD", "VENTA", "VENTAS DUAL", "TEC VENTA Y PUBLICIDAD", "VENTA Y PUBLICIDAD CON MENCION EN ADMINISTRACION", "VENTAS Y PUBLICIDAD CON MENCIÓN ADMINISTRATIVA", "TECNICO EN VENTAS", "TÉCNICO DE NIVEL MEDIO EN VENTAS"],
    "Educación de Párvulos": ["ATENCION DE PARVULO", "ATENCIÓN DE PÁRVULO", "PARVULO", "TECNICO EN PARVULARIO", "PÁRVULO", "TECNICO EN ATENCION DE PARVULOS", "ATENCIÓN A PÁRVULO", "PARVULOS", "EDUCACION PARVULARIA", "EDUCACION APRVULARIA", "TECNICO PARVULO", "ATENCIÓN DE PÁRVULOS", "TECNCO EN PARVULO", "TECNICO EN ATENCIÓN DE PARVULO", "PARVULARIA", "ATENCIÓN A PARVULO", "ATENCION DE PARVULOS", "TÉCNICO EN ATENCIÓN DE PÁRVULO", "EDUCADORA DE PARVULO", "ATENCIÓN DE PARVURO", "TÉCNICO ATENCIÓN DE PARVURO", "TÉCNICO EN PÁRVULO", "EMTP EN EDCUACIÓN PARVULARIA", "TENCINO EN EDUCACION PARVULARIA}", "TENCINO EN EDUCACION PARVULARIA", "TECNICO EN PARVULOS", "ED. PARVULARIA", "PARBULO", "TÉCNICO EN PARVULO", "ASISTENTE DE PARVULO", "TECNICO DE PARVULOS EN DIFERENCIAL", "TECNICO N ATENCION PARVULARIA", "TECNICO EN ATENCION DE PARVULO", "TÉCNICO EN PARVULOS", "EN ATENCIÓN DE PARVULOS", "TÉCNICO EN EDUCACIÓN PARVULARIA", "TÉCNICO EN PÁRVULOS", "ATENVION PARVULARIA", "TÉCNICO PARBULO", "TECNICO EN ATENCIÓN DE PÁRVULOS", "PARVULARIO", "MANIPULACION DE ALIMENTOS Y DE EDUCACION PARV", "EDUCACIÓN PARVULARIA", "ATENCION EN PARVULO", "ATENCION DE PRAVULO"],
    "Enfermería y Cuidado de la Salud": ["TECNICO EN ENFERMERIA", "ENFERMERIA", "ENFERMERÍA", "ATENCION DE ENFERMERIA CON MENCION EN ADULTO MAYOR", "ATENCIÓN DE ENFERMERÍA", "TECNICO ENFERMERIA", "ENFERMERIA CON MENCIÓN GERONTOLOGIA", "ATENCIÓN DE ENFERMERÍA, MENCIÓN ADULTOS MAYORES.", "ENFRMERIA, MENCION ADULTO MAYOR", "ENFEREMERIA", "ENFERMERIA TECNICA", "ENFERMERIA NIVEL MEDIO", "TÉCNICO EN ENFERMERIA", "ENFERMERÍA NIVEL MEDIO", "ENFERMERÍA CON MENCIÓN EN ADULTO MAYOR", "ENFERMERIA CON MENCION EN ADULTO MAYOR", "T. ENFERMERÍA", "AUXILIAR DE ENFERMERÍA ADULTO MAYOR", "ATENCIÓN EN ENFERMERÍA", "TE. EN ENFERMERÍA NIVEL MEDIO", "AUXILIAR ENFERMERIA", "TÉCNICO EN ENFERMERÍA", "ENFERMERÍA MENCIÓN GERONTOLOGÍA", "ATENCION DE ENFERMERIA", "TECNICO EN ENFERMERÍA", "TECNICO NIVEL MEDIO EN ENFERMERIA", "TNS EN ENFERMERÍA", "ATENCION EN ENFERMERIA EN ADULTOS MAYORES", "TECNICO EN ENFERMERIA NIVEL MEDIO", "ENFERMERIA CON GERONTOLOGIA", "ATENCION AL ADULTO MAYOR", "ATENCION A ADULTO MAYOR", "ADULTO MAYOR", "PARAMEDICO", "ASISTENTE LABORATORIO", "ENEFERMERIA", "ENEFERMERIA MENCION ADULTO MAYOR"],
    "Programación e Informática": ["Computación", "EN COMPUTACION", "COMPUTACION", "COMPUTACION APLICADA A LA EMPRESA", "PROGRAMACIÓN DE COMPUTACIÓN", "MANTENCION Y APLICACIONES COMPUTACIONALES", "PROGRAMACION EN COMPUTACION", "COMPUTACIÓN APLICADA A LA EMPRESA", "PROGRAMACIÓN EN COMPUTACIÓN", "COMPUTACIÓN MENCION EN SISTEMAS", "PROGRAMADORA EN COMPUTACIÓN", "MANTENIMIENTO DE EQUIPO COMPUTACIONAL", "PROGRAMADOR EN COMPUTACIÓN", "ADMINISTRACION DE HADWARE Y APLICACIONES COMPU", "TECNICO PROGRAMADOR EN COMPUTACION", "TECNICO EN PROGRAMACION", "TECNICO PROGRAMADOR", "PROGRAMACIÓN", "PROGRAMACION", "PROGRAMADOR", "PROGAMACION", "MANTENCIÓN DE HARDWARE", "INFORMATICA"],
    "Telecomunicaciones y Conectividad": ["CONECTIVIDAD Y REDES", "TECNICO EN CONECTIVIDAD DE REDES", "CONECTIVIDAD EN REDES", "CONECTIVIDAD EB REDES", "TELECOMUNICACION", "TELECOMUNICACIONES"],
    "Electricidad": ["ELECTRICIDAD", "ELECTRICO", "ELECTRICIDAD INDUSTRIAL", "ELECTRCIDAD", "Tecnico de Nivel Medio en Electricidad", "TECNICO EN ELECTRICIDAD", "ELECTRICIDAS", "TECNIC ELECTRICO", "TECNICO ELECTRICO", "TECN. ELECTRICO", "ELECTRICIDADA INDUSTRIAL", "PROYECTO ELECTRICO", "ELECTRISIDAD", "ELECTRICISTA", "TÉCNICO DE NIVEL MEDIO EN ELECTRICIDAD", "ELÉCTRICO", "ELCTRICIDAD", "ELÑECTRICIDAD INDUSTRIAL", "TÉCNICO ELÉCTRICIDAD", "ELECTRICIDAD Y ELECTRÓNICA"],
    "Electrónica": ["ELECTRONICA INDUSTRIAL", "ELECTRONICA", "TECNICO ELECTRONICO", "ELECTRÓNICA", "Tecnico Electrónico", "TNS EN ELECTRONICA", "ELECTRONICO", "TECNICO EN ELECTRONICA", "ELECCTRÓNICA", "ELETRONICA Y ELECTRICIDAD", "TÉCNICO EN NIVEL MEDIO EN ELECTRÓNICA", "TÉCNICO DE MANDO MEDIO EN ELECTRÓNICA", "TECNICO DE NIVEL MEDIO EN ELECTRONICA", "TECNICO ELECTRONICA", "ELECTRONICA Y ELECTRICO", "TÉCNICO EN ELÉCTRONICA", "TECNICO EN ELECTRONICA Y ELECTRICIDAD", "TÉCNICO ELECTRÓNICO", "ELCETRONICA", "ELOCTRONICA", "ELÉCTRONICA"],
    "Electromecánica": ["ELECTROMECANICA", "ELECTRO MECANICA", "ELECTROMECANICO", "TECNICO ELECTROMECANICO DE COMANDO MEDIO", "ELECTROMECÁNICA"],
    "Mecánica Automotriz": ["MECANICA AUTOMOTRIZ", "MECÁNICA AUTOMOTRIZ", "TECNICO MECANICO AUTOMOTRIS", "MECANICO AUTOMOTRIZ", "MECÁNICA", "MECANICA AUTOMOTRIS", "MECANICO", "MECÁNICO AUTOMOTRIZ", "MECANICA DE COMBUSTION", "MECANICO AUTOMATRIZ", "MECANICOAUTOMITRIZ", "MECANICO AUTOMOTRIS", "MENCANICA AUTOMOTRIZ", "TÉCNICO AUTOMOTRIZ"],
    "Mecánica Industrial": ["TECNICO EN MAQUINA HERRAMIENTAS", "MECANICA INDUSTRIAL", "MECANICA INDUSTRTIAL", "MECANICO INDUSTRIAL", "MECANICA INSDUSTRIAL", "MECANICA", "MAQUINA Y HERRAMIENTAS"],
    "Construcciones Metálicas": ["CONSTRUCCIONES METALICAS", "ESTRUCTURAS METALICAS", "TECNICO EN CONSTRUCCIONES METALICAS", "CONTRUCCIONES METALICAS", "ESTRUCTURA METÁLICAS", "CONTRUCCIONES METALICA", "CONSTRUCCIONES METÁLICAS", "ESTRUCTURA METALICA", "CONSTRUCCIÓN METÁLICA", "ESTRUCTURA METÁLICA", "CONSTRUCCION Y METALICA", "SOLDADURA", "CONTRUCCION METALICA", "CONSTRUCCIÓN METALICA", "ESTRUCTURAS METÁLICAS"],
    "Gastronomía y Alimentación": ["ALIMENTACION", "HOTELERIA MENCION GASTRONOMIA", "ALIMENTACIÓN", "COCINA INTERNACIONAL", "ALIMENTACION COLECTIVA", "ALIMENTACIÓN COLECTIVA", "GASTRONOMIA", "TECNICO EN ALIMENTACION COLECTIVA", "TECNICO EN GASTRONOMIA", "GASTRONOMIA MENCION COCINA", "MANIPULACIÓN DE ALIMENTOS", "ALIMENTACIÒN COLECTIVA", "MANIPULADOR DE ALIMENTO", "SERVICIOS DE ALIMENTACIÓN COLECTIVA", "SERVICIOS DE ALIMENTACION COLECTIVA", "GASTRONOMÍA", "COCINA", "TECNICO EN ALIMENTACIÒN", "TECNICO EN ALIMENTACION", "MANIPULACIÓN DE ALIMENTO", "MANIPULACION DE ALIMENTOS", "GASTRONOMIA CONMENCION COCINA", "GASTRONOMIA MENCION EN PASTELERIA", "ALIMENTACION Y DIETETICA", "MANIPULADOR DE ALIMENTOS", "GASTRONOMÍA MENCIÓN EN COCINA", "TECNICO ALIMENTACION COLECTIVA", "MANIPULADORA DE ALIMENTO", "MANUPULACIÓN DE ALIMENTOS", "GASTRONOMIA MENCIÓN COCINA"],
    "Hotelería y Turismo": ["HOTELERIA", "TURISMO Y ECONOMÍA", "SERVICIO AL TURISMO", "TURISMO", "HOTELERIA Y TURISMO", "TECNICO EN TURISMO", "SERVICIO HOTELERO", "SERVICIO DE TURISMO", "HOTELERIA CON MENCION EN GASTRONOMIA", "TURISMO Y HOSTELERIA", "TECNICO HOTELERÍA Y TURISMO", "TURISMO Y HOTELERIA"],
    "Vestuario y Confección Textil": ["TEJIDO INDUSTRIAL", "DISEÑO DE VESTUARIO", "CORTE Y CONFECCION", "DISEÑO Y CONFECCION TEXTIL", "MODA CON MENCION EN ALTA COSTURA", "CORTE Y CONFECCIÓN", "TEJIDOS", "CORTE Y CONFECCIÓN VESTUARIO", "TÉCNICO EN VESTUARIO", "VESTUARIO", "DISEÑO DE VESTUARIO CON MENCIÓN EN CONFECCION INDU", "CORTE Y CONFECCION CON MENCION EN DISEÑO", "DISEÑO Y VESTUARIO", "CORTES Y CONFECCION", "COSTURA", "CORTE Y CONFECCION EN MODELAJE", "CORTE Y CONFECCION TEXTIL", "CORTE CONFECCIONES", "TECNICO EN VESTUARIO"],
    "Agropecuaria y Forestal": ["AGROPECUARIA", "AGRICOLA", "TECN AGRICOLA", "AGROPECUARIO", "AGRONOMIA", "TECNICO AGRICOLA", "TECNICO AGROPECUARIO", "TECNICO AGROPECUARIA", "TECNICO AGRCILA", "AGRONOMÍA", "AGRONOMO", "TÉCNICO AGRÍCOLA", "AGRICOLA AGRPECUARIO", "AGUICULTURA", "TÉCNICO PROFESIONAL AGROPECUARIO", "TECN AGRICOLA", "TEC AGRICOLA", "AGRÍCOLA", "TECNICO AGRICOLA NIVEL MEDIO", "AGRICOLA CON ESPECIALIDAD AGROPECUARIA", "TÉCNICO AGROPECUARIO", "FORESTAL"],
    "Minería y Geología": ["TECNICO METALURGICA EXTRACTIVA", "EXPLOTACION MINERA", "MINERIA", "EXPLOTACIÓN MINERA", "TECNICO EXPLOTACIÓN MINERA", "METALURGIA EXTRACTIVA", "METALURGICA EXTRACTIVA", "GEOLOGIA", "ASISTENTE DE GEOLOGIA", "EXPLOTACION MINERIA", "GEOLOGÍA", "TECNICO EN MINERIA", "ASISTENTE EN GOELOGIA NIVEL MEDIO", "ASISTENTE EN GEOLOGÍA", "ASISTENTE GEOLOGÍA", "EXPLOTACIÓN EN MINERIA", "ASISTENTE EN GEOLOGIA", "MINERÍA", "EXTLOTACION MINERA", "Técnico en explotación minera"],
    "Operaciones Marítimas y Portuarias": ["OPERACION PORTUARIA", "TECN. EN OPERACIONES PORTUARIAS", "TECNICO EN OPERACIONES PORTUARIAS", "Operación Portuaria", "OPERACIONES PORTUARIA", "TECNICO PORTUARIO", "OPERACION PORTUARIAS", "OPERACIONES PORTUARIAS", "PORTUARIA", "PROCESOS PORTUARIOS", "OPERADOR PORTUARIO", "TECNICO EN OPERACIONES PORTUARIA", "OPERACIONES PORRTUARIA", "TECNICO PESQUERO", "MENICA NAVAL"],
    "Aduana y Comercio Exterior": ["COMERCIO EXTERIOR", "ADMINISTRACIÓN DE GESTIÓN ADUANERA", "TÉCNICO EN COMERCIO EXTERIOR", "ADUANERO", "GESTIÓN ADUANERA", "ADMINISTRACIÓN ADUANERA"],
    "Construcción y Edificación": ["DIBUJO TECNICO", "DIBUJO TÉCNICO", "CONSTRUCCIÓN", "CONTRUCCION HABITACIONAL", "TECNICO EN CONSTRUCCION", "EDIFICACION", "CARPINTERIA", "CARPINTERO", "GASFITERÍA"],
    "Trabajo Social y Recreación": ["ATENCIÓN SOCIAL", "ATENCIÓN SOCIAL Y RECREATIVA", "ATENCION SOCIAL Y RECREATIVA", "TÉCNICO EN DEPORTE Y RECREACIÓN Y TRABAJO SOCIAL", "DEPORTE Y RECREACIÓN", "DEPORTE Y RECREACION", "DEPORTES Y RECREACIÒN", "DEPORTES Y RECREACION"]
}

PERTINENCIA_CARRERAS = {
    "Administración y Recursos Humanos": {"TNS EN ASISTENCIA JURÍDICA": "Media", "TNS EN GESTIÓN DE COMERCIO EXTERIOR": "Media", "TNS EN ADMINISTRACIÓN PÚBLICA": "Alta", "TNS EN ADMINISTRACIÓN DE EMPRESAS": "Alta", "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA": "Alta", "TNS EN GESTIÓN CONTABLE": "Media", "TNS EN GESTIÓN DE RECURSOS HUMANOS": "Alta", "TNS EN ADMINISTRACION PUBLICA CON MENCION EN GESTION JURIDICA": "Alta"},
    "Secretariado": {"TNS EN ASISTENCIA JURÍDICA": "Media", "TNS EN ADMINISTRACIÓN PÚBLICA": "Media", "TNS EN ADMINISTRACIÓN DE EMPRESAS": "Media", "TNS EN GESTIÓN DE RECURSOS HUMANOS": "Media", "TNS EN ADMINISTRACION PUBLICA CON MENCION EN GESTION JURIDICA": "Media"},
    "Contabilidad": {"TNS EN ADMINISTRACIÓN PÚBLICA": "Media", "TNS EN ADMINISTRACIÓN DE EMPRESAS": "Media", "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA": "Alta", "TNS EN GESTIÓN CONTABLE": "Alta", "TNS EN ADMINISTRACION PUBLICA CON MENCION EN GESTION JURIDICA": "Media"},
    "Ventas y Comercio": {"TNS EN GESTIÓN DE COMERCIO EXTERIOR": "Media", "TNS EN ADMINISTRACIÓN DE EMPRESAS": "Media", "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA": "Media"},
    "Educación de Párvulos": {"TNS EN EDUCACIÓN PARVULARIA Y PRIMER - SEGUNDO AÑO DE EDUCACIÓN BÁSICA": "Alta", "TNS EN EDUCACIÓN ESPECIAL": "Media", "TNS EN EDUCACION PARVULARIA PRIMER CICLO DE EDUCACION BASICA": "Alta"},
    "Enfermería y Cuidado de la Salud": {"TNS EN ENFERMERÍA CON MENCIÓN EN GERONTOLOGÍA": "Alta", "TNS EN LABORATORIO CLÍNICO, BANCO DE SANGRE E IMAGENOLOGÍA": "Media", "TNS EN ESTETICA CON MENCION EN MASOTERAPIA": "Media", "TNS EN ESTÉTICA Y COSMETOLOGIA MENCIÓN MASOTERAPIA": "Media"},
    "Programación e Informática": {"TNS EN INFORMÁTICA Y APLICACIONES TECNOLÓGICAS": "Alta"},
    "Telecomunicaciones y Conectividad": {"TNS EN INFORMÁTICA Y APLICACIONES TECNOLÓGICAS": "Alta"},
    "Electricidad": {"TNS EN PROYECTOS ELÉCTRICOS DE DISTRIBUCIÓN": "Alta", "TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES": "Media", "TNS EN MANTENIMIENTO ELECTROMECANICO CON MENCION EN ELECTROMOVILIDAD": "Alta"},
    "Electrónica": {"TNS EN PROYECTOS ELÉCTRICOS DE DISTRIBUCIÓN": "Media", "TNS EN INFORMÁTICA Y APLICACIONES TECNOLÓGICAS": "Media", "TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES": "Media", "TNS EN MANTENIMIENTO ELECTROMECANICO CON MENCION EN ELECTROMOVILIDAD": "Media"},
    "Electromecánica": {"TNS EN PROYECTOS ELÉCTRICOS DE DISTRIBUCIÓN": "Media", "TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES": "Alta", "TNS EN MANTENIMIENTO ELECTROMECANICO CON MENCION EN ELECTROMOVILIDAD": "Alta"},
    "Mecánica Automotriz": {"TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES": "Alta", "TNS EN MANTENIMIENTO ELECTROMECANICO CON MENCION EN ELECTROMOVILIDAD": "Alta"},
    "Mecánica Industrial": {"TNS EN FABRICACIÓN Y MONTAJE DE ESTRUCTURAS METÁLICAS": "Media", "TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES": "Media", "TNS EN OBRAS CIVILES": "Media"},
    "Construcciones Metálicas": {"TNS EN FABRICACIÓN Y MONTAJE DE ESTRUCTURAS METÁLICAS": "Alta"},
    "Agropecuaria y Forestal": {"TNS EN AGRÍCOLA": "Alta", "TNS EN VETERINARIA": "Alta"},
    "Minería y Geología": {"TNS EN GEOLOGÍA": "Alta"},
    "Operaciones Marítimas y Portuarias": {"TNS EN GESTIÓN DE COMERCIO EXTERIOR": "Media", "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA": "Media"},
    "Aduana y Comercio Exterior": {"TNS EN GESTIÓN DE COMERCIO EXTERIOR": "Alta", "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA": "Media"},
    "Construcción y Edificación": {"TNS EN OBRAS CIVILES": "Alta"},
    "Trabajo Social y Recreación": {"TNS EN DEPORTES Y RECREACIÓN": "Alta", "TNS EN EDUCACIÓN ESPECIAL": "Media", "TNS EN TRABAJO SOCIAL": "Alta"}
}

# ==========================================
# 2. RUTAS DE ARCHIVOS
# ==========================================
BASE_DIR_DEFAULT = "C:/Users/Usuario/Desktop/Sistema_Analisis_Institucional/2_Datos_Procesados"
RUTA_INTENTOS_DOC = [
    "Sistema_Analisis_Institucional/2_Datos_Procesados/2_Datos_Procesados_Docencia",
    "2_Datos_Procesados/2_Datos_Procesados_Docencia",
    os.path.join(BASE_DIR_DEFAULT, "2_Datos_Procesados_Docencia")
]
RUTA_INTENTOS_VIN = [
    "Sistema_Analisis_Institucional/2_Datos_Procesados/2_Datos_Procesados_Vinculacion",
    "2_Datos_Procesados/2_Datos_Procesados_Vinculacion",
    os.path.join(BASE_DIR_DEFAULT, "2_Datos_Procesados_Vinculacion")
]

CARPETA_DOC = next((ruta for ruta in RUTA_INTENTOS_DOC if os.path.exists(ruta)), None)
CARPETA_VIN = next((ruta for ruta in RUTA_INTENTOS_VIN if os.path.exists(ruta)), None)

if not CARPETA_DOC:
    st.error("❌ No se pudo localizar la carpeta de Docencia. Verifica el directorio.")
    st.stop()

# ==========================================
# 3. FUNCIONES DE LIMPIEZA Y LÓGICA
# ==========================================
def clean_rut(rut_series): return rut_series.astype(str).str.replace(r'[\.\-]', '', regex=True).str.upper().str.strip()

def parse_retiro(row):
    textos = []
    # Prioridad a Situación
    if 'SITUACIÓN' in row.index and pd.notna(row['SITUACIÓN']): textos.append(str(row['SITUACIÓN']).upper().strip())
    elif 'SITUACION' in row.index and pd.notna(row['SITUACION']): textos.append(str(row['SITUACION']).upper().strip())
    
    if 'OBSERVACIÓN' in row.index and pd.notna(row['OBSERVACIÓN']): textos.append(str(row['OBSERVACIÓN']).upper().strip())
    elif 'OBSERVACION' in row.index and pd.notna(row['OBSERVACION']): textos.append(str(row['OBSERVACION']).upper().strip())

    exclusiones = ["ADMINISTRATIVO", "REINCORPORACIÓN", "REINCORPORACION", "EGRESADO", "TITULADO", "FALLECIMIENTO"]
    for text in textos:
        
        if any(exc in text for exc in exclusiones):
            return pd.Series([None, -1, None, None])
        
        # Regla: Solo Definitivo o Temporal
        if "RETIRO DEFINITIVO" in text or "RETIRO TEMPORAL" in text:
            cat = "RETIRO DEFINITIVO" if "RETIRO DEFINITIVO" in text else "RETIRO TEMPORAL"

            # 🔥 CORRECCIÓN 1: Regex sin espacios literales para que siempre atrape los paréntesis
            matches = re.findall(r'\((.*?)\)', text)
            if matches:
                inside = matches[-1]

                # Extraer año (ej. 2021)
                anio_match = re.search(r'(20\d{2})', inside)
                anio_retiro = int(anio_match.group(1)) if anio_match else -1

                # Extraer semestre (ej. 1 SEM o 2 SEM)
                sem_match = re.search(r'([12])\s*SEM', inside)
                sem_retiro = f"{sem_match.group(1)}° Semestre" if sem_match else "No especificado"
                
                # Extraer motivo usando el mismo criterio que en matriculados.py:
                # todo lo que queda después de la primera coma dentro de los paréntesis.
                motivo_raw = inside.split(',', 1)[1].strip() if ',' in inside else inside
                motivo_final = agrupar_motivo_retiro(motivo_raw)

                return pd.Series([cat, anio_retiro, sem_retiro, motivo_final])
            
    return pd.Series([None, -1, None, None])

def evaluar_pertinencia(row):
    esp_raw = normalizar_texto(str(row.get('DATOS ENSEÑANZA MEDIA ESPECIALIDAD', '')))
    car_raw = str(row.get('NOMBRE_CARRERA', '')).upper().strip()
    
    if esp_raw == "" or "SIN ESPECIALIDAD" in esp_raw: return 'Sin Especialidad Previa'
    
    familia_asignada = None
    for familia, keywords in DICT_FAMILIAS_EM.items():
        if any(normalizar_texto(kw) in esp_raw for kw in keywords):
            familia_asignada = familia
            break
            
    if familia_asignada and familia_asignada in PERTINENCIA_CARRERAS:
        alineacion = PERTINENCIA_CARRERAS[familia_asignada].get(car_raw)
        if alineacion == "Alta": return "Pertinencia Alta"
        elif alineacion == "Media": return "Pertinencia Media"
        
    return 'Sin Pertinencia'

@st.cache_data
def cargar_archivo(ruta):
    try:
        if ruta.endswith(".xlsx"): df = pd.read_excel(ruta)
        else:
            df = None
            for s in [';', ',']:
                for enc in ['utf-8-sig', 'latin-1']:
                    try:
                        df = pd.read_csv(ruta, sep=s, encoding=enc, on_bad_lines='skip', engine='python')
                        break
                    except: continue
                if df is not None: break
        if df is not None:
            df.columns = df.columns.str.strip().str.upper()
            if 'RUT' in df.columns: df['RUT'] = clean_rut(df['RUT'])
            col_codigo = next((c for c in ['CARRERA_COD', 'COD_CARRERA'] if c in df.columns), None)
            if col_codigo:
                df['COD_CARRERA'] = df[col_codigo].astype(str).str.extract(r'(\d{3})', expand=False)
                df['NOMBRE_CARRERA'] = df['COD_CARRERA'].map(CARRERAS_MAP).fillna(df.get('CARRERA_DSC', "CARRERA DESCONOCIDA"))
            else: df['NOMBRE_CARRERA'] = df.get('CARRERA_DSC', "CARRERA DESCONOCIDA")
            if 'JORNADA' in df.columns: df['JORNADA'] = df['JORNADA'].astype(str).str.strip().str.upper()
            return df
    except: return None

@st.cache_data
def cargar_nomina_titulados():
    if not CARPETA_VIN: return pd.DataFrame()
    ruta = os.path.join(CARPETA_VIN, "LIMPIO_Nomina_Titulados.xlsx")
    try:
        df = pd.read_excel(ruta)
        df.columns = df.columns.str.strip().str.upper()
        col_rut = next((c for c in df.columns if 'RUT' in c), None)
        if col_rut: df['RUT_TIT'] = clean_rut(df[col_rut])
        col_fecha = next((c for c in df.columns if 'TITULA' in c and 'FECHA' in c), 'TITULACIÓN FECHA')
        if col_fecha in df.columns: df['FECHA_TIT_REAL'] = pd.to_datetime(df[col_fecha], errors='coerce')
        if 'RUT_TIT' in df.columns: return df[['RUT_TIT', 'FECHA_TIT_REAL']].drop_duplicates(subset=['RUT_TIT'])
    except: pass
    return pd.DataFrame()

# ==========================================
# 4. CARGA MASIVA DE DATOS
# ==========================================
archivos_todos = os.listdir(CARPETA_DOC) if os.path.exists(CARPETA_DOC) else []
archivos_matriculados = [f for f in archivos_todos if f.startswith("LIMPIO_ING") and "MAT" in f]

lista_matriculados = []
for arc in archivos_matriculados:
    parts = arc.upper().replace("LIMPIO_ING", "").replace(".XLSX", "").replace(".CSV", "").split("_MAT")
    if len(parts) == 2:
        anio_ing = int(parts[0])
        if anio_ing >= 2021:
            df_temp = cargar_archivo(os.path.join(CARPETA_DOC, arc))
            if df_temp is not None:
                df_temp['ANIO_INGRESO'] = anio_ing
                df_temp['ANIO_MATRICULA'] = int(parts[1])
                lista_matriculados.append(df_temp)

df_matriculados = pd.concat(lista_matriculados, ignore_index=True) if lista_matriculados else pd.DataFrame()
if df_matriculados.empty:
    st.error("❌ No se encontraron archivos de matrícula válidos.")
    st.stop()

max_anio_registrado = df_matriculados['ANIO_MATRICULA'].max()

# ==========================================
# 5. FILTROS Y PROCESAMIENTO
# ==========================================
st.sidebar.header("🔍 Filtros de Análisis")
anios_disp = sorted(df_matriculados['ANIO_INGRESO'].dropna().unique())
opcion_anio = st.sidebar.multiselect("Año de Análisis (Cohorte)", ["Todos"] + anios_disp, default=["Todos"])
sel_anios = anios_disp if "Todos" in opcion_anio else sorted([int(a) for a in opcion_anio])

if not sel_anios: st.stop()

sel_nombres = st.sidebar.multiselect("Carrera", sorted(df_matriculados['NOMBRE_CARRERA'].dropna().unique()), default=sorted(df_matriculados['NOMBRE_CARRERA'].dropna().unique()))
sel_jornadas = st.sidebar.multiselect("Jornada", sorted(df_matriculados['JORNADA'].dropna().unique()), default=sorted(df_matriculados['JORNADA'].dropna().unique()))

df_filt = df_matriculados[df_matriculados['ANIO_INGRESO'].isin(sel_anios) & df_matriculados['NOMBRE_CARRERA'].isin(sel_nombres) & df_matriculados['JORNADA'].isin(sel_jornadas)].copy()
if df_filt.empty:
    st.warning("No hay datos para la selección.")
    st.stop()

# --- PROCESAMIENTO VARIABLES CARACTERIZACIÓN ---
df_filt[['TIPO_RETIRO', 'ANIO_RETIRO_EXTRACT', 'SEMESTRE_RETIRO', 'MOTIVO_LIMPIO']] = df_filt.apply(parse_retiro, axis=1)

df_filt['CONDICION_CUIDADOR'] = df_filt.get('HIJOS TIENE', '').apply(lambda x: 'Sí' if str(x).strip().upper() in ['SI', 'SÍ', '1', 'TRUE'] else 'No')
cat_trabajo = ['DEPENDIENTE', 'INDEPENDIENTE', 'OCACIONAL', 'OCASIONAL']
df_filt['CATEGORIA_OCUPACIONAL'] = df_filt.get('SITUACIÓN OCUPACIONAL CATEGORIA', pd.Series(['No Informa']*len(df_filt))).fillna('No Informa').astype(str).str.strip().str.upper()
df_filt['TRABAJA_REAL'] = df_filt['CATEGORIA_OCUPACIONAL'].apply(lambda x: 'Sí' if x in cat_trabajo else 'No')

df_filt['EDAD'] = pd.to_numeric(df_filt.get('EDAD', np.nan), errors='coerce')
df_filt['RANGO_ETARIO'] = pd.cut(df_filt['EDAD'], bins=[0, 19, 29, 39, 49, 59, 69, 79, 89, 99], labels=['<=19', '20-29', '30-39', '40-49', '50-59', '60-69', '70-79', '80-89', '90-99'])

df_filt['ANIO_EGRESO_NUM'] = pd.to_numeric(df_filt.get('DATOS ENSEÑANZA MEDIA AÑO EGRESO', np.nan), errors='coerce')
df_filt['BRECHA_EGRESO'] = df_filt['ANIO_INGRESO'] - df_filt['ANIO_EGRESO_NUM']
df_filt['BRECHA_CATEGORIA'] = pd.cut(df_filt['BRECHA_EGRESO'], bins=[-np.inf, 0, 2, 5, np.inf], labels=['Directo (0 años)', '1-2 años', '3-5 años', '> 5 años'])

df_filt['PROMEDIO_NEM_CALC'] = pd.to_numeric(df_filt.get('DATOS ENSEÑANZA MEDIA PROMEDIO NOTAS EM', '').astype(str).str.replace(',', '.'), errors='coerce')
df_filt['PERTINENCIA_ESPECIALIDAD'] = df_filt.apply(evaluar_pertinencia, axis=1)

for col, clean_col in [('FINANCIAMIENTO GRATUIDAD APROBADA', 'GRATUIDAD_VALIDADA'), ('FINANCIAMIENTO BECA', 'BECA_VALIDADA'), ('FINANCIAMIENTO CAE', 'CAE_VALIDADO')]:
    df_filt[clean_col] = df_filt.get(col, '').apply(lambda x: 'Sí' if str(x).strip().upper() not in ['NAN', 'NONE', 'NO', 'FALSE', '0', ''] else 'No')

def clasificar_rsh(val):
    try:
        v = float(val)
        if pd.isna(v): return "Sin Registro"
        if v <= 40: return "Tramo 1 0% - 40%"
        elif v <= 50: return "Tramo 2 41% - 50%"
        elif v <= 60: return "Tramo 3 51% - 60%"
        elif v <= 70: return "Tramo 4 61% - 70%"
        elif v <= 80: return "Tramo 5 71% - 80%"
        elif v <= 90: return "Tramo 6 81% - 90%"
        else: return "Tramo 7 91% - 100%"
    except: return "Sin Registro"
df_filt['TRAMO_RSH'] = df_filt.get('REGISTRO SOCIAL DE HOGARES %', pd.Series([np.nan]*len(df_filt))).apply(clasificar_rsh)

# UNIVERSO BASE (Corte de Ingreso)
df_mat_real = df_filt[df_filt['ANIO_INGRESO'] == df_filt['ANIO_MATRICULA']].copy()

# ==========================================
# 6. MOTOR LÓGICO DE TASAS Y DESERCIÓN
# ==========================================
df_mat_real['ES_RETENIDO'] = False
df_mat_real['ES_PROGRESADO'] = False
df_mat_real['ES_DESERTOR_1ER'] = False
df_mat_real['ES_DESERTOR_2DO'] = False
df_mat_real['ES_DESERTOR_3ER'] = False

# Banderas exclusivas para conteo visual en UI
df_mat_real['ES_DECLARADO_2DO'] = False
df_mat_real['ES_DECLARADO_3ER'] = False

df_mat_real['MOTIVO_DESERCION'] = None
df_mat_real['SEMESTRE_DESERCION'] = "No Aplica"
fugas_list = []
for anio in sel_anios:
    idx_cohorte = df_mat_real['ANIO_INGRESO'] == anio
    
    ruts_anio_2 = set(df_matriculados[(df_matriculados['ANIO_INGRESO'] == anio) & (df_matriculados['ANIO_MATRICULA'] == anio + 1)]['RUT'].dropna())
    ruts_anio_3 = set(df_matriculados[(df_matriculados['ANIO_INGRESO'] == anio) & (df_matriculados['ANIO_MATRICULA'] == anio + 2)]['RUT'].dropna())
    
    # 0. Retención y Progresión
    df_mat_real.loc[idx_cohorte, 'ES_RETENIDO'] = df_mat_real.loc[idx_cohorte, 'RUT'].isin(ruts_anio_2)
    df_mat_real.loc[idx_cohorte, 'ES_PROGRESADO'] = df_mat_real.loc[idx_cohorte, 'ES_RETENIDO'] & df_mat_real.loc[idx_cohorte, 'RUT'].isin(ruts_anio_3)
    
    # 1. Deserción 1er Año
    df_fugas_1er = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'] == anio) & (df_filt['TIPO_RETIRO'].notna()) & (df_filt['ANIO_RETIRO_EXTRACT'] == anio)]
    df_mat_real.loc[idx_cohorte, 'ES_DESERTOR_1ER'] = df_mat_real.loc[idx_cohorte, 'RUT'].isin(set(df_fugas_1er['RUT'].dropna()))
    
    # 2. Deserción 2do Año (Se ajusta la búsqueda a isin([anio, anio+1]) como en el código validado)
    df_fugas_2do = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'].isin([anio, anio + 1])) & (df_filt['TIPO_RETIRO'].notna()) & (df_filt['ANIO_RETIRO_EXTRACT'] == anio + 1)]
    ruts_fugas_2do = set(df_fugas_2do['RUT'].dropna())
    df_mat_real.loc[idx_cohorte, 'ES_DECLARADO_2DO'] = df_mat_real.loc[idx_cohorte, 'RUT'].isin(ruts_fugas_2do)
        
    df_mat_real.loc[idx_cohorte, 'ES_DESERTOR_2DO'] = ~df_mat_real.loc[idx_cohorte, 'RUT'].isin(ruts_anio_2)

    # 3. Deserción 3er Año (Se ajusta la búsqueda a isin([anio+1, anio+2]) como en el código validado)
    # 3. Deserción 3er Año (Alineado con el Código de Matriculados)
    df_fugas_3er = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'].isin([anio + 1, anio + 2])) & (df_filt['TIPO_RETIRO'].notna()) & (df_filt['ANIO_RETIRO_EXTRACT'] == anio + 2)]
    ruts_fugas_3er = set(df_fugas_3er['RUT'].dropna())
    df_mat_real.loc[idx_cohorte, 'ES_DECLARADO_3ER'] = df_mat_real.loc[idx_cohorte, 'RUT'].isin(ruts_fugas_3er)
    
    df_mat_real.loc[idx_cohorte, 'ES_DESERTOR_3ER'] = df_mat_real.loc[idx_cohorte, 'ES_RETENIDO'] & (~df_mat_real.loc[idx_cohorte, 'RUT'].isin(ruts_anio_3))
    
    # Almacenar fugas de todos los años para el mapeo categorizado
    if not df_fugas_1er.empty: fugas_list.append(df_fugas_1er)
    if not df_fugas_2do.empty: fugas_list.append(df_fugas_2do)
    if not df_fugas_3er.empty: fugas_list.append(df_fugas_3er)

# Consolidar el archivo de fugas
df_fugas_export = pd.concat(fugas_list, ignore_index=True).drop_duplicates(subset=['RUT', 'ANIO_RETIRO_EXTRACT']) if fugas_list else pd.DataFrame()

# Mapeamos motivos a TODOS los años (1er, 2do y 3er Año)
if not df_fugas_export.empty:
    for idx, row in df_mat_real.iterrows():
        anio_busqueda = None
        if row['ES_DESERTOR_1ER']:
            anio_busqueda = row['ANIO_INGRESO']
        elif row['ES_DESERTOR_2DO']:
            anio_busqueda = row['ANIO_INGRESO'] + 1
        elif row['ES_DESERTOR_3ER']:
            anio_busqueda = row['ANIO_INGRESO'] + 2
            
        if anio_busqueda:
            match = df_fugas_export[(df_fugas_export['RUT'] == row['RUT']) & (df_fugas_export['ANIO_RETIRO_EXTRACT'] == anio_busqueda)]
            if not match.empty:
                df_mat_real.at[idx, 'MOTIVO_DESERCION'] = match.iloc[0]['MOTIVO_LIMPIO']
                df_mat_real.at[idx, 'SEMESTRE_DESERCION'] = match.iloc[0]['SEMESTRE_RETIRO']

# Titulación Lógica (Cálculo directo sobre Matriculados base)
df_titulados = cargar_nomina_titulados()
if not df_titulados.empty:
    df_mat_real = df_mat_real.merge(df_titulados, left_on='RUT', right_on='RUT_TIT', how='left')
    df_mat_real['ES_TITULADO'] = df_mat_real['FECHA_TIT_REAL'].notna()
    meses_duracion = np.where(df_mat_real.get('COD_CARRERA', '') == '184', 6, 30)
    df_mat_real['FECHA_TERMINO_TEORICA'] = pd.to_datetime(df_mat_real['ANIO_INGRESO'].astype(str) + "-03-01") + pd.to_timedelta(meses_duracion * 30.44, unit='D')
    df_mat_real['TITULACION_CALCULABLE'] = pd.Timestamp.now() >= df_mat_real['FECHA_TERMINO_TEORICA']
    df_mat_real['AÑOS_TITULACION'] = (df_mat_real['FECHA_TIT_REAL'] - pd.to_datetime(df_mat_real['ANIO_INGRESO'].astype(str) + "-03-01")).dt.days / 365.25
    limite = np.where(df_mat_real.get('COD_CARRERA', '') == '184', 1.5, 3.5)
    df_mat_real['ES_TITULADO_OPORTUNO'] = df_mat_real['ES_TITULADO'] & (df_mat_real['AÑOS_TITULACION'] <= limite)
else:
    df_mat_real['ES_TITULADO'], df_mat_real['TITULACION_CALCULABLE'], df_mat_real['ES_TITULADO_OPORTUNO'] = False, False, False

DIMENSIONES_DEMO = {"Sexo": "SEXO", "Rango Etario": "RANGO_ETARIO", "Etnia": "PUEBLO ORIGINARIO ETNIA", "Padre/Cuidador": "CONDICION_CUIDADOR", "Trabaja": "TRABAJA_REAL"}
DIMENSIONES_EDU = {"Brecha de Egreso EM": "BRECHA_CATEGORIA", "Tipo Enseñanza": "DATOS ENSEÑANZA MEDIA TIPO ENSEÑANZA", "Régimen": "DATOS ENSEÑANZA MEDIA REGIMEN ENSEÑANZA", "Modalidad": "DATOS ENSEÑANZA MEDIA MODALIDAD"}
DIMENSIONES_FIN = {"Tramo RSH": "TRAMO_RSH", "Gratuidad": "GRATUIDAD_VALIDADA", "Beca": "BECA_VALIDADA", "CAE": "CAE_VALIDADO"}

# ==========================================
# 7. EXPORTACIÓN WORD DETALLADA
# ==========================================
st.sidebar.markdown("---")
st.sidebar.header("📄 Informe Exportable")

if DOCX_DISPONIBLE:
    def exportar_word_hiperdetallado():
        doc = Document()
        doc.add_heading('Informe Detallado de Análisis Institucional CFTe-AP', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generado el: {datetime.datetime.now().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        def imprimir_tabla(df_data, titulo, columnas):
            if df_data.empty: return
            doc.add_heading(titulo, level=3)
            table = doc.add_table(rows=1, cols=len(columnas))
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, col in enumerate(columnas):
                hdr[i].text = col
                hdr[i].paragraphs[0].runs[0].bold = True
                
            for _, row in df_data.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(columnas):
                    cells[i].text = str(row[col])
            doc.add_paragraph()

        for anio in sel_anios:
            df_anio = df_mat_real[df_mat_real['ANIO_INGRESO'] == anio]
            if df_anio.empty: continue
            
            doc.add_heading(f'=== COHORTE {anio} ===', level=1)
            
            # --- Cálculo de métricas alineadas al estándar validado ---
        df_tasas = df_anio.groupby(['JORNADA', 'NOMBRE_CARRERA']).agg(
            Matriculados=('RUT', 'count'),
            Retenidos_2do=('ES_RETENIDO', 'sum'),
            Desc_1er=('ES_DESERTOR_1ER', 'sum'),
            Desc_2do=('ES_DESERTOR_2DO', 'sum'),
            Desc_3er=('ES_DESERTOR_3ER', 'sum'),
            Progresados=('ES_PROGRESADO', 'sum'),
            Titulados=('ES_TITULADO', 'sum')
        ).reset_index()

        df_tasas['% Retención'] = (df_tasas['Retenidos_2do'] / df_tasas['Matriculados'] * 100).round(1).astype(str) + "%"
        
        df_tasas['% Fuga 1er'] = (df_tasas['Desc_1er'] / df_tasas['Matriculados'] * 100).round(1).astype(str) + "%"
        df_tasas['% Fuga 2do'] = (df_tasas['Desc_2do'] / df_tasas['Matriculados'] * 100).round(1).astype(str) + "%"
        
        df_tasas['% Fuga 3er'] = np.where(df_tasas['Retenidos_2do'] > 0, (df_tasas['Desc_3er'] / df_tasas['Retenidos_2do'] * 100).round(1).astype(str) + "%", "0%")
        df_tasas['% Progresión'] = np.where(df_tasas['Retenidos_2do'] > 0, (df_tasas['Progresados'] / df_tasas['Retenidos_2do'] * 100).round(1).astype(str) + "%", "0%")
        df_tasas['% Titulación'] = (df_tasas['Titulados'] / df_tasas['Matriculados'] * 100).round(1).astype(str) + "%"

        columnas_word = [
            'JORNADA', 'NOMBRE_CARRERA', 'Matriculados', 
            'Desc_1er', '% Fuga 1er',
            'Retenidos_2do', '% Retención', 
            'Desc_2do', '% Fuga 2do',
            'Desc_3er', '% Fuga 3er',
            'Progresados', '% Progresión', 
            'Titulados', '% Titulación'
        ]

        imprimir_tabla(df_tasas, "2. Matriz Consolidada de Trayectoria y Deserción", columnas_word)
            
            # Anexo Motivos
        doc.add_heading('3. Anexo: Motivos de Retiro Formalizados', level=2)
        df_fugas_anio = df_fugas_export[(df_fugas_export['ANIO_RETIRO_EXTRACT'] == anio)]
        if not df_fugas_anio.empty:
               df_cross = pd.crosstab(df_fugas_anio['MOTIVO_LIMPIO'], df_fugas_anio['SEMESTRE_RETIRO']).reset_index()
               imprimir_tabla(df_cross, f"Desglose de Retiros Formales Año {anio}", df_cross.columns)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    st.sidebar.download_button(label="📥 Descargar Reporte Completo (Word)", data=exportar_word_hiperdetallado(), file_name=f"Reporte_AI_CFTeAP.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else: st.sidebar.error("Librería python-docx requerida.")

# ==========================================
# 8. FUNCIONES DE VISUALIZACIÓN
# ==========================================
def grafico_con_porcentaje(df_plot, x_col, y_col, title, color_map=None, is_horizontal=False, order=None):
    df_plot = df_plot.copy()
    if order:
        df_plot[x_col if not is_horizontal else y_col] = pd.Categorical(df_plot[x_col if not is_horizontal else y_col], categories=order, ordered=True)
        df_plot = df_plot.sort_values(x_col if not is_horizontal else y_col)

    total = df_plot[y_col].sum() if not is_horizontal else df_plot[x_col].sum()
    if total == 0: return st.info(f"Sin datos para {title}")
    
    val_col = y_col if not is_horizontal else x_col
    cat_col = x_col if not is_horizontal else y_col
    df_plot['%'] = (df_plot[val_col] / total * 100).round(1)
    df_plot['Etiqueta'] = df_plot[val_col].astype(int).astype(str) + " (" + df_plot['%'].astype(str) + "%)"
    
    kwargs = {'x': x_col, 'y': y_col, 'title': title, 'text': 'Etiqueta'}
    if color_map and cat_col in color_map:
        kwargs['color'] = cat_col
        kwargs['color_discrete_map'] = color_map
    else:
        kwargs['color'] = cat_col
        kwargs['color_discrete_sequence'] = COLORES_PASTEL

    if is_horizontal:
        kwargs['orientation'] = 'h'
        fig = px.bar(df_plot, **kwargs)
        if not order: fig.update_layout(yaxis={'categoryorder':'total ascending'})
    else:
        fig = px.bar(df_plot, **kwargs)
    
    fig.update_layout(showlegend=False)
    st.plotly_chart(fig, use_container_width=True)

def render_comparative_subtabs(df_target, col_objetivo, titulo_obj, is_dropout=False, col_declarado=None):
    ts1, ts2, ts3 = st.tabs(["Demografía y Sociedad", "Educación Media", "Financiamiento"])
    def graficar_dims(dims_dict):
        for nombre_dim, col_real in dims_dict.items():
            if col_real not in df_target.columns: continue
            
            # --- 1. Gráfico Base (Usa col_objetivo para el Total de Desertores) ---
            df_grp = df_target.groupby(col_real).agg(Matriculados=('RUT', 'count'), Objetivo=(col_objetivo, 'sum')).reset_index()
            df_grp = df_grp[df_grp['Matriculados'] > 0]
            if df_grp.empty: continue
            
            df_grp['Tasa (%)'] = (df_grp['Objetivo'] / df_grp['Matriculados'] * 100).round(1)
            df_melt = df_grp.melt(id_vars=[col_real, 'Tasa (%)'], value_vars=['Matriculados', 'Objetivo'], var_name='Métrica', value_name='Cantidad')
            df_melt['Etiqueta'] = df_melt.apply(lambda x: f"{int(x['Cantidad'])}" if x['Métrica'] == 'Matriculados' else f"{int(x['Cantidad'])} ({x['Tasa (%)']}%)", axis=1)
            df_melt['Métrica'] = df_melt['Métrica'].replace({'Objetivo': titulo_obj})
            
            mapa = {'Matriculados': COLOR_BASE, titulo_obj: COLOR_NO if is_dropout else COLOR_SI}
            fig = px.bar(df_melt, x=col_real, y='Cantidad', color='Métrica', barmode='group', text='Etiqueta', title=f'Matriculados vs {titulo_obj} por {nombre_dim}', color_discrete_map=mapa)
            st.plotly_chart(fig, use_container_width=True, key=f"base_{titulo_obj}_{nombre_dim}")
            
            # --- 2. Heatmap de Motivos (Usa col_declarado para centrarse solo en los que tienen motivo) ---
            if is_dropout:
                filtro_heatmap = col_declarado if col_declarado else col_objetivo
                df_des = df_target[(df_target[filtro_heatmap] == True) & (df_target['MOTIVO_DESERCION'].notna())]
                
                if not df_des.empty:
                    df_mot = df_des.groupby([col_real, 'MOTIVO_DESERCION']).size().reset_index(name='Cantidad')
                    
                    if not df_mot.empty:
                        fig_m = px.density_heatmap(
                            df_mot,
                            x='MOTIVO_DESERCION',
                            y=col_real,
                            z='Cantidad',
                            histfunc='sum',
                            text_auto=True,
                            title=f'Motivos Declarados segmentados por {nombre_dim}',
                            color_continuous_scale='Teal'
                        )
                        
                        fig_m.update_layout(
                            xaxis_title="Motivo Declarado",
                            yaxis_title=nombre_dim,
                            xaxis_tickangle=-45,
                            margin=dict(b=120)
                        )
                        
                        st.plotly_chart(fig_m, use_container_width=True, key=f"motivos_{titulo_obj}_{nombre_dim}")
    
    st.markdown("---")
    with ts1: graficar_dims(DIMENSIONES_DEMO)
    with ts2: graficar_dims(DIMENSIONES_EDU)
    with ts3: graficar_dims(DIMENSIONES_FIN)

# ==========================================
# 9. RENDERIZADO DE PESTAÑAS (UI)
# ==========================================
t_carac, t_ret, t_des, t_prog, t_tit = st.tabs(["👥 Caracterización", "📈 Retención", "⚠️ Deserción", "🚀 Progresión", "🎓 Titulación"])

# --- TAB 1: CARACTERIZACIÓN ---
with t_carac:
    st.header("Caracterización de Ingreso")
    total_est = len(df_mat_real)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Matriculados", f"{total_est:,}")
    c2.metric("Mujeres / Hombres", f"{(df_mat_real['SEXO']=='F').sum()} / {(df_mat_real['SEXO']=='M').sum()}")
    pct_grat = (df_mat_real['GRATUIDAD_VALIDADA']=='Sí').sum() / total_est * 100 if total_est > 0 else 0
    c3.metric("Gratuidad Aprobada", f"{(df_mat_real['GRATUIDAD_VALIDADA']=='Sí').sum():,}", f"{pct_grat:.1f}%")
    c4.metric("Promedio NEM General", f"{df_mat_real['PROMEDIO_NEM_CALC'].mean():.2f}" if not df_mat_real['PROMEDIO_NEM_CALC'].isna().all() else "N/A")
    st.markdown("---")
    
    sc1, sc2, sc3 = st.tabs(["🌎 Demografía y Sociedad", "📚 Educación Media", "💰 Financiamiento"])
    with sc1:
        cd1, cd2 = st.columns(2)
        with cd1:
            grafico_con_porcentaje(df_mat_real['SEXO'].value_counts().reset_index().rename(columns={'count':'Cantidad'}), 'SEXO', 'Cantidad', 'Distribución por Género', MAPA_SEXO)
            grafico_con_porcentaje(df_mat_real['RANGO_ETARIO'].value_counts().reset_index().rename(columns={'count':'Cantidad'}), 'RANGO_ETARIO', 'Cantidad', 'Rango Etario', order=['<=19', '20-29', '30-39', '40-49', '50-59', '60-69', '70-79', '80-89', '90-99'])
            grafico_con_porcentaje(df_mat_real['NACIONALIDAD'].value_counts().head(10).reset_index().rename(columns={'count':'Cantidad'}), 'Cantidad', 'NACIONALIDAD', 'Nacionalidades (Top 10)', is_horizontal=True)

        with cd2:
            st.markdown("### Perfil Socio-Laboral y Cuidado")
            grafico_con_porcentaje(df_mat_real['TRABAJA_REAL'].value_counts().reset_index().rename(columns={'count':'Cantidad'}), 'TRABAJA_REAL', 'Cantidad', '¿Se encuentra trabajando?', MAPA_SINO)
            df_trabajan = df_mat_real[df_mat_real['TRABAJA_REAL'] == 'Sí']
            if not df_trabajan.empty:
                grafico_con_porcentaje(df_trabajan['CATEGORIA_OCUPACIONAL'].value_counts().reset_index().rename(columns={'count':'Cantidad'}), 'Cantidad', 'CATEGORIA_OCUPACIONAL', 'Clasificación Ocupacional', is_horizontal=True)
            
            df_cruce = df_mat_real.groupby(['TRABAJA_REAL', 'CONDICION_CUIDADOR']).size().reset_index(name='Cantidad')
            df_cruce['%'] = (df_cruce['Cantidad'] / len(df_mat_real) * 100).round(1)
            df_cruce['Etiqueta'] = df_cruce['Cantidad'].astype(str) + " (" + df_cruce['%'].astype(str) + "%)"
            fig_c = px.bar(df_cruce, x='TRABAJA_REAL', y='Cantidad', color='CONDICION_CUIDADOR', barmode='group', text='Etiqueta', title='Cruce: Trabajo vs Cuidador', color_discrete_map=MAPA_SINO)
            st.plotly_chart(fig_c, use_container_width=True)

    with sc2:
        ce1, ce2 = st.columns(2)
        with ce1:
            df_brecha = df_mat_real.dropna(subset=['BRECHA_CATEGORIA'])
            grafico_con_porcentaje(df_brecha['BRECHA_CATEGORIA'].value_counts().reset_index().rename(columns={'count':'Cantidad'}), 'BRECHA_CATEGORIA', 'Cantidad', 'Brecha de Egreso E.M.', order=['Directo (0 años)', '1-2 años', '3-5 años', '> 5 años'])
            
            for col in ['DATOS ENSEÑANZA MEDIA REGIMEN ENSEÑANZA', 'DATOS ENSEÑANZA MEDIA MODALIDAD', 'DATOS ENSEÑANZA MEDIA TIPO ENSEÑANZA']:
                if col in df_mat_real.columns:
                    grafico_con_porcentaje(df_mat_real[col].value_counts().reset_index().rename(columns={'count':'Cantidad'}), 'Cantidad', col, f'Distribución por {col.split()[-1]}', is_horizontal=True)

        with ce2:
            st.markdown("### Rendimiento Previo (NEM)")
            dim_nem = st.selectbox("NEM Promedio según:", ["Sexo", "Rango Etario", "Tipo Enseñanza", "Régimen", "Modalidad"])
            col_nem = DIMENSIONES_DEMO.get(dim_nem) or DIMENSIONES_EDU.get(dim_nem)
            if col_nem and not df_mat_real['PROMEDIO_NEM_CALC'].isna().all():
                df_nem = df_mat_real.groupby(col_nem)['PROMEDIO_NEM_CALC'].mean().reset_index()
                st.plotly_chart(px.bar(df_nem, x=col_nem, y='PROMEDIO_NEM_CALC', text_auto='.2f', title=f'NEM por {dim_nem}', color_discrete_sequence=[COLORES_PASTEL[0]]), use_container_width=True)

            st.markdown("### Especialidades y Pertinencia")
            
            # 1. Top 10 utilizando las agrupaciones (familias) del diccionario
            if 'DATOS ENSEÑANZA MEDIA ESPECIALIDAD' in df_mat_real.columns:
                def obtener_familia_em(esp_raw):
                    # Filtro 1: Descartar celdas completamente vacías o nulas de Excel
                    if pd.isna(esp_raw): 
                        return None
                        
                    esp_norm = normalizar_texto(str(esp_raw))
                    
                    # Filtro 2: Descartar textos que representan celdas sin datos
                    if esp_norm in ["", "NAN"] or "SIN ESPECIALIDAD" in esp_norm or "NINGUNA" in esp_norm: 
                        return None
                        
                    # Filtro 3: Buscar coincidencia en el diccionario
                    for familia, keywords in DICT_FAMILIAS_EM.items():
                        if any(normalizar_texto(kw) in esp_norm for kw in keywords):
                            return familia
                            
                    # Si la celda TIENE datos, pero no está en el diccionario:
                    return "Otras Especialidades"

                # Creamos la columna mapeada para agrupar
                df_mat_real['FAMILIA_ESPECIALIDAD'] = df_mat_real['DATOS ENSEÑANZA MEDIA ESPECIALIDAD'].apply(obtener_familia_em)
                
                # Descartamos los "None" (que son las celdas vacías o "Sin Especialidad")
                df_familias = df_mat_real.dropna(subset=['FAMILIA_ESPECIALIDAD'])
                
                if not df_familias.empty:
                    grafico_con_porcentaje(
                        df_familias['FAMILIA_ESPECIALIDAD'].value_counts().head(10).reset_index().rename(columns={'count':'Cantidad'}), 
                        'Cantidad', 'FAMILIA_ESPECIALIDAD', 'Top 10 Familias de Especialidades E.M.', is_horizontal=True
                    )
            
            # 2. Cruce de Pertinencia agrupado por Carrera CFT
            # Función ultra-estricta para filtrar celdas vacías o sin especialidad real
            def es_especialidad_valida(esp_raw):
                texto = normalizar_texto(str(esp_raw))
                # Lista negra de valores que NO son especialidades técnicas
                invalidos = ["", "NAN", "SIN ESPECIALIDAD", "NINGUNA", "SIN INFORMACION", "NO POSEE", "0", "---", "-", "SIN ESPEC"]
                return texto not in invalidos

            # Filtramos el dataframe original para obtener solo los que TIENEN especialidad válida
            df_con_esp = df_mat_real[df_mat_real['DATOS ENSEÑANZA MEDIA ESPECIALIDAD'].apply(es_especialidad_valida)].copy()
            
            if not df_con_esp.empty:
                st.info(f"De los **{total_est:,}** matriculados, **{len(df_con_esp):,}** cuentan con una especialidad técnica registrada. El siguiente gráfico desglosa la pertinencia de esos {len(df_con_esp):,} registros.")
                
                # Asignamos la familia de especialidad
                df_con_esp['FAMILIA_ESPECIALIDAD'] = df_con_esp['DATOS ENSEÑANZA MEDIA ESPECIALIDAD'].apply(lambda x: next((f for f, keys in DICT_FAMILIAS_EM.items() if any(normalizar_texto(kw) in normalizar_texto(str(x)) for kw in keys)), "Otras Especialidades"))
                
                # Agrupamos cruzando la Carrera con el Nivel de Pertinencia
                df_pert_carrera = df_con_esp.groupby(['NOMBRE_CARRERA', 'PERTINENCIA_ESPECIALIDAD']).size().reset_index(name='Cantidad')
                
                # Cálculo porcentual por carrera
                totales_por_carrera = df_pert_carrera.groupby('NOMBRE_CARRERA')['Cantidad'].transform('sum')
                df_pert_carrera['%'] = (df_pert_carrera['Cantidad'] / totales_por_carrera * 100).round(1)
                df_pert_carrera['Etiqueta'] = df_pert_carrera['Cantidad'].astype(str) + " (" + df_pert_carrera['%'].astype(str) + "%)"
                
                # Gráfico de columnas verticales y agrupadas
                fig_p = px.bar(
                    df_pert_carrera, x='NOMBRE_CARRERA', y='Cantidad', color='PERTINENCIA_ESPECIALIDAD',
                    barmode='group', text='Etiqueta', # barmode='group' separa las barras en lugar de apilarlas
                    title='Alineación Especialidad E.M. vs Carrera CFT',
                    color_discrete_map={'Pertinencia Alta': COLOR_SI, 'Pertinencia Media': '#F9E79F', 'Sin Pertinencia': COLOR_NO}
                )
                
                # Ajustes para que los nombres de las carreras se lean bien en el eje X
                fig_p.update_layout(
                    xaxis_title="",
                    yaxis_title="Cantidad de Estudiantes",
                    xaxis_tickangle=-45, # Inclina los nombres de las carreras para que no se superpongan
                    height=600, # Aumentamos un poco la altura para acomodar el texto inclinado
                    legend_title_text='Nivel de Pertinencia',
                    margin=dict(b=150) # Margen inferior adicional para evitar que se corten los nombres
                )
                
                st.plotly_chart(fig_p, use_container_width=True)
            else:
                st.warning("No se encontraron registros de especialidades técnicas válidas tras el filtrado.")

    with sc3:
        cf1, cf2 = st.columns(2)
        with cf1:
            df_fin = pd.DataFrame({'Beneficio': ['Gratuidad', 'Beca', 'CAE'], 'Cantidad': [(df_mat_real['GRATUIDAD_VALIDADA']=='Sí').sum(), (df_mat_real['BECA_VALIDADA']=='Sí').sum(), (df_mat_real['CAE_VALIDADO']=='Sí').sum()]})
            grafico_con_porcentaje(df_fin, 'Beneficio', 'Cantidad', 'Estudiantes con Beneficios')
        with cf2:
            orden_rsh = ['Tramo 1 0% - 40%', 'Tramo 2 41% - 50%', 'Tramo 3 51% - 60%', 'Tramo 4 61% - 70%', 'Tramo 5 71% - 80%', 'Tramo 6 81% - 90%', 'Tramo 7 91% - 100%', 'Sin Registro']
            df_rsh = df_mat_real['TRAMO_RSH'].value_counts().reindex(orden_rsh).reset_index().rename(columns={'count':'Cantidad'}).fillna(0)
            grafico_con_porcentaje(df_rsh, 'TRAMO_RSH', 'Cantidad', 'Distribución Tramos RSH', order=orden_rsh)

# --- TAB 2: RETENCIÓN ---
with t_ret:
    st.header("Retención al 2do Año")
    df_ret_valid = df_mat_real[df_mat_real['ANIO_INGRESO'] < max_anio_registrado]
    if df_ret_valid.empty: st.info("Cohortes sin ciclo para medir retención.")
    else: 
        t_ing, t_ret = len(df_ret_valid), df_ret_valid['ES_RETENIDO'].sum()
        c1, c2, c3 = st.columns(3)
        c1.metric("Matriculados Base", f"{t_ing:,}")
        c2.metric("Total Retenidos", f"{t_ret:,}")
        c3.metric("Tasa de Retención", f"{round(t_ret/t_ing*100,1) if t_ing>0 else 0}%")
        st.markdown("---")
        
        df_tc = df_ret_valid.groupby(['TRABAJA_REAL', 'CONDICION_CUIDADOR']).agg(Matriculados=('RUT', 'count'), Retenidos=('ES_RETENIDO', 'sum')).reset_index()
        df_tc['Tasa (%)'] = (df_tc['Retenidos'] / df_tc['Matriculados'] * 100).round(1)
        df_tc_melt = df_tc.melt(id_vars=['TRABAJA_REAL', 'CONDICION_CUIDADOR', 'Tasa (%)'], value_vars=['Matriculados', 'Retenidos'], var_name='Métrica', value_name='Cantidad')
        df_tc_melt['Eje_X'] = "Trabaja: " + df_tc_melt['TRABAJA_REAL'] + " | Cuidador: " + df_tc_melt['CONDICION_CUIDADOR']
        df_tc_melt['Etiqueta'] = df_tc_melt.apply(lambda x: f"{int(x['Cantidad'])}" if x['Métrica'] == 'Matriculados' else f"{int(x['Cantidad'])} ({x['Tasa (%)']}%)", axis=1)
        fig_tc = px.bar(df_tc_melt, x='Eje_X', y='Cantidad', color='Métrica', barmode='group', text='Etiqueta', title='Efectividad de Retención: Trabajo vs Cuidador', color_discrete_map={'Matriculados': COLOR_BASE, 'Retenidos': COLOR_SI})
        st.plotly_chart(fig_tc, use_container_width=True)
        st.markdown("---")
        
        render_comparative_subtabs(df_ret_valid, 'ES_RETENIDO', 'Retenidos')

# --- TAB 3: DESERCIÓN ---
    with t_des:
        st.header("Análisis de Deserción Institucional")

        s_1er, s_2do, s_3er = st.tabs(["Deserción 1er Año", "Deserción 2do Año", "Deserción 3er Año"])

        def graficar_desercion_motivos(df_subset, titulo_anual):
            df_mot = df_subset.groupby(['SEMESTRE_RETIRO', 'MOTIVO_LIMPIO']).size().reset_index(name='Cantidad')
            if df_mot.empty: return st.info("Sin retiros formalizados en este período.")
            fig_m = px.bar(
                df_mot,
                y='MOTIVO_LIMPIO',
                x='Cantidad',
                color='SEMESTRE_RETIRO',
                orientation='h',
                barmode='group',
                title=f'Motivos Formalizados por Semestre ({titulo_anual})',
                color_discrete_sequence=['#F5CBA7', '#A3E4D7', '#E5E7E9']
            )
            st.plotly_chart(fig_m, use_container_width=True)

        with s_1er:
            df_d1 = df_mat_real[df_mat_real['ANIO_INGRESO'] < max_anio_registrado]
            if not df_d1.empty:
                t_ing, t_des = len(df_d1), df_d1['ES_DESERTOR_1ER'].sum()
                c1, c2, c3 = st.columns(3)
                c1.metric("Matriculados Base", f"{int(t_ing):,}")
                c2.metric("Desertores Declarados", f"{int(t_des):,}")
                c3.metric("Tasa Deserción", f"{round(t_des/t_ing*100,1) if t_ing>0 else 0}%")
                st.markdown("---")
                df_f1 = df_fugas_export[(df_fugas_export['ANIO_RETIRO_EXTRACT'] == df_fugas_export['ANIO_INGRESO']) & (df_fugas_export['ANIO_INGRESO'] < max_anio_registrado)]
                graficar_desercion_motivos(df_f1, "1er Año")
                st.markdown("### Cruces Sociodemográficos")
                render_comparative_subtabs(df_d1, 'ES_DESERTOR_1ER', 'Desertores (1er Año)', is_dropout=True)

        with s_2do:
            df_d2 = df_mat_real[df_mat_real['ANIO_INGRESO'] < max_anio_registrado]
            if not df_d2.empty:
                t_ing = len(df_d2)
                t_des2 = df_d2['ES_DESERTOR_2DO'].sum()
                t_dec2 = df_d2['ES_DECLARADO_2DO'].sum()
                
                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Matriculados Base", f"{int(t_ing):,}")
                c2.metric("Desertores Totales", f"{int(t_des2):,}")
                c3.metric("Desertores Declarados", f"{int(t_dec2):,}")
                c4.metric("Tasa Deserción Total", f"{round(t_des2/t_ing*100,1) if t_ing>0 else 0}%")
                st.markdown("---")
                st.markdown("### Cruces Sociodemográficos (Deserción Total vs Motivos Declarados)")
                render_comparative_subtabs(df_d2, 'ES_DESERTOR_2DO', 'Desertores Totales (2do Año)', is_dropout=True, col_declarado='ES_DECLARADO_2DO')

        with s_3er:
            df_d3_base = df_mat_real[df_mat_real['ANIO_INGRESO'] <= (max_anio_registrado - 2)]
            df_d3_retenidos = df_d3_base[df_d3_base['ES_RETENIDO'] == True]

            if not df_d3_retenidos.empty:
                t_ret = len(df_d3_retenidos)
                t_des3 = df_d3_retenidos['ES_DESERTOR_3ER'].sum()
                
                # Contamos los declarados sobre la base completa para cuadrar con la extracción de motivos
                t_dec3 = df_d3_base['ES_DECLARADO_3ER'].sum()

                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Retenidos (Base 3er Año)", f"{int(t_ret):,}")
                c2.metric("Desertores Totales", f"{int(t_des3):,}")
                c3.metric("Desertores Declarados", f"{int(t_dec3):,}")
                c4.metric("Tasa Deserción Total", f"{round(t_des3/t_ret*100,1) if t_ret>0 else 0}%")
                st.markdown("---")
                st.markdown("### Cruces Sociodemográficos (Deserción Total vs Motivos Declarados)")
                render_comparative_subtabs(df_d3_retenidos, 'ES_DESERTOR_3ER', 'Desertores Totales (3er Año)', is_dropout=True, col_declarado='ES_DECLARADO_3ER')

# --- TAB 4: PROGRESIÓN ---
with t_prog:
    st.header("Progresión a 3er Año")
    df_prog_valid = df_mat_real[(df_mat_real['ANIO_INGRESO'] <= max_anio_registrado - 2) & (df_mat_real['ES_RETENIDO'] == True)]
    if df_prog_valid.empty: st.info("Cohortes sin ciclo para medir progresión.")
    else:
        t_ret, t_prog = len(df_prog_valid), df_prog_valid['ES_PROGRESADO'].sum()
        c1, c2, c3 = st.columns(3)
        c1.metric("Retenidos Base (2do Año)", f"{int(t_ret):,}")
        c2.metric("Progresados a 3er Año", f"{int(t_prog):,}")
        c3.metric("Tasa de Progresión", f"{round(t_prog/t_ret*100,1) if t_ret>0 else 0}%")
        st.markdown("---")
        render_comparative_subtabs(df_prog_valid, 'ES_PROGRESADO', 'Progresados')

# --- TAB 5: TITULACIÓN ---
with t_tit:
    st.header("Tasas de Titulación")
    df_tit_valid = df_mat_real[df_mat_real['TITULACION_CALCULABLE'] == True]
    
    if df_tit_valid.empty:
        st.info("Ninguna cohorte ha cumplido el plazo teórico (duración de carrera) a la fecha de hoy.")
    else:
        st_t1, st_t2 = st.tabs(["🎓 Titulación General", "⏱️ Titulación Oportuna (1 Año Desfase)"])
        with st_t1: 
            t_m, t_t = len(df_tit_valid), df_tit_valid['ES_TITULADO'].sum()
            cm1, cm2, cm3 = st.columns(3)
            cm1.metric("Matriculados Base", f"{t_m:,}")
            cm2.metric("Total Titulados", f"{t_t:,}")
            cm3.metric("Tasa de Titulación", f"{round(t_t/t_m*100,1) if t_m>0 else 0}%")
            render_comparative_subtabs(df_tit_valid, 'ES_TITULADO', "Titulados Global")
        with st_t2: 
            t_m, t_o = len(df_tit_valid), df_tit_valid['ES_TITULADO_OPORTUNO'].sum()
            co1, co2, co3 = st.columns(3)
            co1.metric("Matriculados Base", f"{t_m:,}")
            co2.metric("Titulados Oportunamente", f"{t_o:,}")
            co3.metric("Tasa Oportuna", f"{round(t_o/t_m*100,1) if t_m>0 else 0}%")
            render_comparative_subtabs(df_tit_valid, 'ES_TITULADO_OPORTUNO', "Titulados Oportunos")
