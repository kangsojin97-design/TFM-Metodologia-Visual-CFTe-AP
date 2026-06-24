import streamlit as st
import pandas as pd
import os
import plotly.express as px
import datetime
import numpy as np
import re
import io

# Intentar importar docx (librería para crear Word)
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

# --- 1. CONFIGURACIÓN E IDENTIDAD INSTITUCIONAL ---
st.set_page_config(page_title="Análisis Institucional CFTe-AP", layout="wide")

COLORES_CLAROS = ['#AED6F1', '#A9DFBF', '#F9E79F', '#F5CBA7', '#E5E7E9']

CARRERAS_MAP = {
    "180": "TNS EN ASISTENCIA JURÍDICA",
    "184": "TNS EN PROYECTOS ELÉCTRICOS DE DISTRIBUCIÓN",
    "201": "TNS EN ENFERMERÍA CON MENCIÓN EN GERONTOLOGÍA",
    "202": "TNS EN INFORMÁTICA Y APLICACIONES TECNOLÓGICAS",
    "204": "TNS EN GESTIÓN DE COMERCIO EXTERIOR",
    "283": "TNS EN EDUCACIÓN PARVULARIA Y PRIMER - SEGUNDO AÑO DE EDUCACIÓN BÁSICA",
    "380": "TNS EN DEPORTES Y RECREACIÓN",
    "554": "TNS EN ADMINISTRACIÓN PÚBLICA",
    "559": "TNS EN AGRÍCOLA",
    "602": "TNS EN EDUCACIÓN ESPECIAL",
    "603": "TNS EN ADMINISTRACIÓN DE EMPRESAS",
    "608": "TNS EN TRABAJO SOCIAL",
    "609": "TNS EN GEOLOGÍA",
    "613": "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA",
    "615": "TNS EN FABRICACIÓN Y MONTAJE DE ESTRUCTURAS METÁLICAS",
    "187": "TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES",
    "206": "TNS EN GESTIÓN CONTABLE",
    "207": "TNS EN GESTIÓN DE RECURSOS HUMANOS",
    "612": "TNS EN LABORATORIO CLÍNICO, BANCO DE SANGRE E IMAGENOLOGÍA",
    "153": "TNS EN OBRAS CIVILES",
    "188": "TNS EN MANTENIMIENTO ELECTROMECANICO CON MENCION EN ELECTROMOVILIDAD",
    "208": "TNS EN ADMINISTRACION PUBLICA CON MENCION EN GESTION JURIDICA",
    "209": "TNS EN ESTETICA CON MENCION EN MASOTERAPIA",
    "210": "TNS EN VETERINARIA",
    "284": "TNS EN EDUCACION PARVULARIA PRIMER CICLO DE EDUCACION BASICA",
    "219": "TNS EN ESTÉTICA Y COSMETOLOGIA MENCIÓN MASOTERAPIA"
}

st.title("🏛️ Dimensión: Docencia - Análisis de Matrícula y Permanencia")
st.markdown("---")

# --- 2. RUTAS ABSOLUTAS ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CARPETA_DOC = os.path.join(BASE_DIR, "2_Datos_Procesados", "2_Datos_Procesados_Docencia")

# --- 3. LECTOR INTELIGENTE Y LIMPIADOR ---
def limpiar_obs_jerarquia(obs):
    obs = str(obs).strip()
    cat = "OTRA SITUACIÓN (O VIGENTE)"
    if "ADMINISTRATIVO" in obs: cat = "RETIRO ADMINISTRATIVO"
    elif "DEFINITIVO" in obs: cat = "RETIRO DEFINITIVO"
    elif "TEMPORAL" in obs: cat = "RETIRO TEMPORAL"
    elif "PÉRDIDA" in obs or "PERDIDA" in obs: cat = "PÉRDIDA DE CARRERA"
    elif "SIN REGISTRO" in obs: cat = "SIN REGISTRO"
    
    motivo = obs
    match = re.search(r'\((.*?)\)', obs)
    if match:
        partes = match.group(1).split(',')
        motivo = partes[-1].strip()
    else:
        if cat in ["RETIRO ADMINISTRATIVO", "RETIRO DEFINITIVO", "RETIRO TEMPORAL", "PÉRDIDA DE CARRERA"]:
            motivo = "MOTIVO NO ESPECIFICADO"
    return cat, motivo

def cargar_archivo(ruta, nombre_archivo=""):
    try:
        if ruta.endswith(".xlsx"): df = pd.read_excel(ruta)
        elif ruta.endswith(".csv"):
            df = None
            for s in [';', ',']:
                for enc in ['utf-8-sig', 'latin-1']:
                    try:
                        df = pd.read_csv(ruta, sep=s, encoding=enc, on_bad_lines='skip', engine='python')
                        break
                    except: continue
                if df is not None: break
        
        if df is not None:
            df.columns = df.columns.str.strip().str.upper()
            if 'RUT' in df.columns:
                df['RUT'] = df['RUT'].astype(str).str.strip().str.upper().str.replace('.', '', regex=False).str.replace('-', '', regex=False)
            
            # Limpieza específica para Convalidaciones (Año y Carrera)
            if 'AÑO' in df.columns: df['ANIO_BENEFICIO'] = pd.to_numeric(df['AÑO'], errors='coerce')
            elif 'ANIO' in df.columns: df['ANIO_BENEFICIO'] = pd.to_numeric(df['ANIO'], errors='coerce')
            elif 'FECHA' in df.columns: df['ANIO_BENEFICIO'] = pd.to_datetime(df['FECHA'], errors='coerce').dt.year
            
            if 'CARRERA' in df.columns and 'COD_CARRERA' not in df.columns and 'CARRERA_COD' not in df.columns:
                df['COD_EXTRAIDO'] = df['CARRERA'].astype(str).str.extract(r'(\d{3})', expand=False)
                df['NOMBRE_CARRERA_CONV'] = df['COD_EXTRAIDO'].map(CARRERAS_MAP).fillna("CARRERA DESCONOCIDA")

            col_codigo = 'CARRERA_COD' if 'CARRERA_COD' in df.columns else None
            if not col_codigo and 'COD_CARRERA' in df.columns: col_codigo = 'COD_CARRERA'
            if col_codigo:
                df['COD_CARRERA'] = df[col_codigo].astype(str).str.extract(r'(\d{3})', expand=False)
                df['NOMBRE_CARRERA'] = df['COD_CARRERA'].map(CARRERAS_MAP).fillna("CARRERA DESCONOCIDA (" + df[col_codigo].astype(str) + ")")
                df['NOMBRE_CARRERA_CONV'] = df['NOMBRE_CARRERA']

            if 'JORNADA' in df.columns:
                df['JORNADA'] = df['JORNADA'].astype(str).str.strip().str.upper()
                
            col_obs = next((col for col in df.columns if 'OBSERVACI' in col), None)
            if col_obs:
                df['OBSERVACION'] = df[col_obs].astype(str).str.strip().str.upper()
                df['OBSERVACION'] = df['OBSERVACION'].str.replace('SITUACIÓN:', '', regex=False).str.replace('SITUACION:', '', regex=False).str.strip()
                df.loc[df['OBSERVACION'] == 'NAN', 'OBSERVACION'] = "SIN REGISTRO"
            else: 
                df['OBSERVACION'] = "SIN REGISTRO"
            return df
    except: return None
    return None

# --- 4. CARGA DE DATOS ---
archivos_todos = os.listdir(CARPETA_DOC) if os.path.exists(CARPETA_DOC) else []
archivos_matriculados = [f for f in archivos_todos if f.startswith("LIMPIO_ING") and "MAT" in f]
archivos_oferta = [f for f in archivos_todos if f.startswith("LIMPIO_OFERTA") or f.startswith("LIMPIO_oferta")]
archivos_conval = [f for f in archivos_todos if "LIMPIO_CONVALIDACION" in f.upper()]

if not archivos_matriculados:
    st.error("❌ No se han encontrado archivos de Matriculados procesados.")
    st.stop()

# Matriculados (Solo desde cohorte 2021+)
lista_matriculados = []
for arc in archivos_matriculados:
    parts = arc.upper().replace("LIMPIO_ING", "").replace(".XLSX", "").replace(".CSV", "").split("_MAT")
    if len(parts) == 2:
        anio_ing = int(parts[0])
        if anio_ing >= 2021:
            df_temp = cargar_archivo(os.path.join(CARPETA_DOC, arc), arc)
            if df_temp is not None:
                df_temp['ANIO_INGRESO'] = anio_ing
                df_temp['ANIO_MATRICULA'] = int(parts[1])
                lista_matriculados.append(df_temp)
df_matriculados = pd.concat(lista_matriculados, ignore_index=True) if lista_matriculados else pd.DataFrame()

# Oferta (Solo desde 2021+)
lista_oferta = []
for arc in archivos_oferta:
    anio_of = "".join(filter(str.isdigit, arc))
    if anio_of and int(anio_of) >= 2021:
        df_temp = cargar_archivo(os.path.join(CARPETA_DOC, arc), arc)
        if df_temp is not None:
            df_temp['ANIO_OFERTA'] = int(anio_of)
            lista_oferta.append(df_temp)
df_oferta = pd.concat(lista_oferta, ignore_index=True) if lista_oferta else pd.DataFrame()

# Convalidaciones (Histórico completo)
lista_conval = []
for arc in archivos_conval:
    df_temp = cargar_archivo(os.path.join(CARPETA_DOC, arc), arc)
    if df_temp is not None:
        lista_conval.append(df_temp)
df_convalidaciones = pd.concat(lista_conval, ignore_index=True) if lista_conval else pd.DataFrame()

# --- Criterio de Corte para Totales ---
# Identifica el último año de matrícula disponible en los datos (ej. 2025)
max_anio_registrado = df_matriculados['ANIO_MATRICULA'].max() if not df_matriculados.empty else datetime.datetime.now().year

# --- 5. FILTROS LATERALES ---
st.sidebar.header("🔍 Filtros de Análisis")
anios_disp = sorted(df_matriculados['ANIO_INGRESO'].dropna().unique()) if 'ANIO_INGRESO' in df_matriculados.columns else []
sel_anios = st.sidebar.multiselect("Año de Análisis", anios_disp, default=anios_disp)

nombres_disp = sorted(df_matriculados['NOMBRE_CARRERA'].dropna().unique()) if 'NOMBRE_CARRERA' in df_matriculados.columns else []
sel_nombres = st.sidebar.multiselect("Carrera", nombres_disp, default=nombres_disp)

jornadas_disp = sorted(df_matriculados['JORNADA'].dropna().unique()) if 'JORNADA' in df_matriculados.columns else []
sel_jornadas = st.sidebar.multiselect("Jornada (No aplica a Convalidación)", jornadas_disp, default=jornadas_disp)

df_filt = df_matriculados.copy()
if not df_filt.empty:
    df_filt = df_filt[df_filt['ANIO_INGRESO'].isin(sel_anios)]
    if sel_nombres: df_filt = df_filt[df_filt['NOMBRE_CARRERA'].isin(sel_nombres)]
    if sel_jornadas: df_filt = df_filt[df_filt['JORNADA'].isin(sel_jornadas)]
    
    if 'OBSERVACION' in df_filt.columns:
        df_filt['CAT_OBS'], df_filt['MOTIVO_OBS'] = zip(*df_filt['OBSERVACION'].map(limpiar_obs_jerarquia))

# Universo Real de 1er Año para Matrícula/Retención/Deserción
df_mat_real = df_filt[df_filt['ANIO_INGRESO'] == df_filt['ANIO_MATRICULA']] if not df_filt.empty else pd.DataFrame()

# --- 6. PESTAÑAS Y CÁLCULO DE INDICADORES ---
tab1, tab2, tab3, tab4, tab5 = st.tabs([" 📚  Tasa de Matrícula", " 🔄  Tasa de Retención", " 📉  Tasa de Deserción", " ✅  Convalidación y Apoyo", " 📈  Tasa de Progresión"])

# Variables globales para reporte Word
df_det_mat = pd.DataFrame()
df_det_ret = pd.DataFrame()
df_det_des = pd.DataFrame()
df_fugas_export = pd.DataFrame()
df_dist_acort = pd.DataFrame()
df_hist_acort = pd.DataFrame()
total_tutorias_cont = 0
df_det_prog = pd.DataFrame()

# ----------------- TAB 1: MATRÍCULA -----------------
with tab1:
    st.header("Oferta vs Matrícula")
    
    if not df_mat_real.empty:
        mat_agrupado = df_mat_real.groupby(['ANIO_INGRESO', 'NOMBRE_CARRERA', 'JORNADA']).size().reset_index(name='Matriculados')
        mat_agrupado.rename(columns={'ANIO_INGRESO': 'Año', 'NOMBRE_CARRERA': 'Carrera', 'JORNADA': 'Jornada'}, inplace=True)
    else:
        mat_agrupado = pd.DataFrame(columns=['Año', 'Carrera', 'Jornada', 'Matriculados'])

    vacantes_list = []
    oferta_filt = df_oferta.copy()
    if not oferta_filt.empty:
        if sel_anios: oferta_filt = oferta_filt[oferta_filt['ANIO_OFERTA'].isin(sel_anios)]
        if sel_nombres: oferta_filt = oferta_filt[oferta_filt['NOMBRE_CARRERA'].isin(sel_nombres)]
            
        for _, row in oferta_filt.iterrows():
            anio_of = row.get('ANIO_OFERTA')
            carrera_of = row.get('NOMBRE_CARRERA', 'CARRERA DESCONOCIDA')
            
            if not sel_jornadas or "DIURNA" in [j.upper() for j in sel_jornadas]:
                vac_d = pd.to_numeric(row.get('OFERTA DIURNA', 0), errors='coerce')
                if pd.notna(vac_d) and vac_d > 0:
                    vacantes_list.append({"Año": anio_of, "Carrera": carrera_of, "Jornada": "DIURNA", "Vacantes": vac_d})
                    
            if not sel_jornadas or "VESPERTINA" in [j.upper() for j in sel_jornadas]:
                vac_v = pd.to_numeric(row.get('OFERTA VESPERTINA', 0), errors='coerce')
                if pd.notna(vac_v) and vac_v > 0:
                    vacantes_list.append({"Año": anio_of, "Carrera": carrera_of, "Jornada": "VESPERTINA", "Vacantes": vac_v})

    df_vacantes = pd.DataFrame(vacantes_list)

    if not mat_agrupado.empty and not df_vacantes.empty:
        df_det_mat = pd.merge(df_vacantes, mat_agrupado, on=['Año', 'Carrera', 'Jornada'], how='outer').fillna(0)
    elif not mat_agrupado.empty:
        df_det_mat = mat_agrupado.copy()
        df_det_mat['Vacantes'] = 0
    elif not df_vacantes.empty:
        df_det_mat = df_vacantes.copy()
        df_det_mat['Matriculados'] = 0
    else:
        df_det_mat = pd.DataFrame(columns=['Año', 'Carrera', 'Jornada', 'Vacantes', 'Matriculados'])

    if not df_det_mat.empty:
        df_det_mat['Vacantes'] = df_det_mat['Vacantes'].astype(int)
        df_det_mat['Matriculados'] = df_det_mat['Matriculados'].astype(int)
        df_det_mat['Tasa Matrícula (%)'] = df_det_mat.apply(
            lambda x: round((x['Matriculados']/x['Vacantes']*100), 1) if x['Vacantes'] > 0 else 0, axis=1
        )
        
        global_mat_vacantes = df_det_mat['Vacantes'].sum()
        global_mat_matriculados = df_det_mat['Matriculados'].sum()
        tasa_tot = round(global_mat_matriculados / global_mat_vacantes * 100, 1) if global_mat_vacantes > 0 else 0
        promedio_tasa = round(df_det_mat["Tasa Matrícula (%)"].mean(), 1)
        
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Total Vacantes Ofertadas", f"{int(global_mat_vacantes)}")
        col2.metric("Total Matriculados Reales", f"{int(global_mat_matriculados)}")
        col3.metric("Tasa Global Acumulada", f"{tasa_tot}%")
        col4.metric("Promedio de las Tasas", f"{promedio_tasa}%")
        st.markdown("---")
        
        if len(sel_anios) == 1:
            res_mat = df_det_mat.groupby('Jornada')[['Vacantes', 'Matriculados']].sum().reset_index()
            df_melt = res_mat.melt(id_vars='Jornada', value_vars=['Vacantes', 'Matriculados'], var_name='Métrica', value_name='Cantidad')
            fig1 = px.bar(df_melt, x="Jornada", y="Cantidad", color="Métrica", barmode='group', color_discrete_sequence=[COLORES_CLAROS[0], COLORES_CLAROS[1]], title=f"Oferta vs Matriculados Reales ({sel_anios[0]}) por Jornada", text='Cantidad')
        else:
            res_mat = df_det_mat.groupby('Año')[['Vacantes', 'Matriculados']].sum().reset_index()
            res_mat['Año'] = res_mat['Año'].astype(str)
            df_melt = res_mat.melt(id_vars='Año', value_vars=['Vacantes', 'Matriculados'], var_name='Métrica', value_name='Cantidad')
            fig1 = px.bar(df_melt, x="Año", y="Cantidad", color="Métrica", barmode='group', color_discrete_sequence=[COLORES_CLAROS[0], COLORES_CLAROS[1]], title="Evolución Histórica Oferta vs Matriculados Reales", text='Cantidad')
        st.plotly_chart(fig1, use_container_width=True)

# ----------------- TAB 2: RETENCIÓN -----------------
with tab2:
    st.header("Tasa de Retención (1er a 2do Año)")
    detalles_ret = []
    for anio in sel_anios:
        for nombre_car in sel_nombres if sel_nombres else df_filt['NOMBRE_CARRERA'].unique():
            for jor in sel_jornadas if sel_jornadas else df_filt['JORNADA'].unique():
                df_base = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'] == anio) & (df_filt['NOMBRE_CARRERA'] == nombre_car) & (df_filt['JORNADA'] == jor)]
                ingresos = len(df_base)
                
                if ingresos > 0:
                    ruts_iniciales = set(df_base['RUT'].dropna())
                    df_sig = df_matriculados[(df_matriculados['ANIO_INGRESO'] == anio) & (df_matriculados['ANIO_MATRICULA'] == anio + 1)]
                    ruts_sig = set(df_sig['RUT'].dropna())
                    retenidos = len(ruts_iniciales & ruts_sig)
                    detalles_ret.append({"Cohorte": anio, "Carrera": nombre_car, "Jornada": jor, "Ingresos (1er Año)": ingresos, "Retenidos (2do Año)": retenidos, "Tasa Retención (%)": round((retenidos/ingresos*100), 1)})
    
    df_det_ret = pd.DataFrame(detalles_ret)

    if not df_det_ret.empty:
        # 1. NUEVO CRITERIO: Filtramos las cohortes que aún no tienen año de comparación
        df_det_ret = df_det_ret[df_det_ret['Cohorte'] < max_anio_registrado]
        
        # 2. AHORA recalculamos los totales globales usando el DataFrame ya filtrado
        global_ret_ingresos = df_det_ret['Ingresos (1er Año)'].sum()
        global_ret_retenidos = df_det_ret['Retenidos (2do Año)'].sum()
        
        # 3. Calculamos las tasas finales
        tasa_tot = round(global_ret_retenidos / global_ret_ingresos * 100, 1) if global_ret_ingresos > 0 else 0
        promedio_tasa = round(df_det_ret["Tasa Retención (%)"].mean(), 1)
        
        # 4. Mostramos las métricas
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Matriculados 1er Año", f"{int(global_ret_ingresos)}")
        col2.metric("Retenidos 2do Año", f"{int(global_ret_retenidos)}")
        col3.metric("Tasa Global Acumulada", f"{tasa_tot}%")
        col4.metric("Promedio de las Tasas", f"{promedio_tasa}%")
        st.markdown("---")
        
        res_ret = df_det_ret.groupby('Cohorte')[['Ingresos (1er Año)', 'Retenidos (2do Año)']].sum().reset_index()
        res_ret['Cohorte'] = res_ret['Cohorte'].astype(str)
        df_melt_ret = res_ret.melt(id_vars='Cohorte', value_vars=['Ingresos (1er Año)', 'Retenidos (2do Año)'], var_name='Métrica', value_name='Cantidad')
        fig2 = px.bar(df_melt_ret, x="Cohorte", y="Cantidad", color="Métrica", barmode='group', color_discrete_sequence=[COLORES_CLAROS[1], COLORES_CLAROS[2]], title="Comparativa Global Ingresos vs Retenidos", text='Cantidad')
        st.plotly_chart(fig2, use_container_width=True)

# ----------------- TAB 3: DESERCIÓN -----------------
        with tab3:
            st.header("Tasa de Deserción y Situación Académica")

            # Subpestañas para organizar el análisis
            subtab_1er, subtab_2do, subtab_3er = st.tabs(["  📉   Deserción al 1er Año", "  📉   Deserción al 2do Año", "  📉   Deserción al 3er Año"])

            detalles_des_1er = []
            detalles_des_2do = []
            detalles_des_3er = []

            fugas_1er_list = []
            fugas_2do_list = []
            fugas_3er_list = []

            import unicodedata
            import re

            def normalizar_texto(texto):
                if not texto or texto == "None": return "MOTIVO NO ESPECIFICADO"
                texto = str(texto).upper().strip()
                texto = ''.join(c for c in unicodedata.normalize('NFD', texto) if unicodedata.category(c) != 'Mn')
                return texto

            LISTA_AGRUPACION = {
                "CAMBIO DE INSTITUCIÓN": ["INGRESO A OTRA INSTITUCIÓN", "CAMBIO DE INSTITUCIÓN", "CAMBIO A OTRA INSTITUCIÓN", "INGRESO A OTRA IES", "INGRESO OTRA IES", "INGRESO A LAS FUERZAS ARMADAS"],
                "CAMBIO DE CARRERA": ["CAMBIO DE CARRERA", "CAMBIO DE CARRERA EN LA INSTITUCIÓN", "INGRESO OTRA CARRERA"],
                "VOCACIONAL": ["NO ME GUSTO LA CARRERA", "DESCONFORMIDAD CON LA CARRERA Y SU METODOLOGÍA", "DISCONFORMIDAD DE LA CARRERA", "INCONFORMIDAD CON LA CARRERA", "LA CARRERA NO CUMPLIÓ MIS EXPECTATIVAS", "LA CARRERA NO CUMPLE MIS EXPECTATIVAS", "VOCACIONAL", "VOCACIONAL INGRESO 2024", "PROBLEMA VOCACIONAL", "VOCACIONAL IGRESO 2024"],
                "PROBLEMAS DE SALUD": ["PROBLEMAS DE SALUD", "SALUD MENTAL", "PROBLEMAS DE SALUD PERSONAL", "PROBLEMAS PÉRSONALES Y DE SALUD", "PROBLEMA DE SALUD", "PROBLAMAS DE SALUD", "PROBLEMAS DE SALUD MENTAL", "EMBARAZO", "EMBARAZO DE ALTO RIESGO"],
                "PROBLEMAS FAMILIARES": ["PROBLEMA DE SALUD DE FAMILIAR", "PROBLEMA DE SALUD FAMILIAR", "PROBLEMAS DE SALUD DE FAMILIAR", "PROBLEMAS FAMILIARES (SALUD)", "SALUD FAMILIAR", "CUIDADO DE HIJO O FAMILIAR", "PROBLEMAS FAMILIARES", "PROBLEMAS FAMILIAR", "FALTA RED DE APOYO FAMILIAR", "FALTA DE REDES DE APOYO FAMILIAR", "FALTA DE REDES DE APOYO FAMILIARES", "FALTA DE RED DE APOYO FAMILIAR", "FALTAS DE REDES DE APOYO FAMILIAR", "FALTA DE REDE DE APOYO FAMILIAR", "FALTA DE REDES DE APOYO FAMILIAR (HIJO)", "FALTA DE REDES DE APOYO FAMILIAR (HIJOS)", "CUIDADO DE HIJO/A", "FALLECIMIENTO DE HIJO", "CUIDADO DE HIJOS", "CUIDADO DE HIJO", "CUIDADO DE HIJA", "FALTA DE REDES DE APOYO CON LOS HIJOS"],
                "TRASLADO": ["TRASLADO", "TRASLADO DE CIUDAD", "TRASLADO DE CUIDAD", "TRASLADO DE FAMILIA", "TRASLDO DE CIUDAD", "CAMBIO DE CUIDAD"],
                "SITUACIÓN ECONÓMICA": ["CAMBIO DE SITUACIÓN ECONÓMICA", "CAMBIO DE SITUACION ECONOMICA", "CAMBUIO DE SITUACIÓN ECONÓMICA", "CAMBIO DE SITUACIÓN ECONJÓMICA", "SIN APOYO ECONÓMICO", "CAMBIO D SITUACIÓN ECONÓMICA", "CAMBIO DE SITUACIÓN ACONÓMICA"],
                "JORNADA LABORAL": ["JORNADA DE TRABAJO", "INCOMPATIBILIDAD DE HORARIOS CON EL TRABAJO", "HORARIO LABORAL", "JORNADA LABORAL", "PROBLEMAS LABORALES", "PROBLEMAS LABORAL", "JORNADA LABORL", "INCOMPATIBILIDAD JORNADA LABORAL V/S ESTUDIOS", "LABORAL", "CAMBIO DE JORNADA LABORAL", "JORMADA LABORAL"],
                "PROBLEMAS PERSONALES": ["PROBLEMAS PERSONALES", "MOTIVOS PERSONALES", "PERSONALES", "PROBLEMA PERSONAL", "PERSONAL", "SITUACIÓN PERSONAL", "MPTIVOS PERSONALES", "CAMBIO DE PLAN", "POR CAMBIO DE PLAN"]
            }
            MAPEO_EXACTO = {normalizar_texto(v): k for k, l in LISTA_AGRUPACION.items() for v in l}

            # --- 1. Función de Extracción Profunda (Situación, Observación y Semestre) ---
            def parse_retiro(row):
                textos = []
                # Prioridad a Situación
                if 'SITUACIÓN' in row.index and pd.notna(row['SITUACIÓN']): textos.append(str(row['SITUACIÓN']).upper().strip())
                elif 'SITUACION' in row.index and pd.notna(row['SITUACION']): textos.append(str(row['SITUACION']).upper().strip())
                
                if 'OBSERVACIÓN' in row.index and pd.notna(row['OBSERVACIÓN']): textos.append(str(row['OBSERVACIÓN']).upper().strip())
                elif 'OBSERVACION' in row.index and pd.notna(row['OBSERVACION']): textos.append(str(row['OBSERVACION']).upper().strip())
                
                exclusiones = ["ADMINISTRATIVO", "REINCORPORACIÓN", "REINCORPORACION", "EGRESADO", "TITULADO", "FALLECIMIENTO"]

                for text in textos:
                    # Exclusión si detecta Administrativo u otras inválidas
                    if any(exc in text for exc in exclusiones):
                        return pd.Series([None, -1, None, None])

                    # Regla: Solo Definitivo o Temporal
                    if "RETIRO DEFINITIVO" in text or "RETIRO TEMPORAL" in text:
                        cat = "RETIRO DEFINITIVO" if "RETIRO DEFINITIVO" in text else "RETIRO TEMPORAL"
                        
                        matches = re.findall(r'\((.*?)\)', text)
                        if matches:
                            inside = matches[-1]
                            
                            # Extraer año (ej. 2021)
                            anio_match = re.search(r'(20\d{2})', inside)
                            anio_retiro = int(anio_match.group(1)) if anio_match else -1
                            
                            # Extraer semestre (ej. 1 SEM o 2 SEM)
                            sem_match = re.search(r'([12])\s*SEM', inside)
                            sem_retiro = f"{sem_match.group(1)}° Semestre" if sem_match else "No especificado"

                            # Extraer motivo
                            motivo_raw = inside.split(',', 1)[1].strip() if ',' in inside else inside
                            motivo_norm = normalizar_texto(motivo_raw)
                            motivo_final = MAPEO_EXACTO.get(motivo_norm, motivo_norm)
                            
                            return pd.Series([cat, anio_retiro, sem_retiro, motivo_final])
                return pd.Series([None, -1, None, None])

            # Aplicar la limpieza general
            df_filt[['TIPO_RETIRO', 'ANIO_RETIRO_EXTRACT', 'SEMESTRE_RETIRO', 'MOTIVO_LIMPIO']] = df_filt.apply(parse_retiro, axis=1)

            # --- 2. Cálculos por Año, Archivos y Cohortes ---
            anios_matricula_disp = df_matriculados['ANIO_MATRICULA'].unique() if not df_matriculados.empty else []

            for anio in sel_anios:
                for nombre_car in sel_nombres if sel_nombres else df_filt['NOMBRE_CARRERA'].unique():
                    for jor in sel_jornadas if sel_jornadas else df_filt['JORNADA'].unique():

                        # UNIVERSO 1ER AÑO: Mismo archivo de ingreso y matrícula
                        df_base = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'] == anio) & (df_filt['NOMBRE_CARRERA'] == nombre_car) & (df_filt['JORNADA'] == jor)]
                        ingresos = len(df_base)

                        if ingresos > 0:
                            ruts_iniciales = set(df_base['RUT'].dropna())

                            # ================= DESERCIÓN 1ER AÑO =================
                            # Solo declarados cuyo año en paréntesis coincida con el año de ingreso
                            df_declarados_1er = df_base[df_base['TIPO_RETIRO'].notna() & (df_base['ANIO_RETIRO_EXTRACT'] == anio)]
                            desertores_1er = df_declarados_1er['RUT'].nunique()
                            detalles_des_1er.append({"Cohorte": anio, "Carrera": nombre_car, "Jornada": jor, "Ingresos (1er Año)": ingresos, "Desertores al 1er Año": desertores_1er, "Tasa (%)": round((desertores_1er/ingresos*100), 1)})
                            
                            if not df_declarados_1er.empty: fugas_1er_list.append(df_declarados_1er)

                            # ================= DESERCIÓN 2DO AÑO =================
                            if (anio + 1) in anios_matricula_disp:
                                df_sig = df_matriculados[(df_matriculados['ANIO_INGRESO'] == anio) & (df_matriculados['ANIO_MATRICULA'] == anio + 1)]
                                ruts_sig = set(df_sig['RUT'].dropna())
                                retenidos_2do = len(ruts_iniciales & ruts_sig)
                                
                                desertores_2do_set = ruts_iniciales - ruts_sig
                                desertores_2do = len(desertores_2do_set)
                                detalles_des_2do.append({"Cohorte": anio, "Carrera": nombre_car, "Jornada": jor, "Ingresos (1er Año)": ingresos, "Desertores al 2do Año": desertores_2do, "Tasa (%)": round((desertores_2do/ingresos*100), 1)})

                                df_cohort_2do = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'].isin([anio, anio+1])) & (df_filt['NOMBRE_CARRERA'] == nombre_car) & (df_filt['JORNADA'] == jor)]
                                df_mot_2 = df_cohort_2do[df_cohort_2do['TIPO_RETIRO'].notna() & (df_cohort_2do['ANIO_RETIRO_EXTRACT'] == (anio + 1))]
                                if not df_mot_2.empty: fugas_2do_list.append(df_mot_2)

                                # ================= DESERCIÓN 3ER AÑO =================
                                if (anio + 2) in anios_matricula_disp:
                                    if retenidos_2do > 0:
                                        df_tercer = df_matriculados[(df_matriculados['ANIO_INGRESO'] == anio) & (df_matriculados['ANIO_MATRICULA'] == anio + 2)]
                                        ruts_tercer = set(df_tercer['RUT'].dropna())
                                        
                                        desertores_3er_set = (ruts_iniciales & ruts_sig) - ruts_tercer
                                        desertores_3er = len(desertores_3er_set)
                                        detalles_des_3er.append({"Cohorte": anio, "Carrera": nombre_car, "Jornada": jor, "Retenidos (2do Año)": retenidos_2do, "Desertores al 3er Año": desertores_3er, "Tasa (%)": round((desertores_3er/retenidos_2do*100), 1)})

                                        df_cohort_3er = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'].isin([anio+1, anio+2])) & (df_filt['NOMBRE_CARRERA'] == nombre_car) & (df_filt['JORNADA'] == jor)]
                                        df_mot_3 = df_cohort_3er[df_cohort_3er['TIPO_RETIRO'].notna() & (df_cohort_3er['ANIO_RETIRO_EXTRACT'] == (anio + 2))]
                                        if not df_mot_3.empty: fugas_3er_list.append(df_mot_3)

            df_det_des_1er = pd.DataFrame(detalles_des_1er)
            df_det_des_2do = pd.DataFrame(detalles_des_2do)
            df_det_des_3er = pd.DataFrame(detalles_des_3er)

            # Consolidar dataframes de motivos y eliminar RUTs duplicados
            def procesar_motivos(fugas_list):
                if not fugas_list: return pd.DataFrame()
                df_export = pd.concat(fugas_list, ignore_index=True)
                # SOLUCIÓN DE CONTABILIDAD: Mantener la misma granularidad del KPI
                return df_export.drop_duplicates(subset=['RUT', 'NOMBRE_CARRERA', 'JORNADA'], keep='first')

            df_fugas_1er_export = procesar_motivos(fugas_1er_list)
            df_fugas_2do_export = procesar_motivos(fugas_2do_list)
            df_fugas_3er_export = procesar_motivos(fugas_3er_list)

            # --- 3. Renderizado de Gráficos ---
            def mostrar_graficos_desercion(df_subset, titulo_pie, titulo_bar):
                if df_subset.empty:
                    st.info("No hay registros de motivos formalizados para esta selección y año.")
                    return
                df_cat = df_subset.groupby('TIPO_RETIRO').size().reset_index(name='Cantidad')
                df_mot = df_subset.groupby(['TIPO_RETIRO', 'MOTIVO_LIMPIO']).size().reset_index(name='Cantidad')
                tot_mot = df_subset.groupby('MOTIVO_LIMPIO').size().reset_index(name='Total').sort_values('Total', ascending=True).tail(10)
                df_mot = df_mot[df_mot['MOTIVO_LIMPIO'].isin(tot_mot['MOTIVO_LIMPIO'])].merge(tot_mot, on='MOTIVO_LIMPIO').sort_values('Total', ascending=True)

                c1, c2 = st.columns(2)
                with c1:
                    fig_p = px.pie(df_cat, names='TIPO_RETIRO', values='Cantidad', title=titulo_pie, color_discrete_sequence=px.colors.qualitative.Pastel)
                    fig_p.update_traces(textposition='inside', textinfo='percent+label')
                    fig_p.update_layout(showlegend=False, height=400)
                    st.plotly_chart(fig_p, use_container_width=True)
                with c2:
                    fig_b = px.bar(df_mot, x='Cantidad', y='MOTIVO_LIMPIO', color='TIPO_RETIRO', orientation='h', title=titulo_bar, color_discrete_sequence=COLORES_CLAROS, barmode='stack')
                    # SOLUCIÓN AL ORDEN: Se agregó yaxis={'categoryorder': 'total ascending'}
                    fig_b.update_layout(yaxis_title=None, showlegend=False, height=400, yaxis={'categoryorder': 'total ascending'})
                    st.plotly_chart(fig_b, use_container_width=True)

            # NUEVO GRÁFICO: Comparación entre Años (Cohortes) y Semestres
            def mostrar_grafico_semestral(df_subset, titulo):
                if df_subset.empty: return
                df_sem = df_subset.groupby(['ANIO_INGRESO', 'SEMESTRE_RETIRO']).size().reset_index(name='Cantidad')
                df_sem['ANIO_INGRESO'] = df_sem['ANIO_INGRESO'].astype(str)
                fig = px.bar(df_sem, x='ANIO_INGRESO', y='Cantidad', color='SEMESTRE_RETIRO', barmode='group',
                             title=titulo, color_discrete_sequence=[COLORES_CLAROS[1], COLORES_CLAROS[3], '#E5E7E9'], text='Cantidad')
                fig.update_layout(xaxis_title="Cohorte (Año de Ingreso)", yaxis_title="Cantidad de Retiros", xaxis={'type': 'category'})
                st.plotly_chart(fig, use_container_width=True)

            # =================== SUB-PESTAÑA: DESERCIÓN 1ER AÑO ===================
            with subtab_1er:
                if not df_det_des_1er.empty:
                    # SOLUCIÓN DE FILTRO: Aplicar '< max_anio_registrado' a todas las métricas para omitir la comparativa en curso
                    df_rep_1 = df_det_des_1er[df_det_des_1er['Cohorte'] < max_anio_registrado]
                    df_fugas_1er_filtrado = df_fugas_1er_export[df_fugas_1er_export['ANIO_INGRESO'] < max_anio_registrado] if not df_fugas_1er_export.empty else pd.DataFrame()

                    if not df_rep_1.empty:
                        col1, col2, col3 = st.columns(3)
                        col1.metric("Matriculados 1er Año", f"{int(df_rep_1['Ingresos (1er Año)'].sum())}")
                        col2.metric("Deserción Declarada (1er Año)", f"{int(df_rep_1['Desertores al 1er Año'].sum())}")
                        tasa_1er_glob = round(df_rep_1['Desertores al 1er Año'].sum() / df_rep_1['Ingresos (1er Año)'].sum() * 100, 1) if df_rep_1['Ingresos (1er Año)'].sum() > 0 else 0
                        col3.metric("Tasa Deserción", f"{tasa_1er_glob}%")
                        st.markdown("---")

                        res_des_1 = df_rep_1.groupby('Cohorte')[['Ingresos (1er Año)', 'Desertores al 1er Año']].sum().reset_index()
                        res_des_1['Cohorte'] = res_des_1['Cohorte'].astype(str)
                        df_melt_des_1 = res_des_1.melt(id_vars='Cohorte', value_vars=['Ingresos (1er Año)', 'Desertores al 1er Año'], var_name='Métrica', value_name='Cantidad')

                        fig3_1 = px.bar(df_melt_des_1, x="Cohorte", y="Cantidad", color="Métrica", barmode='group', color_discrete_sequence=[COLORES_CLAROS[0], COLORES_CLAROS[3]], title="Ingresos vs Desertores (Solo 1er Año)", text='Cantidad')
                        st.plotly_chart(fig3_1, use_container_width=True)
                        st.markdown("---")

                        # Uso del DataFrame filtrado
                        mostrar_graficos_desercion(df_fugas_1er_filtrado, "Tipos de Retiro (1er Año)", "Top 10 Motivos (1er Año)")
                        mostrar_grafico_semestral(df_fugas_1er_filtrado, "Comparativa Semestral de Retiros por Cohorte (1er Año)")

            # =================== SUB-PESTAÑA: DESERCIÓN 2DO AÑO ===================
            with subtab_2do:
                if not df_det_des_2do.empty:
                    df_rep_2 = df_det_des_2do[df_det_des_2do['Cohorte'] < max_anio_registrado]
                    # Aplicando el filtro uniforme a la base gráfica
                    df_fugas_2do_filtrado = df_fugas_2do_export[df_fugas_2do_export['ANIO_INGRESO'] < max_anio_registrado] if not df_fugas_2do_export.empty else pd.DataFrame()

                    if not df_rep_2.empty:
                        c1, c2, c3 = st.columns(3)
                        c1.metric("Ingresos Analizados", f"{int(df_rep_2['Ingresos (1er Año)'].sum())}")
                        c2.metric("Desertores Cruzados al 2do Año", f"{int(df_rep_2['Desertores al 2do Año'].sum())}")
                        c3.metric("Tasa Deserción", f"{round(df_rep_2['Desertores al 2do Año'].sum()/df_rep_2['Ingresos (1er Año)'].sum()*100, 1)}%")
                        st.markdown("---")

                        res_des_2 = df_rep_2.groupby('Cohorte')[['Ingresos (1er Año)', 'Desertores al 2do Año']].sum().reset_index()
                        res_des_2['Cohorte'] = res_des_2['Cohorte'].astype(str)
                        df_melt_des_2 = res_des_2.melt(id_vars='Cohorte', value_vars=['Ingresos (1er Año)', 'Desertores al 2do Año'], var_name='Métrica', value_name='Cantidad')

                        fig3_2 = px.bar(df_melt_des_2, x="Cohorte", y="Cantidad", color="Métrica", barmode='group', color_discrete_sequence=[COLORES_CLAROS[0], COLORES_CLAROS[3]], title="Ingresos vs Desertores al 2do Año", text='Cantidad')
                        st.plotly_chart(fig3_2, use_container_width=True)
                        st.markdown("---")
                        
                        # Uso del DataFrame filtrado
                        mostrar_graficos_desercion(df_fugas_2do_filtrado, "Tipos de Retiro (2do Año)", "Top 10 Motivos (2do Año)")
                        mostrar_grafico_semestral(df_fugas_2do_filtrado, "Comparativa Semestral de Retiros por Cohorte (2do Año)")

            # =================== SUB-PESTAÑA: DESERCIÓN 3ER AÑO ===================
            with subtab_3er:
                if not df_det_des_3er.empty:
                    df_rep_3 = df_det_des_3er[df_det_des_3er['Cohorte'] <= (max_anio_registrado - 2)]
                    # Aplicando el filtro uniforme a la base gráfica
                    df_fugas_3er_filtrado = df_fugas_3er_export[df_fugas_3er_export['ANIO_INGRESO'] <= (max_anio_registrado - 2)] if not df_fugas_3er_export.empty else pd.DataFrame()

                    if not df_rep_3.empty:
                        c1, c2, c3 = st.columns(3)
                        c1.metric("Retenidos en 2do Año", f"{int(df_rep_3['Retenidos (2do Año)'].sum())}")
                        c2.metric("Desertores Cruzados al 3er Año", f"{int(df_rep_3['Desertores al 3er Año'].sum())}")
                        c3.metric("Tasa Deserción", f"{round(df_rep_3['Desertores al 3er Año'].sum()/df_rep_3['Retenidos (2do Año)'].sum()*100, 1)}%")
                        st.markdown("---")

                        res_des_3 = df_rep_3.groupby('Cohorte')[['Retenidos (2do Año)', 'Desertores al 3er Año']].sum().reset_index()
                        res_des_3['Cohorte'] = res_des_3['Cohorte'].astype(str)
                        df_melt_des_3 = res_des_3.melt(id_vars='Cohorte', value_vars=['Retenidos (2do Año)', 'Desertores al 3er Año'], var_name='Métrica', value_name='Cantidad')

                        fig3_3 = px.bar(df_melt_des_3, x="Cohorte", y="Cantidad", color="Métrica", barmode='group', color_discrete_sequence=[COLORES_CLAROS[2], COLORES_CLAROS[3]], title="Retenidos (2do Año) vs Desertores al 3er Año", text='Cantidad')
                        st.plotly_chart(fig3_3, use_container_width=True)
                        st.markdown("---")

                        # Uso del DataFrame filtrado
                        mostrar_graficos_desercion(df_fugas_3er_filtrado, "Tipos de Retiro (3er Año)", "Top 10 Motivos (3er Año)")
                        mostrar_grafico_semestral(df_fugas_3er_filtrado, "Comparativa Semestral de Retiros por Cohorte (3er Año)")

# ----------------- TAB 4: CONVALIDACIÓN Y APOYO -----------------
with tab4:
    st.header("Análisis de Convalidaciones y Apoyos (Gestión CFTe-AP)")
    
    if df_convalidaciones.empty:
        st.warning("⚠️ No se encontró el archivo de Convalidaciones.")
    else:
        # Filtro de AÑO y CARRERA aplicado directamente al archivo de convalidaciones
        df_conv_analisis = df_convalidaciones.copy()
        if sel_anios:
            df_conv_analisis = df_conv_analisis[df_conv_analisis['ANIO_BENEFICIO'].isin(sel_anios)]
        if sel_nombres:
            df_conv_analisis = df_conv_analisis[df_conv_analisis['NOMBRE_CARRERA_CONV'].isin(sel_nombres)]
        
        col_tipo = 'UNNAMED: 6' if 'UNNAMED: 6' in df_conv_analisis.columns else None
        
        if col_tipo:
            # --- 1. SECCIÓN TUTORÍAS (Conteo Total por Año de Uso) ---
            st.subheader("💡 Apoyo Estudiantil: Tutorías Académicas")
            df_tut = df_conv_analisis[df_conv_analisis[col_tipo].astype(str).str.contains("TUTOR", case=False, na=False)].copy()
            if not df_tut.empty:
                df_tut_graph = df_tut.groupby('ANIO_BENEFICIO').size().reset_index(name='Cantidad')
                df_tut_graph = df_tut_graph.sort_values('ANIO_BENEFICIO')
                total_tutorias_cont = len(df_tut)
                
                c1, c2 = st.columns([1, 2])
                c1.metric("Total de Tutorías", total_tutorias_cont)
                c1.markdown("*Incluye registros anteriores a 2021 si los hay.*")
                
                fig_tut = px.bar(df_tut_graph, x='ANIO_BENEFICIO', y='Cantidad', text='Cantidad',
                                 title="Intervenciones de Tutoría por Año", color_discrete_sequence=[COLORES_CLAROS[1]])
                fig_tut.update_layout(xaxis_title="Año en curso")
                c2.plotly_chart(fig_tut, use_container_width=True)
            else:
                st.info("No hay registros de tutorías para los filtros seleccionados.")

            st.markdown("---")

            # --- 2. ACORTAMIENTO DE DURACIÓN (RUT Único por Categoría) ---
            st.subheader("🎓 Estudiantes con Beneficio de Acortamiento de Estudios")
            st.info("Nota: Este análisis no considera el filtro de Jornada. Excluye años anteriores a 2021.")
            
            # Filtro para Acortamiento: Excluir operaciones previas a 2021
            df_conv_2021 = df_conv_analisis[df_conv_analisis['ANIO_BENEFICIO'] >= 2021].copy()
            
            def agrupar_beneficios(val):
                val = str(val).strip().upper()
                if val in ["CONVALIDACIÓN EXTERNA", "CONVALIDACIÓN UTA", "CONVALIDACION EXTERNA", "CONVALIDACION UTA"]: 
                    return "CONVALIDACIÓN"
                if val in ["CONVALIDACIÓN INTERNA", "HOMOLOGACIÓN", "CONVALIDACION INTERNA", "HOMOLOGACION"]: 
                    return "HOMOLOGACIÓN"
                if val in ["ARTICULACIÓN", "ARTICULACION"]: return "ARTICULACIÓN"
                if "R.A.P" in val: return "R.A.P"
                return "EXCLUIR"

            df_conv_2021['CAT_AGRUPADA'] = df_conv_2021[col_tipo].map(agrupar_beneficios)
            df_acort = df_conv_2021[df_conv_2021['CAT_AGRUPADA'] != "EXCLUIR"].copy()
            
            # RUT Único por Categoría para el conteo de estudiantes
            df_acort_unicos = df_acort.drop_duplicates(subset=['RUT', 'CAT_AGRUPADA'])
            
            if not df_acort_unicos.empty:
                col_g1, col_g2 = st.columns(2)
                
                with col_g1:
                    # Distribución Global
                    df_dist_acort = df_acort_unicos.groupby('CAT_AGRUPADA').size().reset_index(name='Estudiantes')
                    df_dist_acort = df_dist_acort.sort_values('Estudiantes', ascending=False)
                    fig_dist = px.bar(df_dist_acort, x='CAT_AGRUPADA', y='Estudiantes', text='Estudiantes',
                                      title="Distribución Total de Beneficios", color='CAT_AGRUPADA',
                                      color_discrete_sequence=px.colors.qualitative.Pastel)
                    fig_dist.update_layout(showlegend=False, xaxis_title=None)
                    st.plotly_chart(fig_dist, use_container_width=True)

                with col_g2:
                    # Comparativa por Año de Beneficio
                    df_hist_acort = df_acort_unicos.groupby(['ANIO_BENEFICIO', 'CAT_AGRUPADA']).size().reset_index(name='Estudiantes')
                    df_hist_acort['ANIO_BENEFICIO'] = df_hist_acort['ANIO_BENEFICIO'].astype(str)
                    fig_hist = px.bar(df_hist_acort, x='ANIO_BENEFICIO', y='Estudiantes', color='CAT_AGRUPADA',
                                      barmode='group', text='Estudiantes', title="Evolución por Año de Otorgamiento",
                                      color_discrete_sequence=px.colors.qualitative.Pastel)
                    fig_hist.update_layout(xaxis_title="Año del Beneficio")
                    st.plotly_chart(fig_hist, use_container_width=True)
            else:
                st.info("No se encontraron registros de acortamiento para las categorías y filtros seleccionados.")
        else:
            st.error("No se encontró la columna de beneficios (Unnamed: 6).")

# ----------------- TAB 5: PROGRESIÓN -----------------
with tab5:
    st.header("Tasa de Progresión (2do a 3er Año)")
    st.info("📌 Esta métrica compara los estudiantes que fueron retenidos en su segundo año (Año X+1) versus aquellos de esa misma cohorte que continúan matriculados en el tercer año (Año X+2).")
    detalles_prog = []

    for anio in sel_anios:
        for nombre_car in sel_nombres if sel_nombres else df_filt['NOMBRE_CARRERA'].unique():
            for jor in sel_jornadas if sel_jornadas else df_filt['JORNADA'].unique():
                # Base: Universo Real de 1er año
                df_base = df_filt[(df_filt['ANIO_INGRESO'] == anio) & (df_filt['ANIO_MATRICULA'] == anio) & (df_filt['NOMBRE_CARRERA'] == nombre_car) & (df_filt['JORNADA'] == jor)]
                ingresos = len(df_base)

                if ingresos > 0:
                    ruts_iniciales = set(df_base['RUT'].dropna())
                    
                    # Retenidos (2do año): Estudiantes matriculados en Año + 1
                    df_segundo_anio = df_matriculados[(df_matriculados['ANIO_INGRESO'] == anio) & (df_matriculados['ANIO_MATRICULA'] == anio + 1)]
                    ruts_segundo_anio = set(df_segundo_anio['RUT'].dropna())
                    
                    # Intersección para obtener los retenidos reales
                    retenidos = len(ruts_iniciales & ruts_segundo_anio)
                    
                    # Tercer año: Estudiantes matriculados en Año + 2
                    df_tercer_anio = df_matriculados[(df_matriculados['ANIO_INGRESO'] == anio) & (df_matriculados['ANIO_MATRICULA'] == anio + 2)]
                    ruts_tercer_anio = set(df_tercer_anio['RUT'].dropna())
                    
                    # Progresados: Estudiantes iniciales que estuvieron en 2do año y llegaron a 3er año
                    progresados = len(ruts_iniciales & ruts_segundo_anio & ruts_tercer_anio)

                    # Cálculo de la tasa sobre los retenidos, no sobre los ingresos
                    tasa_prog = round((progresados / retenidos * 100), 1) if retenidos > 0 else 0.0

                    detalles_prog.append({
                        "Cohorte": anio,
                        "Carrera": nombre_car,
                        "Jornada": jor,
                        "Retenidos (2do Año)": retenidos,
                        "Progresión (3er Año)": progresados,
                        "Tasa Progresión (%)": tasa_prog
                    })

    df_det_prog = pd.DataFrame(detalles_prog)

    if not df_det_prog.empty:
        # 1. FILTRO DE CORTE: Solo cohortes que ya tienen año de comparación para 3er año (Año + 2)
        df_det_prog = df_det_prog[df_det_prog['Cohorte'] <= (max_anio_registrado - 2)]
        
        # Validación extra: Verificar que queden datos después de filtrar los años recientes
        if not df_det_prog.empty:
            # 2. Recálculo de totales con los datos correctos
            global_prog_retenidos = df_det_prog['Retenidos (2do Año)'].sum()
            global_prog_progresados = df_det_prog['Progresión (3er Año)'].sum()
            tasa_tot_prog = round(global_prog_progresados / global_prog_retenidos * 100, 1) if global_prog_retenidos > 0 else 0
            promedio_tasa_prog = round(df_det_prog["Tasa Progresión (%)"].mean(), 1)

            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Retenidos 2do Año", f"{int(global_prog_retenidos)}")
            col2.metric("Alcanzan 3er Año", f"{int(global_prog_progresados)}")
            col3.metric("Tasa Global Acumulada", f"{tasa_tot_prog}%")
            col4.metric("Promedio de las Tasas", f"{promedio_tasa_prog}%")
            st.markdown("---")

            res_prog = df_det_prog.groupby('Cohorte')[['Retenidos (2do Año)', 'Progresión (3er Año)']].sum().reset_index()
            res_prog['Cohorte'] = res_prog['Cohorte'].astype(str)
            df_melt_prog = res_prog.melt(id_vars='Cohorte', value_vars=['Retenidos (2do Año)', 'Progresión (3er Año)'], var_name='Métrica', value_name='Cantidad')
            
            fig_prog = px.bar(df_melt_prog, x="Cohorte", y="Cantidad", color="Métrica", barmode='group', 
                              color_discrete_sequence=[COLORES_CLAROS[2], '#D7BDE2'], 
                              title="Comparativa Global Retenidos vs Progresión a 3er Año", text='Cantidad')
            st.plotly_chart(fig_prog, use_container_width=True)
        else:
            st.info("⚠️ Los filtros actuales corresponden a cohortes recientes que aún no cumplen el ciclo de 3 años para medir la progresión.")
    else:
        st.warning("⚠️ No se encontraron datos suficientes para calcular la progresión al tercer año con los filtros actuales.")

# --- 7. GENERACIÓN DEL INFORME WORD ---
st.markdown("---")
st.header("📄 Informe de Análisis Institucional Exportable")

if DOCX_DISPONIBLE:
    def crear_documento_word():
        doc = Document()
        title = doc.add_heading('Informe de Análisis Institucional CFTe-AP', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Fecha de generación: {datetime.datetime.now().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Representatividad
        doc.add_heading('1. Resumen de Representatividad (Desde 2021)', level=1)
        
        df_hist_real = df_matriculados[df_matriculados['ANIO_INGRESO'] == df_matriculados['ANIO_MATRICULA']] if not df_matriculados.empty else []
        doc.add_paragraph(f"Total Histórico CFTe-AP de 1er Año (Sin filtro): {len(df_hist_real)} registros")
        doc.add_paragraph(f"Universo Analizado de 1er Año (Filtro Actual): {len(df_mat_real)} registros")
        
        if not df_mat_real.empty:
            doc.add_heading('Distribución Global por Jornada (1er Año)', level=2)
            df_jor = df_mat_real['JORNADA'].value_counts().reset_index()
            df_jor.columns = ['Jornada', 'Cantidad']
            df_jor['%'] = round((df_jor['Cantidad'] / df_jor['Cantidad'].sum() * 100), 1)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Jornada', 'Cantidad', '%'
            for _, row in df_jor.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text, row_cells[1].text, row_cells[2].text = str(row['Jornada']), str(row['Cantidad']), f"{row['%']}%"

        # =========================================================================
            # 2. TASA DE MATRÍCULA (OFERTA VS REAL)
            # =========================================================================
            if not df_det_mat.empty:
                doc.add_heading('2. Tasa de Matrícula (Oferta vs Real)', level=1)
                for anio in sorted(df_det_mat['Año'].unique()):
                    doc.add_heading(f'Año: {int(anio)}', level=2)
                    df_anio = df_det_mat[df_det_mat['Año'] == anio]

                    # --- NUEVO: Resumen Consolidado por Carrera ---
                    doc.add_heading('Resumen Consolidado por Carrera', level=3)
                    df_carrera = df_anio.groupby('Carrera')[['Vacantes', 'Matriculados']].sum().reset_index()
                    
                    table_car = doc.add_table(rows=1, cols=4)
                    table_car.style = 'Table Grid'
                    hdr_car = table_car.rows[0].cells
                    hdr_car[0].text, hdr_car[1].text, hdr_car[2].text, hdr_car[3].text = 'Carrera', 'Vacantes', 'Matriculados', 'Tasa (%)'
                    
                    for cell in hdr_car:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs: run.bold = True
                                
                    tot_vac_c, tot_mat_c = 0, 0
                    for _, row in df_carrera.iterrows():
                        cells = table_car.add_row().cells
                        cells[0].text = str(row['Carrera'])
                        cells[1].text = str(int(row['Vacantes']))
                        cells[2].text = str(int(row['Matriculados']))
                        tasa_c = round((row['Matriculados'] / row['Vacantes'] * 100), 1) if row['Vacantes'] > 0 else 0
                        cells[3].text = f"{tasa_c}%"
                        tot_vac_c += row['Vacantes']
                        tot_mat_c += row['Matriculados']
                        
                    row_tot_c = table_car.add_row().cells
                    row_tot_c[0].text = "TOTAL AÑO"
                    row_tot_c[1].text = str(int(tot_vac_c))
                    row_tot_c[2].text = str(int(tot_mat_c))
                    row_tot_c[3].text = f"{round((tot_mat_c/tot_vac_c*100) if tot_vac_c > 0 else 0, 1)}%"
                    for cell in row_tot_c:
                        for paragraph in cell.paragraphs:
                            if paragraph.runs: paragraph.runs[0].bold = True
                    doc.add_paragraph()

                    # --- ORIGINAL: Detalle por Jornada ---
                    doc.add_heading('Detalle por Jornada', level=3)
                    for jor in sorted(df_anio['Jornada'].unique()):
                        doc.add_heading(f'Jornada: {jor}', level=4)
                        df_jor_data = df_anio[df_anio['Jornada'] == jor].sort_values(by='Carrera')

                        table = doc.add_table(rows=1, cols=4)
                        table.style = 'Table Grid'
                        hdr = table.rows[0].cells
                        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Carrera', 'Vacantes', 'Matriculados', 'Tasa (%)'

                        for cell in hdr:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs: paragraph.runs[0].bold = True

                        tot_vac, tot_mat = 0, 0
                        for _, row in df_jor_data.iterrows():
                            cells = table.add_row().cells
                            cells[0].text = str(row['Carrera'])
                            cells[1].text = str(int(row['Vacantes']))
                            cells[2].text = str(int(row['Matriculados']))
                            cells[3].text = f"{row['Tasa Matrícula (%)']}%"
                            tot_vac += row['Vacantes']
                            tot_mat += row['Matriculados']

                        row_tot = table.add_row().cells
                        row_tot[0].text = "SUBTOTAL JORNADA"
                        row_tot[1].text = str(int(tot_vac))
                        row_tot[2].text = str(int(tot_mat))
                        row_tot[3].text = f"{round((tot_mat/tot_vac*100) if tot_vac > 0 else 0, 1)}%"
                        for cell in row_tot:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs: paragraph.runs[0].bold = True
                        doc.add_paragraph()

            # =========================================================================
            # 3. TASA DE RETENCIÓN (1ER A 2DO AÑO)
            # =========================================================================
            if not df_det_ret.empty:
                df_rep_ret = df_det_ret[df_det_ret['Cohorte'] < max_anio_registrado]
                if not df_rep_ret.empty:
                    doc.add_heading('3. Tasa de Retención (1er a 2do Año)', level=1)
                    for anio in sorted(df_rep_ret['Cohorte'].unique()):
                        doc.add_heading(f'Cohorte: {int(anio)}', level=2)
                        df_anio = df_rep_ret[df_rep_ret['Cohorte'] == anio]

                        # --- NUEVO: Resumen Consolidado por Carrera ---
                        doc.add_heading('Resumen Consolidado por Carrera', level=3)
                        df_carrera = df_anio.groupby('Carrera')[['Ingresos (1er Año)', 'Retenidos (2do Año)']].sum().reset_index()
                        
                        table_car = doc.add_table(rows=1, cols=4)
                        table_car.style = 'Table Grid'
                        hdr_car = table_car.rows[0].cells
                        hdr_car[0].text, hdr_car[1].text, hdr_car[2].text, hdr_car[3].text = 'Carrera', 'Ingresos 1er Año', 'Retenidos 2do Año', 'Tasa (%)'
                        
                        for cell in hdr_car:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs: run.bold = True
                                    
                        tot_ing_c, tot_ret_c = 0, 0
                        for _, row in df_carrera.iterrows():
                            cells = table_car.add_row().cells
                            cells[0].text = str(row['Carrera'])
                            cells[1].text = str(int(row['Ingresos (1er Año)']))
                            cells[2].text = str(int(row['Retenidos (2do Año)']))
                            tasa_c = round((row['Retenidos (2do Año)'] / row['Ingresos (1er Año)'] * 100), 1) if row['Ingresos (1er Año)'] > 0 else 0
                            cells[3].text = f"{tasa_c}%"
                            tot_ing_c += row['Ingresos (1er Año)']
                            tot_ret_c += row['Retenidos (2do Año)']
                            
                        row_tot_c = table_car.add_row().cells
                        row_tot_c[0].text = "TOTAL COHORTE"
                        row_tot_c[1].text = str(int(tot_ing_c))
                        row_tot_c[2].text = str(int(tot_ret_c))
                        row_tot_c[3].text = f"{round((tot_ret_c/tot_ing_c*100) if tot_ing_c > 0 else 0, 1)}%"
                        for cell in row_tot_c:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs: paragraph.runs[0].bold = True
                        doc.add_paragraph()

                        # --- ORIGINAL: Detalle por Jornada ---
                        doc.add_heading('Detalle por Jornada', level=3)
                        for jor in sorted(df_anio['Jornada'].unique()):
                            doc.add_heading(f'Jornada: {jor}', level=4)
                            df_jor_data = df_anio[df_anio['Jornada'] == jor].sort_values(by='Carrera')

                            table = doc.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            hdr = table.rows[0].cells
                            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Carrera', 'Ingresos 1er Año', 'Retenidos', 'Tasa (%)'

                            for cell in hdr:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True

                            tot_ing, tot_ret = 0, 0
                            for _, row in df_jor_data.iterrows():
                                cells = table.add_row().cells
                                cells[0].text = str(row['Carrera'])
                                cells[1].text = str(int(row['Ingresos (1er Año)']))
                                cells[2].text = str(int(row['Retenidos (2do Año)']))
                                cells[3].text = f"{row['Tasa Retención (%)']}%"
                                tot_ing += row['Ingresos (1er Año)']
                                tot_ret += row['Retenidos (2do Año)']

                            row_tot = table.add_row().cells
                            row_tot[0].text = "SUBTOTAL JORNADA"
                            row_tot[1].text = str(int(tot_ing))
                            row_tot[2].text = str(int(tot_ret))
                            row_tot[3].text = f"{round((tot_ret/tot_ing*100) if tot_ing > 0 else 0, 1)}%"
                            for cell in row_tot:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True
                            doc.add_paragraph()

            # =========================================================================
            # 4. TASA DE DESERCIÓN (CON AJUSTES PREVIOS)
            # =========================================================================
            if not df_det_des_1er.empty:
                df_rep_des_1 = df_det_des_1er[df_det_des_1er['Cohorte'] < max_anio_registrado]
                if not df_rep_des_1.empty:
                    doc.add_heading('4. Tasa de Deserción', level=1)
                    doc.add_heading('4.1 Deserción al 1er Año', level=2)
                    for anio in sorted(df_rep_des_1['Cohorte'].unique()):
                        doc.add_heading(f'Cohorte: {int(anio)}', level=3)
                        df_anio = df_rep_des_1[df_rep_des_1['Cohorte'] == anio]

                        doc.add_heading('Resumen Consolidado por Carrera', level=4)
                        df_carrera = df_anio.groupby('Carrera')[['Ingresos (1er Año)', 'Desertores al 1er Año']].sum().reset_index()
                        
                        table_car = doc.add_table(rows=1, cols=4)
                        table_car.style = 'Table Grid'
                        hdr_car = table_car.rows[0].cells
                        hdr_car[0].text, hdr_car[1].text, hdr_car[2].text, hdr_car[3].text = 'Carrera', 'Ingresos 1er Año', 'Desertores 1er Año', 'Tasa (%)'
                        
                        for cell in hdr_car:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs: run.bold = True
                                    
                        tot_ing_c, tot_des_c = 0, 0
                        for _, row in df_carrera.iterrows():
                            cells = table_car.add_row().cells
                            cells[0].text = str(row['Carrera'])
                            cells[1].text = str(int(row['Ingresos (1er Año)']))
                            cells[2].text = str(int(row['Desertores al 1er Año']))
                            tasa_c = round((row['Desertores al 1er Año'] / row['Ingresos (1er Año)'] * 100), 1) if row['Ingresos (1er Año)'] > 0 else 0
                            cells[3].text = f"{tasa_c}%"
                            tot_ing_c += row['Ingresos (1er Año)']
                            tot_des_c += row['Desertores al 1er Año']
                            
                        row_tot_c = table_car.add_row().cells
                        row_tot_c[0].text = "TOTAL COHORTE"
                        row_tot_c[1].text = str(int(tot_ing_c))
                        row_tot_c[2].text = str(int(tot_des_c))
                        row_tot_c[3].text = f"{round((tot_des_c/tot_ing_c*100) if tot_ing_c > 0 else 0, 1)}%"
                        for cell in row_tot_c:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs: paragraph.runs[0].bold = True
                        doc.add_paragraph()

                        doc.add_heading('Detalle por Jornada', level=4)
                        for jor in sorted(df_anio['Jornada'].unique()):
                            doc.add_heading(f'Jornada: {jor}', level=5)
                            df_jor_data = df_anio[df_anio['Jornada'] == jor].sort_values(by='Carrera')
                            
                            table = doc.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            hdr = table.rows[0].cells
                            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Carrera', 'Ingresos 1er Año', 'Desertores 1er Año', 'Tasa (%)'
                            
                            for cell in hdr:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs: run.bold = True
                                        
                            tot_ing, tot_des = 0, 0
                            for _, row in df_jor_data.iterrows():
                                cells = table.add_row().cells
                                cells[0].text = str(row['Carrera'])
                                cells[1].text = str(int(row['Ingresos (1er Año)']))
                                cells[2].text = str(int(row['Desertores al 1er Año']))
                                cells[3].text = f"{row['Tasa (%)']}%"
                                tot_ing += row['Ingresos (1er Año)']
                                tot_des += row['Desertores al 1er Año']
                                
                            row_tot = table.add_row().cells
                            row_tot[0].text = "SUBTOTAL JORNADA"
                            row_tot[1].text = str(int(tot_ing))
                            row_tot[2].text = str(int(tot_des))
                            row_tot[3].text = f"{round((tot_des/tot_ing*100) if tot_ing > 0 else 0, 1)}%"
                            for cell in row_tot:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True
                            doc.add_paragraph()

            if not df_det_des_2do.empty:
                df_rep_des_2 = df_det_des_2do[df_det_des_2do['Cohorte'] < max_anio_registrado]
                if not df_rep_des_2.empty:
                    doc.add_heading('4.2 Deserción al 2do Año', level=2)
                    for anio in sorted(df_rep_des_2['Cohorte'].unique()):
                        doc.add_heading(f'Cohorte: {int(anio)}', level=3)
                        df_anio = df_rep_des_2[df_rep_des_2['Cohorte'] == anio]

                        doc.add_heading('Resumen Consolidado por Carrera', level=4)
                        df_carrera = df_anio.groupby('Carrera')[['Ingresos (1er Año)', 'Desertores al 2do Año']].sum().reset_index()
                        
                        table_car = doc.add_table(rows=1, cols=4)
                        table_car.style = 'Table Grid'
                        hdr_car = table_car.rows[0].cells
                        hdr_car[0].text, hdr_car[1].text, hdr_car[2].text, hdr_car[3].text = 'Carrera', 'Ingresos 1er Año', 'Desertores 2do Año', 'Tasa (%)'
                        
                        for cell in hdr_car:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs: run.bold = True
                                    
                        tot_ing_c, tot_des_c = 0, 0
                        for _, row in df_carrera.iterrows():
                            cells = table_car.add_row().cells
                            cells[0].text = str(row['Carrera'])
                            cells[1].text = str(int(row['Ingresos (1er Año)']))
                            cells[2].text = str(int(row['Desertores al 2do Año']))
                            tasa_c = round((row['Desertores al 2do Año'] / row['Ingresos (1er Año)'] * 100), 1) if row['Ingresos (1er Año)'] > 0 else 0
                            cells[3].text = f"{tasa_c}%"
                            tot_ing_c += row['Ingresos (1er Año)']
                            tot_des_c += row['Desertores al 2do Año']
                            
                        row_tot_c = table_car.add_row().cells
                        row_tot_c[0].text = "TOTAL COHORTE"
                        row_tot_c[1].text = str(int(tot_ing_c))
                        row_tot_c[2].text = str(int(tot_des_c))
                        row_tot_c[3].text = f"{round((tot_des_c/tot_ing_c*100) if tot_ing_c > 0 else 0, 1)}%"
                        for cell in row_tot_c:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs: paragraph.runs[0].bold = True
                        doc.add_paragraph()

                        doc.add_heading('Detalle por Jornada', level=4)
                        for jor in sorted(df_anio['Jornada'].unique()):
                            doc.add_heading(f'Jornada: {jor}', level=5)
                            df_jor_data = df_anio[df_anio['Jornada'] == jor].sort_values(by='Carrera')
                            
                            table = doc.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            hdr = table.rows[0].cells
                            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Carrera', 'Ingresos 1er Año', 'Desertores 2do Año', 'Tasa (%)'
                            
                            for cell in hdr:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs: run.bold = True
                                        
                            tot_ing, tot_des = 0, 0
                            for _, row in df_jor_data.iterrows():
                                cells = table.add_row().cells
                                cells[0].text = str(row['Carrera'])
                                cells[1].text = str(int(row['Ingresos (1er Año)']))
                                cells[2].text = str(int(row['Desertores al 2do Año']))
                                cells[3].text = f"{row['Tasa (%)']}%"
                                tot_ing += row['Ingresos (1er Año)']
                                tot_des += row['Desertores al 2do Año']
                                
                            row_tot = table.add_row().cells
                            row_tot[0].text = "SUBTOTAL JORNADA"
                            row_tot[1].text = str(int(tot_ing))
                            row_tot[2].text = str(int(tot_des))
                            row_tot[3].text = f"{round((tot_des/tot_ing*100) if tot_ing > 0 else 0, 1)}%"
                            for cell in row_tot:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True
                            doc.add_paragraph()

            if not df_det_des_3er.empty:
                df_rep_des_3 = df_det_des_3er[df_det_des_3er['Cohorte'] <= (max_anio_registrado - 2)]
                if not df_rep_des_3.empty:
                    doc.add_heading('4.3 Deserción al 3er Año', level=2)
                    for anio in sorted(df_rep_des_3['Cohorte'].unique()):
                        doc.add_heading(f'Cohorte: {int(anio)}', level=3)
                        df_anio = df_rep_des_3[df_rep_des_3['Cohorte'] == anio]

                        doc.add_heading('Resumen Consolidado por Carrera', level=4)
                        df_carrera = df_anio.groupby('Carrera')[['Retenidos (2do Año)', 'Desertores al 3er Año']].sum().reset_index()
                        
                        table_car = doc.add_table(rows=1, cols=4)
                        table_car.style = 'Table Grid'
                        hdr_car = table_car.rows[0].cells
                        hdr_car[0].text, hdr_car[1].text, hdr_car[2].text, hdr_car[3].text = 'Carrera', 'Retenidos 2do Año', 'Desertores 3er Año', 'Tasa (%)'
                        
                        for cell in hdr_car:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs: run.bold = True
                                    
                        tot_ret_c, tot_des_c = 0, 0
                        for _, row in df_carrera.iterrows():
                            cells = table_car.add_row().cells
                            cells[0].text = str(row['Carrera'])
                            cells[1].text = str(int(row['Retenidos (2do Año)']))
                            cells[2].text = str(int(row['Desertores al 3er Año']))
                            tasa_c = round((row['Desertores al 3er Año'] / row['Retenidos (2do Año)'] * 100), 1) if row['Retenidos (2do Año)'] > 0 else 0
                            cells[3].text = f"{tasa_c}%"
                            tot_ret_c += row['Retenidos (2do Año)']
                            tot_des_c += row['Desertores al 3er Año']
                            
                        row_tot_c = table_car.add_row().cells
                        row_tot_c[0].text = "TOTAL COHORTE"
                        row_tot_c[1].text = str(int(tot_ret_c))
                        row_tot_c[2].text = str(int(tot_des_c))
                        row_tot_c[3].text = f"{round((tot_des_c/tot_ret_c*100) if tot_ret_c > 0 else 0, 1)}%"
                        for cell in row_tot_c:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs: paragraph.runs[0].bold = True
                        doc.add_paragraph()

                        doc.add_heading('Detalle por Jornada', level=4)
                        for jor in sorted(df_anio['Jornada'].unique()):
                            doc.add_heading(f'Jornada: {jor}', level=5)
                            df_jor_data = df_anio[df_anio['Jornada'] == jor].sort_values(by='Carrera')
                            
                            table = doc.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            hdr = table.rows[0].cells
                            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Carrera', 'Retenidos 2do Año', 'Desertores 3er Año', 'Tasa (%)'
                            
                            for cell in hdr:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs: run.bold = True
                                        
                            tot_ret, tot_des = 0, 0
                            for _, row in df_jor_data.iterrows():
                                cells = table.add_row().cells
                                cells[0].text = str(row['Carrera'])
                                cells[1].text = str(int(row['Retenidos (2do Año)']))
                                cells[2].text = str(int(row['Desertores al 3er Año']))
                                cells[3].text = f"{row['Tasa (%)']}%"
                                tot_ret += row['Retenidos (2do Año)']
                                tot_des += row['Desertores al 3er Año']
                                
                            row_tot = table.add_row().cells
                            row_tot[0].text = "SUBTOTAL JORNADA"
                            row_tot[1].text = str(int(tot_ret))
                            row_tot[2].text = str(int(tot_des))
                            row_tot[3].text = f"{round((tot_des/tot_ret*100) if tot_ret > 0 else 0, 1)}%"
                            for cell in row_tot:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True
                            doc.add_paragraph()

            # =========================================================================
            # 5. ANEXO MOTIVOS DE RETIRO (CON TABLAS DETALLADAS DE DOBLE ENTRADA)
            # =========================================================================
            doc.add_heading('5. Anexo: Motivos de Retiro Clasificados', level=1)

            def generar_anexo_motivos(doc, df_fugas, titulo_seccion, limite_anio):
                # Copia segura para manipular columnas sin alterar los dataframes originales
                df_valido = df_fugas[df_fugas['ANIO_INGRESO'] <= limite_anio].copy() if not df_fugas.empty else pd.DataFrame()
                
                if not df_valido.empty:
                    doc.add_heading(titulo_seccion, level=2)
                    doc.add_paragraph(f"Total de estudiantes con retiro formalizado: {len(df_valido)}")
                    
                    # --- EXTRAER SEMESTRE DESDE HISTORIAL/OBSERVACIONES ---
                    def extraer_semestre(row):
                        textos = []
                        if 'SITUACIÓN' in row.index and pd.notna(row['SITUACIÓN']): textos.append(str(row['SITUACIÓN']).upper())
                        elif 'SITUACION' in row.index and pd.notna(row['SITUACION']): textos.append(str(row['SITUACION']).upper())
                        if 'OBSERVACIÓN' in row.index and pd.notna(row['OBSERVACIÓN']): textos.append(str(row['OBSERVACIÓN']).upper())
                        elif 'OBSERVACION' in row.index and pd.notna(row['OBSERVACION']): textos.append(str(row['OBSERVACION']).upper())
                        
                        for text in textos:
                            if re.search(r'(1ER|PRIMER|1°|1)\s*SEM', text) or re.search(r'\(\s*20\d{2}\s*,\s*1\s*,', text):
                                return '1er Semestre'
                            if re.search(r'(2DO|SEGUNDO|2°|2)\s*SEM', text) or re.search(r'\(\s*20\d{2}\s*,\s*2\s*,', text):
                                return '2do Semestre'
                        return 'No Especificado'

                    df_valido['SEMESTRE_EXTRACT'] = df_valido.apply(extraer_semestre, axis=1)
                    
                    # Forzar consistencia de nombres en la jornada
                    if 'JORNADA' in df_valido.columns:
                        df_valido['JORNADA_NORM'] = df_valido['JORNADA'].str.upper().str.strip()
                    else:
                        df_valido['JORNADA_NORM'] = 'NO ESPECIFICADO'

                    for anio in sorted(df_valido['ANIO_INGRESO'].unique()):
                        doc.add_heading(f'Cohorte: {int(anio)}', level=3)
                        df_anio = df_valido[df_valido['ANIO_INGRESO'] == anio]
                        
                        for tipo in sorted(df_anio['TIPO_RETIRO'].unique()):
                            doc.add_heading(f'Tipo: {tipo}', level=4)
                            df_tipo = df_anio[df_anio['TIPO_RETIRO'] == tipo]
                            
                            # ---------------------------------------------------------
                            # TABLA 1 NUEVA: MOTIVOS POR JORNADA (CRUZADA)
                            # ---------------------------------------------------------
                            doc.add_heading('Distribución de Motivos por Jornada', level=5)
                            
                            df_cross_jor = pd.crosstab([df_tipo['NOMBRE_CARRERA'], df_tipo['MOTIVO_LIMPIO']], df_tipo['JORNADA_NORM']).reset_index()
                            # Garantizar la existencia de ambas columnas requeridas
                            for col in ['DIURNA', 'VESPERTINA']:
                                if col not in df_cross_jor.columns:
                                    df_cross_jor[col] = 0
                            
                            # Calcular total y ordenar por carrera y luego por total (mayor a menor)
                            df_cross_jor['TOTAL'] = df_cross_jor['DIURNA'] + df_cross_jor['VESPERTINA']
                            df_cross_jor = df_cross_jor[['NOMBRE_CARRERA', 'MOTIVO_LIMPIO', 'DIURNA', 'VESPERTINA', 'TOTAL']].sort_values(by=['NOMBRE_CARRERA', 'TOTAL'], ascending=[True, False])
                            
                            table_jor = doc.add_table(rows=1, cols=5)
                            table_jor.style = 'Table Grid'
                            hdr_jor = table_jor.rows[0].cells
                            hdr_jor[0].text, hdr_jor[1].text, hdr_jor[2].text, hdr_jor[3].text, hdr_jor[4].text = 'Carrera', 'Motivo declarado', 'Jornada Diurna', 'Jornada Vespertina', 'Total'
                            
                            for cell in hdr_jor:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs: run.bold = True
                                        
                            tot_diurna, tot_vesp, tot_global_j = 0, 0, 0
                            for _, row in df_cross_jor.iterrows():
                                cells = table_jor.add_row().cells
                                cells[0].text = str(row['NOMBRE_CARRERA'])
                                cells[1].text = str(row['MOTIVO_LIMPIO'])
                                cells[2].text = str(int(row['DIURNA']))
                                cells[3].text = str(int(row['VESPERTINA']))
                                cells[4].text = str(int(row['TOTAL']))
                                tot_diurna += row['DIURNA']
                                tot_vesp += row['VESPERTINA']
                                tot_global_j += row['TOTAL']
                                
                            row_tot_jor = table_jor.add_row().cells
                            row_tot_jor[0].text = "TOTAL"
                            row_tot_jor[1].text = ""
                            row_tot_jor[2].text = str(int(tot_diurna))
                            row_tot_jor[3].text = str(int(tot_vesp))
                            row_tot_jor[4].text = str(int(tot_global_j))
                            for cell in row_tot_jor:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True
                            doc.add_paragraph()
                            
                            # ---------------------------------------------------------
                            # TABLA 2 NUEVA: MOTIVOS POR SEMESTRE (CRUZADA)
                            # ---------------------------------------------------------
                            doc.add_heading('Distribución de Motivos por Semestre', level=5)
                            
                            df_cross_sem = pd.crosstab([df_tipo['NOMBRE_CARRERA'], df_tipo['MOTIVO_LIMPIO']], df_tipo['SEMESTRE_EXTRACT']).reset_index()
                            # Garantizar la existencia de los dos semestres requeridos
                            for col in ['1er Semestre', '2do Semestre']:
                                if col not in df_cross_sem.columns:
                                    df_cross_sem[col] = 0
                                    
                            # Calcular total y ordenar por carrera y luego por total (mayor a menor)
                            df_cross_sem['TOTAL'] = df_cross_sem['1er Semestre'] + df_cross_sem['2do Semestre']
                            df_cross_sem = df_cross_sem[['NOMBRE_CARRERA', 'MOTIVO_LIMPIO', '1er Semestre', '2do Semestre', 'TOTAL']].sort_values(by=['NOMBRE_CARRERA', 'TOTAL'], ascending=[True, False])
                            
                            table_sem = doc.add_table(rows=1, cols=5)
                            table_sem.style = 'Table Grid'
                            hdr_sem = table_sem.rows[0].cells
                            hdr_sem[0].text, hdr_sem[1].text, hdr_sem[2].text, hdr_sem[3].text, hdr_sem[4].text = 'Carrera', 'Motivo declarado', '1er Semestre', '2do Semestre', 'Total'
                            
                            for cell in hdr_sem:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs: run.bold = True
                                        
                            tot_sem1, tot_sem2, tot_global_s = 0, 0, 0
                            for _, row in df_cross_sem.iterrows():
                                cells = table_sem.add_row().cells
                                cells[0].text = str(row['NOMBRE_CARRERA'])
                                cells[1].text = str(row['MOTIVO_LIMPIO'])
                                cells[2].text = str(int(row['1er Semestre']))
                                cells[3].text = str(int(row['2do Semestre']))
                                cells[4].text = str(int(row['TOTAL']))
                                tot_sem1 += row['1er Semestre']
                                tot_sem2 += row['2do Semestre']
                                tot_global_s += row['TOTAL']
                                
                            row_tot_sem = table_sem.add_row().cells
                            row_tot_sem[0].text = "TOTAL"
                            row_tot_sem[1].text = ""
                            row_tot_sem[2].text = str(int(tot_sem1))
                            row_tot_sem[3].text = str(int(tot_sem2))
                            row_tot_sem[4].text = str(int(tot_global_s))
                            for cell in row_tot_sem:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True
                            doc.add_paragraph()

            # Llamadas dinámicas cruzadas para los 3 momentos del flujo institucional
            generar_anexo_motivos(doc, df_fugas_1er_export, '5.1 Motivos de Deserción al 1er Año', max_anio_registrado - 1)
            generar_anexo_motivos(doc, df_fugas_2do_export, '5.2 Motivos de Deserción al 2do Año', max_anio_registrado - 1)
            generar_anexo_motivos(doc, df_fugas_3er_export, '5.3 Motivos de Deserción al 3er Año', max_anio_registrado - 2)

            # =========================================================================
            # 6. GESTIÓN DE CONVALIDACIONES Y TUTORÍAS (SIN MODIFICACIONES)
            # =========================================================================
            if not df_dist_acort.empty or total_tutorias_cont > 0:
                doc.add_heading('6. Gestión de Convalidaciones y Tutorías', level=1)
                doc.add_paragraph(f"Total de Intervenciones de Tutoría registradas: {total_tutorias_cont}")

                if not df_dist_acort.empty:
                    doc.add_heading('Estudiantes Únicos por Beneficio de Acortamiento (Acumulado Global)', level=2)
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr = table.rows[0].cells
                    hdr[0].text, hdr[1].text = 'Tipo de Beneficio', 'N° Estudiantes (Únicos)'

                    for _, row in df_dist_acort.iterrows():
                        cells = table.add_row().cells
                        cells[0].text = str(row['CAT_AGRUPADA'])
                        cells[1].text = str(row['Estudiantes'])
                    doc.add_paragraph()

                    if not df_hist_acort.empty:
                        doc.add_heading('Estudiantes Únicos Desglosado por Año de Beneficio', level=2)
                        for anio in sorted(df_hist_acort['ANIO_BENEFICIO'].unique()):
                            doc.add_heading(f'Año: {int(float(anio))}', level=3)
                            df_anio_acort = df_hist_acort[df_hist_acort['ANIO_BENEFICIO'] == anio]

                            table_anio = doc.add_table(rows=1, cols=2)
                            table_anio.style = 'Table Grid'
                            hdr_anio = table_anio.rows[0].cells
                            hdr_anio[0].text, hdr_anio[1].text = 'Tipo de Beneficio', 'N° Estudiantes (Únicos)'

                            tot_est = 0
                            for _, row in df_anio_acort.iterrows():
                                cells_anio = table_anio.add_row().cells
                                cells_anio[0].text = str(row['CAT_AGRUPADA'])
                                cells_anio[1].text = str(row['Estudiantes'])
                                tot_est += int(row['Estudiantes'])

                            row_tot = table_anio.add_row().cells
                            row_tot[0].text = "TOTAL"
                            row_tot[1].text = str(tot_est)
                            doc.add_paragraph()

            # =========================================================================
            # 7. TASA DE PROGRESIÓN (2DO A 3ER AÑO)
            # =========================================================================
            if not df_det_prog.empty:
                df_rep_prog = df_det_prog[df_det_prog['Cohorte'] <= (max_anio_registrado - 2)]
                if not df_rep_prog.empty:
                    doc.add_heading('7. Tasa de Progresión (2do a 3er Año)', level=1)
                    for anio in sorted(df_rep_prog['Cohorte'].unique()):
                        doc.add_heading(f'Cohorte: {int(anio)}', level=2)
                        df_anio = df_rep_prog[df_rep_prog['Cohorte'] == anio]

                        # --- NUEVO: Resumen Consolidado por Carrera ---
                        doc.add_heading('Resumen Consolidado por Carrera', level=3)
                        df_carrera = df_anio.groupby('Carrera')[['Retenidos (2do Año)', 'Progresión (3er Año)']].sum().reset_index()
                        
                        table_car = doc.add_table(rows=1, cols=4)
                        table_car.style = 'Table Grid'
                        hdr_car = table_car.rows[0].cells
                        hdr_car[0].text, hdr_car[1].text, hdr_car[2].text, hdr_car[3].text = 'Carrera', 'Retenidos 2do Año', 'Progresión 3er Año', 'Tasa (%)'
                        
                        for cell in hdr_car:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs: run.bold = True
                                    
                        tot_ret_c, tot_prog_c = 0, 0
                        for _, row in df_carrera.iterrows():
                            cells = table_car.add_row().cells
                            cells[0].text = str(row['Carrera'])
                            cells[1].text = str(int(row['Retenidos (2do Año)']))
                            cells[2].text = str(int(row['Progresión (3er Año)']))
                            tasa_c = round((row['Progresión (3er Año)'] / row['Retenidos (2do Año)'] * 100), 1) if row['Retenidos (2do Año)'] > 0 else 0
                            cells[3].text = f"{tasa_c}%"
                            tot_ret_c += row['Retenidos (2do Año)']
                            tot_prog_c += row['Progresión (3er Año)']
                            
                        row_tot_c = table_car.add_row().cells
                        row_tot_c[0].text = "TOTAL COHORTE"
                        row_tot_c[1].text = str(int(tot_ret_c))
                        row_tot_c[2].text = str(int(tot_prog_c))
                        row_tot_c[3].text = f"{round((tot_prog_c/tot_ret_c*100) if tot_ret_c > 0 else 0, 1)}%"
                        for cell in row_tot_c:
                            for paragraph in cell.paragraphs:
                                if paragraph.runs: paragraph.runs[0].bold = True
                        doc.add_paragraph()

                        # --- ORIGINAL: Detalle por Jornada ---
                        doc.add_heading('Detalle por Jornada', level=3)
                        for jor in sorted(df_anio['Jornada'].unique()):
                            doc.add_heading(f'Jornada: {jor}', level=4)
                            df_jor_data = df_anio[df_anio['Jornada'] == jor].sort_values(by='Carrera')

                            table = doc.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            hdr = table.rows[0].cells
                            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Carrera', 'Retenidos 2do Año', 'Progresión 3er Año', 'Tasa (%)'

                            for cell in hdr:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs: run.bold = True

                            for _, row in df_jor_data.iterrows():
                                cells = table.add_row().cells
                                cells[0].text = str(row['Carrera'])
                                cells[1].text = str(int(row['Retenidos (2do Año)']))
                                cells[2].text = str(int(row['Progresión (3er Año)']))
                                cells[3].text = f"{row['Tasa Progresión (%)']}%"

                            row_tot = table.add_row().cells
                            row_tot[0].text = "SUBTOTAL JORNADA"
                            row_tot[1].text = str(int(df_jor_data['Retenidos (2do Año)'].sum()))
                            row_tot[2].text = str(int(df_jor_data['Progresión (3er Año)'].sum()))
                            tot_ret_j = df_jor_data['Retenidos (2do Año)'].sum()
                            tot_prog_j = df_jor_data['Progresión (3er Año)'].sum()
                            row_tot[3].text = f"{round((tot_prog_j/tot_ret_j*100) if tot_ret_j > 0 else 0, 1)}%"
                            for cell in row_tot:
                                for paragraph in cell.paragraphs:
                                    if paragraph.runs: paragraph.runs[0].bold = True
                            doc.add_paragraph()

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    st.write("Descarga el reporte consolidado basado en los filtros actuales.")
    docx_file = crear_documento_word()
    st.download_button(
        label="📥 Descargar Informe en Word",
        data=docx_file,
        file_name=f"Reporte_AI_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.error("⚠️ Para descargar el informe necesitas la librería 'python-docx'. Instálala en tu terminal ejecutando: pip install python-docx")



