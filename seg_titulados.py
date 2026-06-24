import streamlit as st
import pandas as pd
import os
import plotly.express as px
from docx import Document        
from docx.shared import Inches   
import io                        
import time

# --- 1. CONFIGURACIÓN E IDENTIDAD INSTITUCIONAL ---
st.set_page_config(page_title="Análisis Institucional CFTe-AP", layout="wide")
COLORES = ['#4A90E2', '#50E3C2', '#F5A623', '#F8E71C', '#9B9B9B', '#D3D3D3']

st.title("🏛️ Dimensión: Vinculación con el Medio y Resultados")
st.write("Seguimiento de Titulados para Aseguramiento de la Calidad (CNA / PDI)")
st.markdown("---")

# --- 2. RUTAS ABSOLUTAS ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CARPETA_VIN = os.path.join(BASE_DIR, "2_Datos_Procesados", "2_Datos_Procesados_Vinculacion")

# --- 3. LECTOR INTELIGENTE Y TODOTERRENO ---
def cargar_archivo(ruta):
    try:
        if ruta.endswith(".xlsx"):
            return pd.read_excel(ruta)
        for s in [';', ',']:
            for enc in ['utf-8-sig', 'latin-1']:
                try:
                    df = pd.read_csv(ruta, sep=s, encoding=enc, on_bad_lines='skip', engine='python')
                    if 'Carrera_dsc' in df.columns or 'Rut' in df.columns: return df
                except: continue
    except: return None
    return None

# --- 4. CARGA, CONSOLIDACIÓN Y CRUCE DE DATOS ---
if not os.path.exists(CARPETA_VIN):
    st.error(f"❌ No se encontró la carpeta: {CARPETA_VIN}")
    st.stop()

archivos_todos = os.listdir(CARPETA_VIN)
archivos_encuesta = [f for f in archivos_todos if f.startswith("LIMPIO_") and "Nomina" not in f]
archivos_nomina = [f for f in archivos_todos if "Nomina" in f]

if not archivos_encuesta:
    st.warning("⚠️ No se han encontrado archivos de Encuestas procesados en la carpeta.")
    st.stop()

# Diccionario de homologación para unificar nombres de columnas
DICCIONARIO_HOMOLOGACION = {
    'Datos de Empleabilidad ¿Cuanto tiempo demoro en encontrar empleo desde su titulación o inicio de actividades?': 'Datos de Empleabilidad',
    'Datos de Empleabilidad ¿Cuánto tiempo demoró en encontrar su primer empleo desde su titulación?': 'Datos de Empleabilidad',
    'Continuidad de Estudios ¿En qué institución estudia?': 'Institución de Continuidad',
    'Si estas continuando estudios ¿En qué institución realiza su continuidad de estudios?': 'Institución de Continuidad',
    'Continuidad de Estudios Tipo de Estudio': 'Tipo de Continuidad',
    'Si estas continuando estudios Tipo de estudio': 'Tipo de Continuidad',
    'Indique 3 áreas de interés en la cual le gustaría desarrollar cursos de especialización o diplomados dictados por el CFT Estatal de Arica y Parinacota Opción 1': 'Áreas de Interés Formación Continua',
    'Áreas de Interés Indique 1 área de interés en la cual le gustaría desarrollar Cursos certificados, taller, capacitaciones.': 'Áreas de Interés Formación Continua',
    'Datos de Empleabilidad Nombre de la empresa dónde trabaja': 'Datos de Empleabilidad Nombre de la empresa en la que trabaja',
    'Contacto Fono/Celular': 'Contacto Celular'
}

lista_encuestas = []
for arc in archivos_encuesta:
    temp_df = cargar_archivo(os.path.join(CARPETA_VIN, arc))
    if temp_df is not None:
        temp_df['Origen_Encuesta'] = 'Actual' if '2023' not in arc else 'Antigua'
        temp_df = temp_df.rename(columns=DICCIONARIO_HOMOLOGACION)
        lista_encuestas.append(temp_df)

df_full_encuesta = pd.concat(lista_encuestas, ignore_index=True)

# ==============================================================
# NUEVO FILTRO: EXCLUSIÓN DE ESTUDIANTES PREVIOS A 2021 (ENCUESTA)
# ==============================================================
col_ingreso = [c for c in df_full_encuesta.columns if 'AÑO INGRESO' in str(c).upper()]
if col_ingreso:
    # Convertimos a número, los que den error serán 0. Filtramos solo los >= 2021
    df_full_encuesta['Año_Ingreso_Num'] = pd.to_numeric(df_full_encuesta[col_ingreso[0]], errors='coerce').fillna(0).astype(int)
    df_full_encuesta = df_full_encuesta[df_full_encuesta['Año_Ingreso_Num'] >= 2021]


# --- REGLAS DE NEGOCIO Y LIMPIEZA DE RESPUESTAS ---

# A. Limpieza de Renta
mapa_renta = {
    "MENOS DEL SUELDO MÍNIMO": "MENOS DEL SUELDO MÍNIMO",
    "400.000 - 500.000": "SUELDO MÍNIMO",
    "SUELDO MÍNIMO ($510.636)": "SUELDO MÍNIMO",
    "500.001 - 700.000": "ENTRE SUELDO MÍNIMO Y 700.000",
    "511.000 A 700.000": "ENTRE SUELDO MÍNIMO Y 700.000",
    "700.001 - 900.000": "700.001 A 900.000",
    "701.000 A 900.000": "700.001 A 900.000",
    "900.001 - 1.100.000": "900.001 A 1.100.000",
    "900.001 A 1.100.000": "900.001 A 1.100.000",
    "MÁS DE 1.100.000": "MÁS DE 1.100.000",
    "MAS DE 1.100.000": "MÁS DE 1.100.000",
    "6": None  
}
if 'Datos de Empleabilidad Renta' in df_full_encuesta.columns:
    df_full_encuesta['Datos de Empleabilidad Renta'] = df_full_encuesta['Datos de Empleabilidad Renta'].astype(str).str.upper().str.strip()
    df_full_encuesta['Datos de Empleabilidad Renta'] = df_full_encuesta['Datos de Empleabilidad Renta'].replace(mapa_renta)

# B. Limpieza de Tiempo de Inserción
mapa_tiempo = {
    "AUN NO ENCUENTRA TRABAJO": "AÚN NO ENCUENTRA TRABAJO",
    "MAS DE 12 MESES": "MÁS DE 12 MESES"
}
if 'Datos de Empleabilidad' in df_full_encuesta.columns:
    df_full_encuesta['Datos de Empleabilidad'] = df_full_encuesta['Datos de Empleabilidad'].astype(str).str.upper().str.strip()
    df_full_encuesta['Datos de Empleabilidad'] = df_full_encuesta['Datos de Empleabilidad'].replace(mapa_tiempo)

# D.1 Limpieza Institución Continuidad
mapa_inst = {
    "SANTO TOMAS": "SANTO TOMÁS",
    "UTA": "UNIVERSIDAD DE TARAPACÁ",
    "UNAP": "UNIVERSIDAD ARTURO PRAT",
    "UNAP ARTURO PRAT": "UNIVERSIDAD ARTURO PRAT"
}
if 'Institución de Continuidad' in df_full_encuesta.columns:
    df_full_encuesta['Institución de Continuidad'] = df_full_encuesta['Institución de Continuidad'].astype(str).str.upper().str.strip()
    df_full_encuesta['Institución de Continuidad'] = df_full_encuesta['Institución de Continuidad'].replace(mapa_inst)

# D.2 Limpieza Tipo de Estudio
mapa_tipo = {
    "IP": "INSTITUTO PROFESIONAL",
    "DIPLOMADO/POSTÍTULO": "DIPLOMADO",
    "MAGISTER": "MAGISTER/MASTER"
}
if 'Tipo de Continuidad' in df_full_encuesta.columns:
    df_full_encuesta['Tipo de Continuidad'] = df_full_encuesta['Tipo de Continuidad'].astype(str).str.upper().str.strip()
    df_full_encuesta['Tipo de Continuidad'] = df_full_encuesta['Tipo de Continuidad'].replace(mapa_tipo)
    
    mask_split = df_full_encuesta['Tipo de Continuidad'] == "CARRERA UNIVERSITARIA O IP (INSTITUTO PROFESIONAL)"
    idx_split = df_full_encuesta[mask_split].index
    if len(idx_split) > 0:
        mitad = len(idx_split) // 2
        df_full_encuesta.loc[idx_split[:mitad], 'Tipo de Continuidad'] = "CARRERA UNIVERSITARIA"
        df_full_encuesta.loc[idx_split[mitad:], 'Tipo de Continuidad'] = "INSTITUTO PROFESIONAL"

# NUEVO SISTEMA DE DESEMPATE POR RUT 
df_full_encuesta['Prioridad_Origen'] = df_full_encuesta['Origen_Encuesta'].apply(lambda x: 2 if x == 'Actual' else 1)
df_full_encuesta = df_full_encuesta.sort_values(by='Prioridad_Origen', ascending=True)

col_rut = [c for c in df_full_encuesta.columns if 'RUT' in c.upper()]
if col_rut:
    df_full_encuesta = df_full_encuesta.drop_duplicates(subset=[col_rut[0]], keep='last')

df_full_encuesta = df_full_encuesta.drop(columns=['Prioridad_Origen'])

# Normalizar años de Titulación
if 'Titulación Fecha' in df_full_encuesta.columns:
    df_full_encuesta['Año_T'] = pd.to_datetime(df_full_encuesta['Titulación Fecha'], errors='coerce').dt.year.fillna(0).astype(int)
else:
    df_full_encuesta['Año_T'] = 2024

# --- 5. FILTROS LATERALES ---
st.sidebar.header("🔍 Filtros de Gestión")

df_nom_maestro = None
if archivos_nomina:
    df_nom_maestro = cargar_archivo(os.path.join(CARPETA_VIN, archivos_nomina[0]))
    
    # ==============================================================
    # NUEVO FILTRO: EXCLUSIÓN PRE-2021 EN MAESTRO DE FILTROS LATERALES
    # ==============================================================
    if df_nom_maestro is not None:
        col_ing_maestro = [c for c in df_nom_maestro.columns if 'AÑO INGRESO' in str(c).upper()]
        if col_ing_maestro:
            df_nom_maestro['Año_Ingreso_Num'] = pd.to_numeric(df_nom_maestro[col_ing_maestro[0]], errors='coerce').fillna(0).astype(int)
            df_nom_maestro = df_nom_maestro[df_nom_maestro['Año_Ingreso_Num'] >= 2021]

# --- A. OBTENER AÑOS ---
años_lista = set()
if df_nom_maestro is not None:
    col_f_n = [c for c in df_nom_maestro.columns if 'TITULACIÓN FECHA' in str(c).upper()]
    if col_f_n:
        años_nom = pd.to_datetime(df_nom_maestro[col_f_n[0]], errors='coerce').dt.year.dropna().unique()
        años_lista.update(años_nom.astype(int))

años_lista.update(df_full_encuesta['Año_T'].unique())
años_disp = sorted([a for a in años_lista if a > 0])
sel_años = st.sidebar.multiselect("Año de Titulación", años_disp, default=años_disp)

# --- B. OBTENER CARRERAS ---
carreras_lista = set()

if df_nom_maestro is not None:
    col_c_n = [c for c in df_nom_maestro.columns if 'CARRERA' in str(c).upper() and 'COD' not in str(c).upper()]
    if col_c_n:
        c_nom = df_nom_maestro[col_c_n[0]].astype(str).str.upper().str.strip().str.replace(r'^\d+\s+', '', regex=True).unique()
        carreras_lista.update(c_nom)

if 'Carrera_dsc' in df_full_encuesta.columns:
    c_enc = df_full_encuesta['Carrera_dsc'].astype(str).str.upper().str.strip().unique()
    carreras_lista.update(c_enc)

carreras_disp = sorted([c for c in carreras_lista if c not in ['NAN', 'NONE', '']])
sel_carreras = st.sidebar.multiselect("Carrera", carreras_disp, default=carreras_disp)

# --- C. APLICAR FILTROS A LA ENCUESTA ---
df_full_encuesta['Carrera_dsc_LIMPIA'] = df_full_encuesta['Carrera_dsc'].astype(str).str.upper().str.strip()
df = df_full_encuesta[
    (df_full_encuesta['Año_T'].isin(sel_años)) & 
    (df_full_encuesta['Carrera_dsc_LIMPIA'].isin(sel_carreras))
].copy()
total_respuestas = len(df)

# --- 6. CÁLCULO DE UNIVERSO (NÓMINA - CONTEO DIRECTO) ---
st.subheader("📈 Representatividad de la Muestra")

if archivos_nomina:
    df_nom_raw = cargar_archivo(os.path.join(CARPETA_VIN, archivos_nomina[0]))
    
    if df_nom_raw is not None:
        
        # ==============================================================
        # NUEVO FILTRO: EXCLUSIÓN PRE-2021 EN EL CÁLCULO DEL UNIVERSO
        # ==============================================================
        col_ingreso_raw = [c for c in df_nom_raw.columns if 'AÑO INGRESO' in str(c).upper()]
        if col_ingreso_raw:
            df_nom_raw['Año_Ingreso_Num'] = pd.to_numeric(df_nom_raw[col_ingreso_raw[0]], errors='coerce').fillna(0).astype(int)
            df_nom_raw = df_nom_raw[df_nom_raw['Año_Ingreso_Num'] >= 2021]

        col_rut_nom = [c for c in df_nom_raw.columns if 'RUT' in str(c).upper() or 'RUN' in str(c).upper()]
        col_rut_enc = [c for c in df.columns if 'RUT' in str(c).upper() or 'RUN' in str(c).upper()]
        
        if col_rut_nom and col_rut_enc:
            # UNIVERSO ABSOLUTO (Ahora sólo considera desde 2021)
            universo_absoluto = df_nom_raw[col_rut_nom[0]].nunique()
            total_universo = universo_absoluto 
          
            col_fecha_nom = [c for c in df_nom_raw.columns if 'TITULACIÓN FECHA' in str(c).upper()]
            if col_fecha_nom:
              df_nom_raw['Año_T_Nom'] = pd.to_datetime(df_nom_raw[col_fecha_nom[0]], errors='coerce').dt.year.fillna(0).astype(int)
            else:
              df_nom_raw['Año_T_Nom'] = 2024 

            col_carrera_nom = [c for c in df_nom_raw.columns if 'CARRERA' in str(c).upper() and 'COD' not in str(c).upper()]

            if col_carrera_nom:
              df_nom_raw['TEMP_CARRERA'] = df_nom_raw[col_carrera_nom[0]].astype(str).str.upper().str.strip()
              df_nom_raw['TEMP_CARRERA'] = df_nom_raw['TEMP_CARRERA'].str.replace(r'^\d+\s+', '', regex=True)

              if sel_carreras:
                  sel_carreras_upper = [str(c).upper().strip() for c in sel_carreras]
                  filtro_carrera = df_nom_raw['TEMP_CARRERA'].isin(sel_carreras_upper)
                  filtro_ano = df_nom_raw['Año_T_Nom'].isin(sel_años)
                  df_nom_filt = df_nom_raw[filtro_carrera & filtro_ano]
                  total_universo = df_nom_filt[col_rut_nom[0]].nunique()

            # Respuestas únicas 
            total_respuestas_unicas = df[col_rut_enc[0]].nunique()
            tasa = (total_respuestas_unicas / total_universo * 100) if total_universo > 0 else 0
            
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Universo Histórico (Desde 2021)", f"{universo_absoluto} Pers.")
            c2.metric("Universo Filtrado", f"{total_universo} Pers.")
            c3.metric("Respuestas Únicas", f"{total_respuestas_unicas} Pers.")
            c4.metric("Tasa de Respuesta", f"{tasa:.1f}%")
            
            st.progress(tasa / 100 if tasa <= 100 else 1.0)
            
            if total_universo == 0 and universo_absoluto > 0:
                st.warning("⚠️ El Universo Filtrado es cero. Las carreras de la Nómina no se llaman igual que las de la Encuesta.")
                st.info(f"👉 Nombres en la Nómina: {list(df_nom_raw['TEMP_CARRERA'].dropna().unique()[:5])}...")
                st.info(f"👉 Nombres que estás filtrando: {sel_carreras[:5]}")
                
        else:
            st.error("⚠️ No se encontró la columna de identificación (RUT o RUN) en alguno de los archivos.")
    else:
        st.error("No se pudo leer el archivo de Nómina.")
else:
    st.info("ℹ️ Cargue el archivo de Nómina ('LIMPIO_Nomina_Titulados') para calcular la Tasa de Respuesta.")

st.markdown("---")

# --- 7. PESTAÑAS DE ANÁLISIS DE RESULTADOS ---
if total_respuestas == 0:
    st.warning("⚠️ No hay encuestas para mostrar con los filtros seleccionados.")
    st.stop()

t1, t2, t3 = st.tabs(["💼 Empleabilidad y Renta", "🌟 Satisfacción y Calidad", "🎓 Progresión y Formación"])

def bloque_indicador(titulo, columna, orden_estricto, total_especifico=None):
    if columna not in df.columns:
        return
    st.subheader(titulo)
    datos = df[columna].dropna().astype(str).str.upper().str.strip()
    datos = datos[~datos.isin(['NAN', 'NONE', '', 'N/A'])]
    
    total_validas = len(datos)
    if total_validas == 0:
        st.info("No hay respuestas registradas para este indicador.")
        return

    denominador = total_especifico if total_especifico is not None else total_respuestas
    conteo = datos.value_counts().reindex(orden_estricto, fill_value=0)
    df_res = pd.DataFrame({'Cantidad': conteo.values, 'Porcentaje': [(v/total_validas*100) for v in conteo.values]}, index=conteo.index)
    df_res['Porcentaje'] = df_res['Porcentaje'].round(1).astype(str) + '%'

    c_t, c_g = st.columns([1, 2])
    with c_t:
        st.metric("Respuestas Válidas", f"{total_validas} / {denominador}")
        st.dataframe(df_res, use_container_width=True)
    with c_g:
        fig = px.bar(x=conteo.index, y=conteo.values, color=conteo.index, color_discrete_sequence=COLORES, text_auto=True)
        fig.update_layout(showlegend=False, height=300, xaxis_title=None, xaxis={'categoryorder':'array', 'categoryarray': orden_estricto})
        st.plotly_chart(fig, use_container_width=True)

with t1:
    ord_r = ["MENOS DEL SUELDO MÍNIMO", "SUELDO MÍNIMO", "ENTRE SUELDO MÍNIMO Y 700.000", "700.001 A 900.000", "900.001 A 1.100.000", "MÁS DE 1.100.000"]
    bloque_indicador("Análisis de Rentas Promedio", 'Datos de Empleabilidad Renta', ord_r)
    st.markdown("---")
    
    ord_t = ["AÚN NO ENCUENTRA TRABAJO", "MENOS DE 6 MESES", "DE 6 A 12 MESES", "MÁS DE 12 MESES", "YA SE ENCONTRABA TRABAJANDO"]
    bloque_indicador("Datos de Empleabilidad", 'Datos de Empleabilidad', ord_t)

with t2:
    st.info("ℹ️ **Nota Metodológica:** Estos indicadores de satisfacción evalúan el impacto de la formación y no se encontraban presentes en el instrumento aplicado durante 2021-2023.")
    
    total_solo_titulacion = len(df[df['Origen_Encuesta'] == 'Actual'])
    
    ord_s = ["MUY BUENA", "BUENA", "SUFICIENTE", "DEFICIENTE"]
    bloque_indicador("Satisfacción con la Formación CFTe-AP", 
                     'Formación y expectativas respecto a su desempeño laboral Como es su formación respecto a su desempeño en su trabajo', 
                     ord_s, 
                     total_especifico=total_solo_titulacion)
    
    st.markdown("---")
    
    ord_w = ["MUY DE ACUERDO", "DE ACUERDO", "EN DESACUERDO", "MUY EN DESACUERDO"]
    bloque_indicador("Nivel de Satisfacción en el Trabajo Actual", 
                     'Formación y expectativas respecto a su desempeño laboral Logra cumplir sus expectativas laborales', 
                     ord_w, 
                     total_especifico=total_solo_titulacion)
    
with t3:
    st.subheader("📚 Continuidad de Estudios Superiores")
    col_inst = 'Institución de Continuidad'
    col_tipo = 'Tipo de Continuidad'
    
    if col_inst in df.columns:
        ord_inst = ["INACAP", "SANTO TOMÁS", "UNIVERSIDAD DE TARAPACÁ", "UNIVERSIDAD ARTURO PRAT", "OTRO"]
        bloque_indicador("Institución de Continuidad", col_inst, ord_inst)
        
    if col_tipo in df.columns:
        ord_tipo = ["INSTITUTO PROFESIONAL", "CARRERA UNIVERSITARIA", "DIPLOMADO", "MAGISTER/MASTER", "DOCTORADO/POSTDOCTORADO"]
        bloque_indicador("Tipo de Estudio", col_tipo, ord_tipo)

    st.markdown("---")
    st.subheader("🎯 Demandas de Formación Continua (Top 5)")
    col_i = 'Áreas de Interés Formación Continua'
    if col_i in df.columns:
        def agrupador_final(t):
            t = str(t).upper().strip()
            if t in ['NAN', '', 'NONE', 'N/A', 'NINGUNO', 'NO', 'NADA']: return None
            if 'NO TIENE INTERÉS' in t or 'NO ME INTERESA' in t: return None
            
            if 'EXCEL' in t: return 'EXCEL (TODOS LOS NIVELES)'
            if 'SEÑA' in t or 'INCLUSI' in t or 'LEGUAJE' in t: return 'LENGUA DE SEÑAS E INCLUSIÓN'
            if 'INGLE' in t or 'INGLÉ' in t: return 'INGLÉS TÉCNICO'
            if 'CONTAB' in t or 'FINAN' in t: return 'CONTABILIDAD Y TRIBUTARIA'
            if 'SALUD' in t or 'CURA' in t or 'PRIMEROS' in t: return 'SALUD Y PRIMEROS AUXILIOS'
            if 'RRHH' in t or 'HUMANO' in t: return 'RECURSOS HUMANOS'
            if 'EDUCAC' in t or 'PEDAGOG' in t: return 'EDUCACIÓN / PEDAGOGÍA'
            
            return t
        
        df['INT_FINAL'] = df[col_i].apply(agrupador_final)
        res_int = df['INT_FINAL'].dropna().value_counts()
        
        if not res_int.empty:
            ci1, ci2 = st.columns([1, 2])
            with ci1:
                st.dataframe(pd.DataFrame({'Cant': res_int.values, '%': (res_int.values/res_int.sum()*100).round(1).astype(str)+'%'}, index=res_int.index), use_container_width=True)
            with ci2:
                fig_int = px.bar(x=res_int.head(5).index, y=res_int.head(5).values, color=res_int.head(5).index, text_auto=True, color_discrete_sequence=COLORES)
                fig_int.update_layout(showlegend=False, height=350, margin=dict(t=10, b=10), xaxis_title=None, yaxis_title="Cantidad")
                st.plotly_chart(fig_int, use_container_width=True)

# --- 7.5. GENERADOR DE REPORTES WORD ---
def generar_word_institucional(df_filtrado, stats):
    doc = Document()
    
    doc.add_heading('Informe de Análisis Institucional: Vinculación con el Medio', 0)
    doc.add_paragraph(f"Fecha de generación: {time.strftime('%d/%m/%Y')}")
    doc.add_paragraph("---")

    doc.add_heading('1. Resumen de Representatividad', level=1)
    doc.add_paragraph(f"• Total Histórico (Desde 2021): {stats.get('historico', 0)} personas", style='List Bullet')
    doc.add_paragraph(f"• Universo Filtrado: {stats.get('universo', 0)} personas", style='List Bullet')
    doc.add_paragraph(f"• Respuestas Únicas: {stats.get('respuestas', 0)} personas", style='List Bullet')
    doc.add_paragraph(f"• Tasa de Respuesta Global: {stats.get('tasa', 0):.1f}%", style='List Bullet')
    doc.add_paragraph()

    def escribir_bloque_puntos(doc_obj, titulo, columna, dataframe, total_referencia, top=None):
        if columna not in dataframe.columns:
            return
        
        datos = dataframe[columna].dropna().astype(str).str.upper().str.strip()
        datos = datos[~datos.isin(['NAN', 'NONE', '', 'N/A'])]
        
        validas = len(datos)
        if validas == 0:
            return
            
        conteo = datos.value_counts()
        if top:
            conteo = conteo.head(top)
        
        tasa_indicador = (validas / total_referencia * 100) if total_referencia > 0 else 0
        
        p = doc_obj.add_paragraph()
        p.add_run(f"{titulo} ({validas}/{total_referencia} - {tasa_indicador:.1f}%)").bold = True
        
        table = doc_obj.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Categoría'
        hdr_cells[1].text = 'Cantidad'
        hdr_cells[2].text = '%'
        
        for cat, val in conteo.items():
            porcentaje_cat = (val / validas * 100) if validas > 0 else 0
            row_cells = table.add_row().cells
            row_cells[0].text = str(cat)
            row_cells[1].text = str(val)
            row_cells[2].text = f"{porcentaje_cat:.1f}%"
            
        doc_obj.add_paragraph() 

    doc.add_heading('2. Empleabilidad y Renta', level=1)
    escribir_bloque_puntos(doc, 'a) Tasa de Empleabilidad', 'Datos de Empleabilidad', df_filtrado, stats['total_global'])
    escribir_bloque_puntos(doc, 'b) Rentas promedio', 'Datos de Empleabilidad Renta', df_filtrado, stats['total_global'])

    doc.add_heading('3. Continuidad de Estudios Superiores', level=1)
    escribir_bloque_puntos(doc, 'a) Institución de Continuidad', 'Institución de Continuidad', df_filtrado, stats['total_global'])
    escribir_bloque_puntos(doc, 'b) Tipo de estudio', 'Tipo de Continuidad', df_filtrado, stats['total_global'])
    
    if 'INT_FINAL' in df_filtrado.columns:
        escribir_bloque_puntos(doc, 'c) Demandas de Formación Continua (top 5)', 'INT_FINAL', df_filtrado, stats['total_global'], top=5)

    doc.add_heading('4. Satisfacción y Calidad', level=1)
    df_actual = df_filtrado[df_filtrado['Origen_Encuesta'] == 'Actual']
    
    col_sat1 = 'Formación y expectativas respecto a su desempeño laboral Como es su formación respecto a su desempeño en su trabajo'
    col_sat2 = 'Formación y expectativas respecto a su desempeño laboral Logra cumplir sus expectativas laborales'
    
    escribir_bloque_puntos(doc, 'a) Satisfacción con la Formación CFTe-AP', col_sat1, df_actual, stats['total_actual'])
    escribir_bloque_puntos(doc, 'b) Nivel de satisfacción en el trabajo actual', col_sat2, df_actual, stats['total_actual'])

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# --- 8. ZONA DE DESCARGA Y REPORTES ---
st.markdown("---")
st.subheader("📥 Exportación de Evidencias")
col_btn1, col_btn2 = st.columns(2)

with col_btn1:
    csv_data = df.to_csv(index=False).encode('utf-8-sig')
    st.download_button(
        label="📊 Descargar Base de Datos Filtrada (CSV)", 
        data=csv_data, 
        file_name="Evidencia_Titulados_Filtrada.csv", 
        mime="text/csv", 
        use_container_width=True
    )

with col_btn2:
    estadisticas = {
        'historico': universo_absoluto if 'universo_absoluto' in locals() else 0,
        'universo': total_universo if 'total_universo' in locals() else 0,
        'respuestas': total_respuestas_unicas if 'total_respuestas_unicas' in locals() else 0,
        'tasa': tasa if 'tasa' in locals() else 0,
        'total_global': total_respuestas,
        'total_actual': len(df[df['Origen_Encuesta'] == 'Actual'])
    }
    
    archivo_word = generar_word_institucional(df, estadisticas)
    
    st.download_button(
        label="📝 Descargar Informe de Análisis Específico",
        data=archivo_word,
        file_name=f"Informe_Detallado_Titulados_{time.strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )