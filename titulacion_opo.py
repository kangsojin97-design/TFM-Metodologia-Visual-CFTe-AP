import streamlit as st
import pandas as pd
import numpy as np
import os
import plotly.express as px
from io import BytesIO

# Intentar importar docx para la exportación a Word
try:
    from docx import Document
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

# --- 1. CONFIGURACIÓN E IDENTIDAD INSTITUCIONAL ---
st.set_page_config(page_title="Análisis Institucional CFTe-AP", layout="wide")

COLORES_CLAROS = ['#AED6F1', '#A9DFBF', '#F9E79F', '#F5CBA7', '#E5E7E9']
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

CARRERAS_MAP = {
    "180": "TNS EN ASISTENCIA JURÍDICA",
    "184": "TNS EN PROYECTOS ELÉCTRICOS DE DISTRIBUCIÓN",
    "201": "TNS EN ENFERMERÍA CON MENCIÓN EN GERONTOLOGÍA",
    "202": "TNS EN INFORMÁTICA Y APLICACIONES TECNOLÓGICAS",
    "204": "TNS EN GESTIÓN DE COMERCIO EXTERIOR",
    "283": "TNS EN EDUCACIÓN PARVULARIA Y PRIMER - SEGUNDO AÑO DE EDUCACIÓN BÁSICA",
    "380": "TNS EN DEPORTES Y RECREACIÓN",
    "554": "TNS EN ADMINISTRACIÓN PÚBLICA",
    "559": "TNS EN AGRÍCOLA",
    "602": "TNS EN EDUCACIÓN ESPECIAL",
    "603": "TNS EN ADMINISTRACIÓN DE EMPRESAS",
    "608": "TNS EN TRABAJO SOCIAL",
    "609": "TNS EN GEOLOGÍA",
    "613": "TNS EN CONTROL DE GESTIÓN Y LOGÍSTICA",
    "615": "TNS EN FABRICACIÓN Y MONTAJE DE ESTRUCTURAS METÁLICAS",
    "187": "TNS EN MANTENIMIENTO ELECTROMECÁNICO DE EQUIPOS MÓVILES",
    "206": "TNS EN GESTIÓN CONTABLE",
    "207": "TNS EN GESTIÓN DE RECURSOS HUMANOS",
    "612": "TNS EN LABORATORIO CLÍNICO, BANCO DE SANGRE E IMAGENOLOGÍA",
    "153": "TNS EN OBRAS CIVILES",
    "188": "TNS EN MANTENIMIENTO ELECTROMECANICO CON MENCION EN ELECTROMOVILIDAD",
    "208": "TNS EN ADMINISTRACION PUBLICA CON MENCION EN GESTION JURIDICA",
    "209": "TNS EN ESTETICA CON MENCION EN MASOTERAPIA",
    "210": "TNS EN VETERINARIA",
    "284": "TNS EN EDUCACION PARVULARIA PRIMER CICLO DE EDUCACION BASICA",
    "219": "TNS EN ESTÉTICA Y COSMETOLOGIA MENCIÓN MASOTERAPIA"
}

st.title("🏛️ Dimensión: Docencia - Análisis de Finalización y Titulación")
st.markdown("---")

# --- 2. RUTAS ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CARPETA_VIN = os.path.join(BASE_DIR, "2_Datos_Procesados", "2_Datos_Procesados_Vinculacion")
CARPETA_DOC = os.path.join(BASE_DIR, "2_Datos_Procesados", "2_Datos_Procesados_Docencia")

# --- 3. FUNCIONES UTILITARIAS Y LECTOR ---
@st.cache_data
def cargar_archivo(ruta):
    try:
        df = None
        if ruta.endswith(".xlsx"): df = pd.read_excel(ruta)
        else:
            for s in [';', ',']:
                for enc in ['utf-8-sig', 'latin-1']:
                    try:
                        df = pd.read_csv(ruta, sep=s, encoding=enc, on_bad_lines='skip', engine='python')
                        break
                    except: continue
                if df is not None: break
        
        if df is not None:
            df.columns = df.columns.str.strip().str.upper()
            
            # Limpieza RUT
            col_rut = next((c for c in df.columns if 'RUT' in c or 'RUN' in c), None)
            if col_rut:
                df.rename(columns={col_rut: 'RUT'}, inplace=True)
                df['RUT'] = df['RUT'].astype(str).str.replace(r'[\.\-]', '', regex=True).str.upper().str.strip()
            
            # Limpieza Cohorte
            if 'COHORTE' in df.columns:
                df['ANIO_COHORTE'] = pd.to_numeric(df['COHORTE'], errors='coerce')

            # Limpieza Carrera (Asegurando COD_LIMP)
            col_carrera = next((c for c in ['CARRERA_COD', 'COD_CARRERA', 'CARRERA'] if c in df.columns), None)
            if col_carrera:
                df['COD_LIMP'] = df[col_carrera].astype(str).str.extract(r'(\d{3})', expand=False)
                df['NOMBRE_CARRERA'] = df['COD_LIMP'].map(CARRERAS_MAP).fillna("CARRERA DESCONOCIDA")

            # Fechas
            variaciones_ingreso = ['AÑO INGRESO', 'AÑO DE INGRESO', 'ANIO INGRESO', 'FECHA INGRESO', 'FECHA_INGRESO', 'INGRESO']
            col_ing = next((c for c in variaciones_ingreso if c in df.columns), None)
            if col_ing:
                df.rename(columns={col_ing: 'FECHA_INGRESO_CALC'}, inplace=True)

            if 'TITULACIÓN FECHA' in df.columns:
                df['TITULACIÓN FECHA'] = pd.to_datetime(df['TITULACIÓN FECHA'], errors='coerce')
                
        return df
    except: return None

def evaluar_op_interna(row):
    ing, tit, cod = row.get('ANIO_COHORTE'), row.get('TITULACIÓN FECHA'), str(row.get('COD_LIMP', ''))
    if pd.isna(ing) or pd.isna(tit): return False
    ing_dt = pd.to_datetime(f"{int(ing)}-03-01")
    anios = (tit - ing_dt).days / 365.25
    limite = 4.0 if cod == '184' else 3.5
    return anios <= limite

def agregar_df_a_word(doc, df):
    df_export = df.reset_index() if df.index.name is not None else df.copy()
    table = doc.add_table(rows=1, cols=len(df_export.columns))
    table.style = 'Table Grid'
    for i, column in enumerate(df_export.columns):
        table.rows[0].cells[i].text = str(column).replace('_', ' ')
    for index, row in df_export.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            if pd.isna(value): row_cells[i].text = "0"
            elif isinstance(value, float): row_cells[i].text = str(round(value, 1))
            else: row_cells[i].text = str(value)

def generar_reporte_completo(df_ret, df_op, df_tut, df_conv, sel_anios, sel_carreras):
    if not DOCX_DISPONIBLE: return None
    doc = Document()
    doc.add_heading('Informe Integral de Titulación y Servicios CFTe-AP', level=0)
    doc.add_paragraph(f"Cohortes analizadas: {', '.join(map(str, sel_anios))}")
    
    doc.add_heading('1. Análisis de Titulación sobre Retención', level=1)
    if df_ret is not None and not df_ret.empty: agregar_df_a_word(doc, df_ret)
    else: doc.add_paragraph("Sin datos.")
    
    doc.add_heading('2. Análisis de Titulación Oportuna (Global)', level=1)
    if df_op is not None and not df_op.empty: agregar_df_a_word(doc, df_op)
    else: doc.add_paragraph("Sin datos.")
        
    doc.add_heading('3. Efectividad de Servicios Estudiantiles', level=1)
    
    doc.add_heading('3.1 Tutorías Académicas', level=2)
    if df_tut is not None and not df_tut.empty:
        for anio in sorted(df_tut['ANIO_COHORTE'].unique()):
            doc.add_heading(f'Cohorte {int(anio)}', level=3)
            df_anio_tut = df_tut[df_tut['ANIO_COHORTE'] == anio].drop(columns=['ANIO_COHORTE'])
            agregar_df_a_word(doc, df_anio_tut)
    else: doc.add_paragraph("No hay registros de tutorías.")
    
    doc.add_heading('3.2 Beneficio de Acortamiento de Estudios', level=2)
    if df_conv is not None and not df_conv.empty:
        for anio in sorted(df_conv['ANIO_COHORTE'].unique()):
            doc.add_heading(f'Cohorte {int(anio)}', level=3)
            df_anio_conv = df_conv[df_conv['ANIO_COHORTE'] == anio].drop(columns=['ANIO_COHORTE'])
            agregar_df_a_word(doc, df_anio_conv)
    else: doc.add_paragraph("No hay registros de acortamiento de estudios.")
    
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# --- 4. CARGA DE DATOS ---
df_titulados = cargar_archivo(os.path.join(CARPETA_VIN, "LIMPIO_Nomina_Titulados.xlsx"))
df_convalidaciones = cargar_archivo(os.path.join(CARPETA_DOC, "LIMPIO_Convalidaciones.xlsx"))

if df_titulados is not None: df_titulados = df_titulados[df_titulados['ANIO_COHORTE'] >= 2021]
if df_convalidaciones is not None: df_convalidaciones = df_convalidaciones[df_convalidaciones['ANIO_COHORTE'] >= 2021]

df_retenidos_list = []
if os.path.exists(CARPETA_DOC):
    archivos_doc = [f for f in os.listdir(CARPETA_DOC) if f.startswith("LIMPIO_ING")]
    anios_cohorte_doc = sorted({int(f.split('_')[1].replace('ING', '')) for f in archivos_doc if 'ING' in f and int(f.split('_')[1].replace('ING', '')) >= 2021})
    for anio in anios_cohorte_doc:
        f_ing = next((f for f in archivos_doc if f.startswith(f"LIMPIO_ING{anio}_MAT{anio}")), None)
        f_ret = next((f for f in archivos_doc if f.startswith(f"LIMPIO_ING{anio}_MAT{anio+1}")), None)
        if f_ing and f_ret:
            d_i, d_r = cargar_archivo(os.path.join(CARPETA_DOC, f_ing)), cargar_archivo(os.path.join(CARPETA_DOC, f_ret))
            if d_i is not None and d_r is not None:
                ret = d_i[d_i['RUT'].isin(d_r['RUT'])].copy()
                ret['ANIO_COHORTE'] = anio
                df_retenidos_list.append(ret)

df_retenidos_total = pd.concat(df_retenidos_list, ignore_index=True) if df_retenidos_list else None

if df_retenidos_total is None or df_titulados is None:
    st.error("❌ Error al cargar datos base. Verifique los archivos.")
    st.stop()

# --- 5. FILTROS GLOBALES ---
st.sidebar.header("🔍 Filtros Globales")
carreras_disp = sorted(CARRERAS_MAP.values())
sel_carreras = st.sidebar.multiselect("Carreras", carreras_disp, default=carreras_disp)

anios_disp = sorted(df_retenidos_total['ANIO_COHORTE'].dropna().unique(), reverse=True)
sel_anios = st.sidebar.multiselect("Año de Cohorte", anios_disp, default=anios_disp)

df_ret_f = df_retenidos_total[(df_retenidos_total['NOMBRE_CARRERA'].isin(sel_carreras)) & (df_retenidos_total['ANIO_COHORTE'].isin(sel_anios))]
df_tit_f = df_titulados[(df_titulados['NOMBRE_CARRERA'].isin(sel_carreras)) & (df_titulados['ANIO_COHORTE'].isin(sel_anios))]

tabla_retencion = pd.DataFrame()
tabla_oportunidad = pd.DataFrame()
res_tut = pd.DataFrame()
res_conv = pd.DataFrame()

# --- 6. PESTAÑAS ---
tab1, tab2, tab3 = st.tabs(["🎓 Titulación (Retención)", "🕒 Oportunidad", "📊 Efectividad Servicios"])

with tab1:
    st.subheader("Análisis de Titulación sobre Retención")
    res_ret = df_ret_f.groupby('ANIO_COHORTE').size().reset_index(name='TOTAL RETENIDOS')
    res_tit = df_tit_f.groupby('ANIO_COHORTE').size().reset_index(name='TOTAL TITULADOS')
    tabla_retencion = pd.merge(res_ret, res_tit, on='ANIO_COHORTE', how='left').fillna(0)
    tabla_retencion['TASA TITULACIÓN (%)'] = np.where(tabla_retencion['TOTAL RETENIDOS'] > 0, (tabla_retencion['TOTAL TITULADOS'] / tabla_retencion['TOTAL RETENIDOS'] * 100), 0).round(1)
    
    c1, c2, c3 = st.columns(3)
    c1.metric("Total Retenidos", int(tabla_retencion['TOTAL RETENIDOS'].sum()))
    c2.metric("Total Titulados", int(tabla_retencion['TOTAL TITULADOS'].sum()))
    tasa_g = round((tabla_retencion['TOTAL TITULADOS'].sum() / tabla_retencion['TOTAL RETENIDOS'].sum() * 100), 1) if tabla_retencion['TOTAL RETENIDOS'].sum() > 0 else 0
    c3.metric("Tasa Global", f"{tasa_g}%")
    
    st.dataframe(tabla_retencion.set_index('ANIO_COHORTE'), use_container_width=True)
    st.plotly_chart(px.bar(tabla_retencion, x='ANIO_COHORTE', y='TASA TITULACIÓN (%)', text='TASA TITULACIÓN (%)', title="Tasa de Titulación por Año de Cohorte", color_discrete_sequence=['#A9DFBF']), use_container_width=True)

with tab2:
    st.subheader("Análisis de Oportunidad de Titulación (Global)")
    df_tit_f_op = df_tit_f.copy()
    df_tit_f_op['ES_OPORTUNO'] = df_tit_f_op.apply(evaluar_op_interna, axis=1)
    
    if not df_tit_f_op.empty:
        tabla_oportunidad = df_tit_f_op.groupby('ANIO_COHORTE').agg(
            TOTAL_TITULADOS=('RUT', 'nunique'),
            OPORTUNOS=('ES_OPORTUNO', 'sum')
        ).reset_index()
        tabla_oportunidad['NO OPORTUNOS'] = tabla_oportunidad['TOTAL_TITULADOS'] - tabla_oportunidad['OPORTUNOS']
        tabla_oportunidad['TASA OPORTUNIDAD (%)'] = np.where(tabla_oportunidad['TOTAL_TITULADOS'] > 0, (tabla_oportunidad['OPORTUNOS'] / tabla_oportunidad['TOTAL_TITULADOS'] * 100), 0).round(1)
        
        m1, m2, m3 = st.columns(3)
        tot_tit = int(tabla_oportunidad['TOTAL_TITULADOS'].sum())
        tot_op = int(tabla_oportunidad['OPORTUNOS'].sum())
        m1.metric("Titulados Analizados", tot_tit)
        m2.metric("Titulados Oportunos", tot_op)
        t_global_op = round((tot_op / tot_tit * 100), 1) if tot_tit > 0 else 0
        m3.metric("Tasa Oportunidad Global", f"{t_global_op}%")
        
        st.dataframe(tabla_oportunidad.set_index('ANIO_COHORTE'), use_container_width=True)
        
        df_melt_op = tabla_oportunidad.melt(id_vars='ANIO_COHORTE', value_vars=['OPORTUNOS', 'NO OPORTUNOS'], var_name='ESTADO', value_name='CANTIDAD')
        st.plotly_chart(px.bar(df_melt_op, x='ANIO_COHORTE', y='CANTIDAD', color='ESTADO', barmode='group', text_auto=True, title="Distribución de Titulación Oportuna", color_discrete_sequence=COLORES_CLAROS), use_container_width=True)

with tab3:
    st.subheader("📊 Efectividad de Servicios y Oportunidad")

    if df_convalidaciones is not None:
        # FILTRO PURO: Directamente desde el archivo de convalidaciones
        df_conv_f = df_convalidaciones[
            (df_convalidaciones['NOMBRE_CARRERA'].isin(sel_carreras)) & 
            (df_convalidaciones['ANIO_COHORTE'].isin(sel_anios))
        ].copy()
        
        # CRUCE CON TITULADOS: Solo cruzamos con la nómina para traer fechas de titulación
        df_tit_unica = df_tit_f[['RUT', 'TITULACIÓN FECHA', 'COD_LIMP']].drop_duplicates(subset=['RUT'])
        df_cruce = pd.merge(df_conv_f, df_tit_unica, on='RUT', how='left')
        df_cruce['ES_OPORTUNO'] = df_cruce.apply(evaluar_op_interna, axis=1)
        
        col_tipo = 'UNNAMED: 6' if 'UNNAMED: 6' in df_cruce.columns else next((c for c in df_cruce.columns if any(x in c for x in ['TIPO', 'MOTIVO', 'DESCRIPCIÓN', 'OBSERVACIÓN'])), df_cruce.columns[0])

        # --- 1. TUTORÍAS ---
        st.markdown("### 📘 1. Apoyo Estudiantil: Tutorías Académicas")
        df_tut_raw = df_cruce[df_cruce[col_tipo].astype(str).str.contains('TUTOR', case=False, na=False)].copy()
        
        if not df_tut_raw.empty:
            # REGLA CRÍTICA: Eliminar duplicados para que un RUT se cuente solo una vez por cohorte
            df_tut_unicos = df_tut_raw.drop_duplicates(subset=['RUT', 'ANIO_COHORTE'])
            
            res_tut = df_tut_unicos.groupby('ANIO_COHORTE').agg(
                ESTUDIANTES_CON_SERVICIO=('RUT', 'nunique'),
                TITULADOS=('TITULACIÓN FECHA', lambda x: x.notna().sum()),
                OPORTUNOS=('ES_OPORTUNO', 'sum')
            ).reset_index()

            res_tut['TASA TITULACIÓN (%)'] = np.where(res_tut['ESTUDIANTES_CON_SERVICIO'] > 0, (res_tut['TITULADOS'] / res_tut['ESTUDIANTES_CON_SERVICIO'] * 100), 0).round(1)
            # Tasa Oportuna estrictamente sobre los Titulados
            res_tut['TASA OPORTUNIDAD (%)'] = np.where(res_tut['TITULADOS'] > 0, (res_tut['OPORTUNOS'] / res_tut['TITULADOS'] * 100), 0).round(1)

            t1, t2, t3 = st.columns(3)
            t1.metric("Estudiantes con Tutoría", int(res_tut['ESTUDIANTES_CON_SERVICIO'].sum()))
            t2.metric("Lograron Titularse", int(res_tut['TITULADOS'].sum()))
            t3.metric("Titulados Oportunamente", int(res_tut['OPORTUNOS'].sum()))

            tab_t_1, tab_t_2 = st.tabs(["A) Detalle General", "B) Detalle Oportunidad"])
            with tab_t_1:
                st.dataframe(res_tut[['ANIO_COHORTE', 'ESTUDIANTES_CON_SERVICIO', 'TITULADOS', 'TASA TITULACIÓN (%)']].set_index('ANIO_COHORTE'), use_container_width=True)
            with tab_t_2:
                st.dataframe(res_tut[['ANIO_COHORTE', 'TITULADOS', 'OPORTUNOS', 'TASA OPORTUNIDAD (%)']].set_index('ANIO_COHORTE'), use_container_width=True)
            
            fig_tut = px.bar(res_tut, x='ANIO_COHORTE', y=['ESTUDIANTES_CON_SERVICIO', 'TITULADOS', 'OPORTUNOS'], barmode='group', text_auto=True, title="Cantidades Absolutas (Tutorías)", color_discrete_sequence=['#E5E7E9', '#AED6F1', '#A9DFBF'])
            st.plotly_chart(fig_tut, use_container_width=True)
        else: st.info("No hay registros de tutorías para las opciones seleccionadas.")

        st.markdown("---")

        # --- 2. CONVALIDACIONES ---
        st.markdown("### 🎓 2. Beneficio de Acortamiento de Estudios")
        def clasificar_conv(val):
            v = str(val).upper().strip()
            if any(x in v for x in ["CONVALIDACIÓN EXTERNA", "CONVALIDACIÓN UTA", "CONVALIDACION EXTERNA", "CONVALIDACION UTA"]): return "CONVALIDACIÓN"
            if any(x in v for x in ["HOMOLOGACIÓN", "CONVALIDACIÓN INTERNA", "CONVALIDACION INTERNA", "HOMOLOGACION"]): return "HOMOLOGACIÓN"
            if "ARTICULACIÓN" in v or "ARTICULACION" in v: return "ARTICULACIÓN"
            if "R.A.P" in v: return "R.A.P"
            return "EXCLUIR"

        df_cruce['CATEGORIA'] = df_cruce[col_tipo].apply(clasificar_conv)
        df_conv_raw = df_cruce[df_cruce['CATEGORIA'] != "EXCLUIR"].copy()

        if not df_conv_raw.empty:
            # REGLA CRÍTICA: Eliminar duplicados para que un RUT se cuente solo una vez por categoría
            df_conv_unicos = df_conv_raw.drop_duplicates(subset=['RUT', 'ANIO_COHORTE', 'CATEGORIA'])
            
            res_conv = df_conv_unicos.groupby(['ANIO_COHORTE', 'CATEGORIA']).agg(
                ESTUDIANTES_CON_SERVICIO=('RUT', 'nunique'),
                TITULADOS=('TITULACIÓN FECHA', lambda x: x.notna().sum()),
                OPORTUNOS=('ES_OPORTUNO', 'sum')
            ).reset_index()

            res_conv['TASA TITULACIÓN (%)'] = np.where(res_conv['ESTUDIANTES_CON_SERVICIO'] > 0, (res_conv['TITULADOS'] / res_conv['ESTUDIANTES_CON_SERVICIO'] * 100), 0).round(1)
            # Tasa Oportuna estrictamente sobre los Titulados
            res_conv['TASA OPORTUNIDAD (%)'] = np.where(res_conv['TITULADOS'] > 0, (res_conv['OPORTUNOS'] / res_conv['TITULADOS'] * 100), 0).round(1)

            c1, c2, c3 = st.columns(3)
            c1.metric("Estudiantes con Acortamiento", int(res_conv['ESTUDIANTES_CON_SERVICIO'].sum()))
            c2.metric("Lograron Titularse", int(res_conv['TITULADOS'].sum()))
            c3.metric("Titulados Oportunamente", int(res_conv['OPORTUNOS'].sum()))

            tab_c_1, tab_c_2, tab_c_3 = st.tabs(["A) Tasas de Titulación", "B) Tasas de Oportunidad", "C) Totales Absolutos"])
            with tab_c_1:
                st.dataframe(res_conv.pivot(index='ANIO_COHORTE', columns='CATEGORIA', values='TASA TITULACIÓN (%)').fillna(0), use_container_width=True)
            with tab_c_2:
                st.dataframe(res_conv.pivot(index='ANIO_COHORTE', columns='CATEGORIA', values='TASA OPORTUNIDAD (%)').fillna(0), use_container_width=True)
            with tab_c_3:
                st.dataframe(res_conv[['ANIO_COHORTE', 'CATEGORIA', 'ESTUDIANTES_CON_SERVICIO', 'TITULADOS', 'OPORTUNOS']].set_index(['ANIO_COHORTE', 'CATEGORIA']), use_container_width=True)

            fig_conv_tot = px.bar(res_conv, x='ANIO_COHORTE', y='ESTUDIANTES_CON_SERVICIO', color='CATEGORIA', text='ESTUDIANTES_CON_SERVICIO', barmode='group', title="Cantidades Absolutas por Categoría", color_discrete_sequence=COLORES_CLAROS)
            st.plotly_chart(fig_conv_tot, use_container_width=True)
        else: st.info("No hay registros válidos de Acortamiento de Estudios.")
    else: st.warning("⚠️ No se encontró el archivo 'LIMPIO_Convalidaciones.xlsx'.")

# --- 7. EXPORTACIÓN DEL REPORTE MAESTRO ---
st.sidebar.markdown("---")
st.sidebar.header("📥 Exportación General")
if DOCX_DISPONIBLE:
    docx_file = generar_reporte_completo(tabla_retencion, tabla_oportunidad, res_tut, res_conv, sel_anios, sel_carreras)
    if docx_file:
        st.sidebar.download_button(
            label="📄 Descargar Reporte Integral (Word)", 
            data=docx_file, 
            file_name="Reporte_Titulacion_Integral.docx", 
            mime=DOCX_MIME
        )
else:
    st.sidebar.error("Librería 'python-docx' no instalada.")